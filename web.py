import gzip
import json
import hmac
import os
import hashlib
import secrets
import threading
import logging
import time as _time
from http.server import HTTPServer, BaseHTTPRequestHandler
from http.cookies import SimpleCookie
from urllib.parse import urlparse, parse_qs

from storage.database import get_connection, _is_postgres, get_unprocessed_news, update_news_status, init_db

logger = logging.getLogger(__name__)
PORT = int(os.getenv("PORT", 8080))

# Pre-compressed dashboard HTML (computed once at module load, see bottom)
_DASHBOARD_HTML_GZIP = None
_DASHBOARD_HTML_BYTES = None

# Secret key for signing cookies — stable across redeploys via env var
_COOKIE_SECRET = os.getenv("COOKIE_SECRET", "igronews-default-secret-key-2024")

# Users: {username: {"hash": password_hash, "role": "admin"|"editor"|"viewer"}}
USERS = {
    "admin": {"hash": hashlib.sha256("admin123".encode()).hexdigest(), "role": "admin"},
}

# Role permissions
ROLE_PERMISSIONS = {
    "admin": {"read", "write", "approve", "settings", "users", "flags", "delete", "pipeline"},
    "editor": {"read", "write", "approve", "pipeline"},
    "viewer": {"read"},
}

def _user_role(username: str) -> str:
    entry = USERS.get(username, {})
    if isinstance(entry, dict):
        return entry.get("role", "editor")
    return "editor"  # legacy compat

def _user_has_perm(username: str, perm: str) -> bool:
    role = _user_role(username)
    return perm in ROLE_PERMISSIONS.get(role, set())


def _sign_cookie(username: str) -> str:
    """Создаёт подписанную куку: username.expiry.signature"""
    expiry = int(_time.time()) + 86400 * 7  # 7 дней
    payload = f"{username}:{expiry}"
    sig = hmac.new(_COOKIE_SECRET.encode(), payload.encode(), hashlib.sha256).hexdigest()[:24]
    return f"{payload}:{sig}"


def _verify_cookie(value: str) -> str | None:
    """Проверяет подписанную куку. Возвращает username или None."""
    try:
        parts = value.split(":")
        if len(parts) != 3:
            return None
        username, expiry_str, sig = parts
        expiry = int(expiry_str)
        if _time.time() > expiry:
            return None
        payload = f"{username}:{expiry_str}"
        expected = hmac.new(_COOKIE_SECRET.encode(), payload.encode(), hashlib.sha256).hexdigest()[:24]
        if not hmac.compare_digest(sig, expected):
            return None
        if username not in USERS:
            return None
        return username
    except Exception:
        return None


class AdminHandler(BaseHTTPRequestHandler):

    def _get_session_user(self):
        cookie_header = self.headers.get("Cookie", "")
        cookie = SimpleCookie(cookie_header)
        token = cookie.get("session")
        if token:
            return _verify_cookie(token.value)
        return None

    def _require_auth(self):
        """Returns True if authorized, False if redirected to login."""
        if self._get_session_user():
            return True
        path = urlparse(self.path).path
        if path == "/login":
            return True
        if path.startswith("/api/"):
            self._json({"error": "unauthorized"}, 401)
        else:
            self.send_response(302)
            self.send_header("Location", "/login")
            self.end_headers()
        return False

    def do_GET(self):
        path = urlparse(self.path).path
        # Static files that don't require auth
        if path == "/robots.txt":
            self.send_response(200)
            self.send_header("Content-Type", "text/plain")
            self.end_headers()
            self.wfile.write(b"User-agent: *\nDisallow: /\n")
            return
        if path == "/favicon.ico":
            self.send_response(204)
            self.end_headers()
            return
        if path == "/login":
            self._serve_login()
            return
        if path == "/logout":
            self._do_logout()
            return
        if path == "/api/diag":
            self._serve_diag()
            return
        if not self._require_auth():
            return

        routes = {
            "/": self._serve_dashboard,
            "/api/stats": lambda: self._json(self._get_stats()),
            "/api/pipeline/status": lambda: self._json(self._get_pipeline_status()),
            "/api/news": lambda: self._json(self._get_news()),
            "/api/sources": lambda: self._json(self._get_sources()),
            "/api/prompts": lambda: self._json(self._get_prompts()),
            "/api/settings": lambda: self._json(self._get_settings()),
            "/api/users": lambda: self._json(self._get_users()),
            "/api/health": lambda: self._json(self._get_health()),
            "/api/dashboard_groups": lambda: self._dashboard_groups(),
            "/api/sources_stats": lambda: self._json(self._get_sources_stats()),
            "/api/db_info": lambda: self._json(self._get_db_info()),
            "/api/articles": lambda: self._json(self._get_articles()),
            "/api/queue": lambda: self._json(self._get_queue()),
            "/api/analytics": lambda: self._json(self._get_analytics()),
            "/api/prompt_versions": lambda: self._json(self._get_prompt_versions()),
            "/api/viral": lambda: self._json(self._get_viral()),
            "/api/logs": lambda: self._json(self._get_logs()),
            "/api/rate_stats": lambda: self._json(self._get_rate_stats()),
            "/api/cache_stats": lambda: self._json(self._get_cache_stats()),
            "/api/editorial": lambda: self._json(self._get_editorial()),
            "/api/moderation_list": lambda: self._json(self._get_moderation_list()),
            "/api/digests": lambda: self._json(self._get_digests()),
            "/api/viral_triggers": lambda: self._json(self._get_viral_triggers()),
            "/api/final": lambda: self._json(self._get_final()),
            # Phase 0: feature flags, observability, dashboard v2
            "/api/feature_flags": lambda: self._json(self._get_feature_flags()),
            "/api/cost_summary": lambda: self._json(self._get_cost_summary()),
            "/api/config_audit": lambda: self._json(self._get_config_audit()),
            "/api/ops_dashboard": lambda: self._json(self._get_ops_dashboard()),
            # Phase 3: analytics funnel, source intelligence
            "/api/analytics/funnel": lambda: self._json(self._get_funnel_analytics()),
            "/api/analytics/cost_by_source": lambda: self._json(self._get_cost_by_source()),
            "/api/analytics/prompt_insights": lambda: self._json(self._get_prompt_insights()),
            # Phase 4: storylines, source health plus, threshold simulator
            "/api/storylines": lambda: self._json(self._get_storylines()),
            "/api/source_health_plus": lambda: self._json(self._get_source_health_plus()),
        }

        # DOCX download (GET with query param)
        if path == "/api/articles/docx":
            from urllib.parse import parse_qs
            qs = parse_qs(urlparse(self.path).query)
            aid = qs.get("id", [""])[0]
            if aid:
                self._serve_docx(aid)
            else:
                self._json({"error": "id required"}, 400)
            return

        # Bulk DOCX (ZIP with multiple DOCX)
        if path == "/api/articles/docx_bulk":
            from urllib.parse import parse_qs
            qs = parse_qs(urlparse(self.path).query)
            ids = qs.get("ids", [""])[0]
            if ids:
                self._serve_docx_bulk(ids.split(","))
            else:
                self._json({"error": "ids required"}, 400)
            return

        # Event chain (temporal clusters)
        if path == "/api/event_chain":
            qs = parse_qs(urlparse(self.path).query)
            news_id = qs.get("news_id", [""])[0]
            if news_id:
                self._json(self._get_event_chain_by_id(news_id))
            else:
                self._json({"error": "news_id required"}, 400)
            return

        handler = routes.get(path)
        if handler:
            handler()
        else:
            self._json({"error": "not found"}, 404)

    def do_POST(self):
        path = urlparse(self.path).path
        body = self._read_body()

        if path == "/api/login":
            self._do_login(body)
            return
        if not self._require_auth():
            return

        routes = {
            "/api/process": self._run_process,
            "/api/process_one": lambda: self._process_one(body),
            "/api/export_sheets": lambda: self._export_sheets(body),
            "/api/sources/add": lambda: self._add_source(body),
            "/api/sources/edit": lambda: self._edit_source(body),
            "/api/sources/delete": lambda: self._delete_source(body),
            "/api/prompts/save": lambda: self._save_prompts(body),
            "/api/settings/save": lambda: self._save_settings(body),
            "/api/test_llm": lambda: self._test_llm(body),
            "/api/test_keyso": lambda: self._test_keyso(body),
            "/api/reparse": lambda: self._reparse_source(body),
            "/api/test_sheets": lambda: self._test_sheets(body),
            "/api/quick_tags": lambda: self._quick_tags(body),
            "/api/review": lambda: self._run_review(body),
            "/api/review_batch": lambda: self._review_batch(body),
            "/api/approve": lambda: self._approve_news(body),
            "/api/reject": lambda: self._reject_news(body),
            "/api/users/add": lambda: self._add_user(body),
            "/api/users/delete": lambda: self._delete_user(body),
            "/api/users/change_password": lambda: self._change_password(body),
            "/api/news/bulk_status": lambda: self._bulk_status(body),
            "/api/news/delete": lambda: self._delete_news(body),
            "/api/test_parse": lambda: self._test_parse(body),
            "/api/setup_headers": lambda: self._setup_headers(body),
            "/api/reparse_all": lambda: self._reparse_all(body),
            "/api/rewrite": lambda: self._rewrite_news(body),
            "/api/merge": lambda: self._merge_news(body),
            "/api/news/detail": lambda: self._news_detail(body),
            "/api/export_sheets_bulk": lambda: self._export_sheets_bulk(body),
            "/api/analyze_news": lambda: self._analyze_news(body),
            "/api/batch_rewrite": lambda: self._batch_rewrite(body),
            "/api/articles/save": lambda: self._save_article(body),
            "/api/articles/update": lambda: self._update_article(body),
            "/api/articles/delete": lambda: self._delete_article(body),
            "/api/articles/rewrite": lambda: self._rewrite_article(body),
            "/api/articles/improve": lambda: self._improve_article(body),
            "/api/articles/detail": lambda: self._article_detail(body),
            "/api/articles/schedule": lambda: self._schedule_article(body),
            "/api/prompt_versions/save": lambda: self._save_prompt_version(body),
            "/api/prompt_versions/activate": lambda: self._activate_prompt_version(body),
            "/api/generate_digest": lambda: self._generate_digest(body),
            "/api/digest/generate": lambda: self._generate_and_save_digest(body),
            "/api/event_chain": lambda: self._get_event_chain(body),
            "/api/queue/cancel": lambda: self._cancel_queue_task(body),
            "/api/queue/cancel_all": lambda: self._cancel_all_queue(body),
            "/api/queue/clear_done": lambda: self._clear_done_queue(body),
            "/api/queue/retry": lambda: self._retry_queue_tasks(body),
            "/api/viral_triggers/save": lambda: self._save_viral_trigger(body),
            "/api/viral_triggers/delete": lambda: self._delete_viral_trigger(body),
            "/api/run_auto_review": lambda: self._run_auto_review(body),
            "/api/queue/rewrite": lambda: self._queue_batch_rewrite(body),
            "/api/queue/sheets": lambda: self._queue_sheets_export(body),
            "/api/translate_title": lambda: self._translate_title(body),
            "/api/ai_recommend": lambda: self._ai_recommend(body),
            "/api/cache/clear": lambda: self._clear_cache(body),
            "/api/export_all_processed": lambda: self._export_all_processed(body),
            "/api/export_ready_all": lambda: self._export_ready_all(body),
            "/api/pipeline/full_auto": lambda: self._pipeline_full_auto(body),
            "/api/pipeline/no_llm": lambda: self._pipeline_no_llm(body),
            "/api/pipeline/stop": lambda: self._pipeline_stop(body),
            "/api/moderation": lambda: self._get_moderation(body),
            "/api/moderation/rewrite": lambda: self._moderation_rewrite(body),
            "/api/seo_check": lambda: self._seo_check(body),
            # Phase 0: feature flags, decision trace
            "/api/feature_flags/toggle": lambda: self._toggle_feature_flag(body),
            "/api/decision_trace": lambda: self._get_decision_trace(body),
            # Phase 2: content versioning, multi-output
            "/api/articles/versions": lambda: self._get_article_versions(body),
            "/api/articles/multi_output": lambda: self._generate_multi_output(body),
            "/api/articles/regenerate_field": lambda: self._regenerate_field(body),
            # Phase 4: threshold simulator, rescore
            "/api/simulate_thresholds": lambda: self._simulate_thresholds(body),
            "/api/rescore": lambda: self._rescore_news(body),
            "/api/health/heal": lambda: self._heal_source(body),
        }
        handler = routes.get(path)
        if handler:
            handler()
        else:
            self._json({"error": "not found"}, 404)

    # --- Helpers ---
    def _read_body(self):
        length = int(self.headers.get("Content-Length", 0))
        if length == 0:
            return {}
        try:
            return json.loads(self.rfile.read(length))
        except Exception:
            return {}

    def _accepts_gzip(self):
        return "gzip" in self.headers.get("Accept-Encoding", "")

    def _json(self, data, code=200):
        body = json.dumps(data, ensure_ascii=False, default=str).encode()
        self.send_response(code)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Cache-Control", "no-cache, no-store")
        # Gzip if body > 1KB and client supports it
        if len(body) > 1024 and self._accepts_gzip():
            body = gzip.compress(body, compresslevel=6)
            self.send_header("Content-Encoding", "gzip")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    # --- Auth ---
    def _do_login(self, body):
        username = body.get("username", "")
        password = body.get("password", "")
        pw_hash = hashlib.sha256(password.encode()).hexdigest()
        entry = USERS.get(username, {})
        stored_hash = entry.get("hash", entry) if isinstance(entry, dict) else entry
        if stored_hash == pw_hash:
            signed = _sign_cookie(username)
            self.send_response(200)
            self.send_header("Content-Type", "application/json")
            self.send_header("Set-Cookie", f"session={signed}; Path=/; HttpOnly; SameSite=Lax; Max-Age=604800")
            self.end_headers()
            self.wfile.write(json.dumps({"status": "ok", "role": _user_role(username)}).encode())
        else:
            self._json({"status": "error", "message": "Invalid credentials"}, 401)

    def _do_logout(self):
        self.send_response(302)
        self.send_header("Set-Cookie", "session=; Path=/; Max-Age=0")
        self.send_header("Location", "/login")
        self.end_headers()

    def _get_users(self):
        return [{"username": u, "role": _user_role(u)} for u in USERS.keys()]

    def _add_user(self, body):
        username = body.get("username", "")
        password = body.get("password", "")
        role = body.get("role", "editor")
        if role not in ROLE_PERMISSIONS:
            role = "editor"
        if not username or not password:
            self._json({"status": "error", "message": "Username and password required"})
            return
        user = self._get_session_user()
        if not _user_has_perm(user, "users"):
            self._json({"status": "error", "message": "Недостаточно прав"}, 403)
            return
        USERS[username] = {"hash": hashlib.sha256(password.encode()).hexdigest(), "role": role}
        self._json({"status": "ok", "users": self._get_users()})

    def _delete_user(self, body):
        username = body.get("username", "")
        if username == "admin":
            self._json({"status": "error", "message": "Cannot delete admin"})
            return
        user = self._get_session_user()
        if not _user_has_perm(user, "users"):
            self._json({"status": "error", "message": "Недостаточно прав"}, 403)
            return
        USERS.pop(username, None)
        self._json({"status": "ok", "users": self._get_users()})

    def _require_perm(self, perm: str) -> bool:
        """Check permission. Returns True if allowed, sends 403 and returns False if denied."""
        user = self._get_session_user()
        if not user:
            self._json({"error": "unauthorized"}, 401)
            return False
        if not _user_has_perm(user, perm):
            self._json({"error": "Недостаточно прав"}, 403)
            return False
        return True

    def _serve_diag(self):
        """Публичный диагностический endpoint — проверка БД и парсинга."""
        import json as _json
        import config
        diag = {}
        try:
            # DB connection
            db_url = config.DATABASE_URL
            diag["db_type"] = "PostgreSQL" if db_url.startswith("postgres") else "SQLite"
            diag["db_url_set"] = bool(db_url and db_url != "sqlite:///news.db")

            conn = get_connection()
            cur = conn.cursor()
            try:
                diag["db_connected"] = True

                # Counts
                cur.execute("SELECT COUNT(*) FROM news")
                diag["total_news"] = cur.fetchone()[0]

                cur.execute("SELECT status, COUNT(*) FROM news GROUP BY status")
                status_counts = {}
                for row in cur.fetchall():
                    if _is_postgres():
                        status_counts[row[0]] = row[1]
                    else:
                        status_counts[row[0]] = row[1]
                diag["by_status"] = status_counts

                cur.execute("SELECT COUNT(*) FROM news_analysis")
                diag["total_analyzed"] = cur.fetchone()[0]

                cur.execute("SELECT COUNT(*) FROM news_analysis WHERE reviewed_at IS NOT NULL AND reviewed_at != ''")
                diag["total_reviewed"] = cur.fetchone()[0]

                # Last parsed
                cur.execute("SELECT MAX(parsed_at) FROM news")
                row = cur.fetchone()
                diag["last_parsed"] = str(row[0]) if row and row[0] else "never"

                # Sources parsed
                cur.execute("SELECT source, COUNT(*) FROM news GROUP BY source ORDER BY COUNT(*) DESC")
                src_counts = {}
                for row in cur.fetchall():
                    if _is_postgres():
                        src_counts[row[0]] = row[1]
                    else:
                        src_counts[row[0]] = row[1]
                diag["sources"] = src_counts

                diag["configured_sources"] = len(config.SOURCES)

                # Scoring diagnostics: check recent news for missing data
                ph = "%s" if _is_postgres() else "?"
                cur.execute(f"""
                    SELECT n.id, n.title, n.source, n.status,
                           LENGTH(n.plain_text) as text_len,
                           COALESCE(a.total_score, -1) as score
                    FROM news n
                    LEFT JOIN news_analysis a ON n.id = a.news_id
                    ORDER BY n.parsed_at DESC
                    LIMIT 10
                """)
                if _is_postgres():
                    cols = [d[0] for d in cur.description]
                    recent = [dict(zip(cols, r)) for r in cur.fetchall()]
                else:
                    recent = [dict(r) for r in cur.fetchall()]
                diag["recent_news"] = [{
                    "title": r["title"][:60],
                    "source": r["source"],
                    "status": r["status"],
                    "text_len": r["text_len"] or 0,
                    "score": r["score"],
                } for r in recent]

                # Count news with no analysis at all
                cur.execute("SELECT COUNT(*) FROM news n LEFT JOIN news_analysis a ON n.id = a.news_id WHERE a.news_id IS NULL")
                diag["no_analysis"] = cur.fetchone()[0]

                # Count with score = 0
                cur.execute("SELECT COUNT(*) FROM news_analysis WHERE total_score = 0")
                diag["score_zero"] = cur.fetchone()[0]

                # Source-level audit: text_len=0 and score=0 per source
                cur.execute("""
                    SELECT n.source,
                           COUNT(*) as total,
                           SUM(CASE WHEN LENGTH(COALESCE(n.plain_text, '')) = 0 THEN 1 ELSE 0 END) as no_text,
                           SUM(CASE WHEN LENGTH(COALESCE(n.description, '')) = 0 THEN 1 ELSE 0 END) as no_desc,
                           SUM(CASE WHEN COALESCE(a.total_score, 0) = 0 THEN 1 ELSE 0 END) as score_zero
                    FROM news n
                    LEFT JOIN news_analysis a ON n.id = a.news_id
                    GROUP BY n.source
                    ORDER BY n.source
                """)
                if _is_postgres():
                    cols = [d[0] for d in cur.description]
                    audit_rows = [dict(zip(cols, r)) for r in cur.fetchall()]
                else:
                    audit_rows = [dict(r) for r in cur.fetchall()]
                diag["source_audit"] = audit_rows

            finally:
                cur.close()
        except Exception as e:
            diag["db_connected"] = False
            diag["error"] = str(e)

        self.send_response(200)
        self.send_header("Content-Type", "application/json")
        self.end_headers()
        self.wfile.write(_json.dumps(diag, indent=2, ensure_ascii=False).encode())

    def _serve_login(self):
        html = """<!DOCTYPE html>
<html lang="ru"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>IgroNews Login</title>
<link rel="icon" href="data:image/svg+xml,<svg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 100 100'><rect width='100' height='100' rx='20' fill='%23192734'/><text x='50' y='38' text-anchor='middle' font-size='28' font-family='sans-serif' font-weight='bold' fill='%231da1f2'>IGR</text><text x='50' y='70' text-anchor='middle' font-size='20' font-family='sans-serif' fill='%2317bf63'>NEWS</text><circle cx='82' cy='20' r='8' fill='%23e0245e'/></svg>">
<style>
* { margin:0;padding:0;box-sizing:border-box; }
body { font-family:-apple-system,sans-serif; background:#0f1923; color:#e1e8ed; display:flex; justify-content:center; align-items:center; height:100vh; }
.login { background:#192734; padding:40px; border-radius:12px; width:340px; }
.login h2 { color:#1da1f2; margin-bottom:20px; text-align:center; }
.login input { width:100%; padding:10px 14px; margin-bottom:12px; background:#22303c; border:1px solid #38444d; color:#e1e8ed; border-radius:6px; font-size:0.95em; }
.login input:focus { outline:none; border-color:#1da1f2; }
.login button { width:100%; padding:10px; background:#1da1f2; color:#fff; border:none; border-radius:6px; font-size:1em; cursor:pointer; }
.login button:hover { background:#1a91da; }
.error { color:#e0245e; font-size:0.85em; margin-bottom:10px; display:none; }
</style></head><body>
<div class="login">
  <div style="text-align:center;margin-bottom:10px">
    <svg width="48" height="48" viewBox="0 0 100 100">
      <rect width="100" height="100" rx="20" fill="#1da1f2"/>
      <text x="50" y="40" text-anchor="middle" font-size="30" font-family="sans-serif" font-weight="bold" fill="#fff">IGR</text>
      <text x="50" y="72" text-anchor="middle" font-size="22" font-family="sans-serif" fill="#fff" opacity="0.8">NEWS</text>
      <circle cx="85" cy="18" r="9" fill="#e0245e"/>
    </svg>
  </div>
  <h2>IgroNews</h2>
  <div class="error" id="err">Invalid credentials</div>
  <input id="username" placeholder="Username" autofocus>
  <input id="password" type="password" placeholder="Password">
  <button onclick="login()">Login</button>
</div>
<script>
document.getElementById('password').addEventListener('keypress', e => { if(e.key==='Enter') login(); });
async function login() {
  const r = await fetch('/api/login', {method:'POST', headers:{'Content-Type':'application/json'},
    body: JSON.stringify({username: document.getElementById('username').value, password: document.getElementById('password').value})
  });
  if (r.ok) { window.location.href = '/'; }
  else { document.getElementById('err').style.display = 'block'; }
}
</script></body></html>"""
        self.send_response(200)
        self.send_header("Content-Type", "text/html; charset=utf-8")
        self.end_headers()
        self.wfile.write(html.encode())

    def _get_health(self):
        from checks.health import get_sources_health
        return get_sources_health()

    # --- Data ---
    def _get_stats(self):
        conn = get_connection()
        cur = conn.cursor()
        try:
            # Single query instead of 7+2 separate queries
            cur.execute("SELECT status, COUNT(*) FROM news GROUP BY status")
            stats = {}
            total = 0
            for row in cur.fetchall():
                s, c = (row[0], row[1]) if _is_postgres() else (row[0], row[1])
                stats[s] = c
                total += c
            stats["total"] = total
            # Ensure all expected statuses exist
            for s in ["new", "in_review", "duplicate", "approved", "processed", "rejected", "ready", "moderation"]:
                stats.setdefault(s, 0)
            cur.execute("SELECT COUNT(*) FROM news_analysis")
            stats["analyzed"] = cur.fetchone()[0]
            return stats
        finally:
            cur.close()

    def _get_news(self):
        conn = get_connection()
        cur = conn.cursor()
        try:
            return self._get_news_impl(cur)
        finally:
            cur.close()

    def _get_news_impl(self, cur):
        qs = parse_qs(urlparse(self.path).query)
        limit = int(qs.get("limit", [100])[0])
        offset = int(qs.get("offset", [0])[0])
        status_filter = qs.get("status", [None])[0]
        source_filter = qs.get("source", [None])[0]
        date_from = qs.get("date_from", [None])[0]
        date_to = qs.get("date_to", [None])[0]
        llm_filter = qs.get("llm", [None])[0]

        ph = "%s" if _is_postgres() else "?"
        conditions = []
        params = []
        # Need LEFT JOIN for LLM filter on count query too
        need_join_for_count = False
        if status_filter:
            conditions.append(f"n.status = {ph}")
            params.append(status_filter)
        else:
            # Default: show only enrichment-relevant statuses
            conditions.append("n.status IN ('approved', 'processed', 'ready')")
        if source_filter:
            conditions.append(f"n.source = {ph}")
            params.append(source_filter)
        if date_from:
            conditions.append(f"n.parsed_at >= {ph}")
            params.append(date_from)
        if date_to:
            conditions.append(f"n.parsed_at <= {ph}")
            params.append(date_to + "T23:59:59")
        if llm_filter:
            need_join_for_count = True
            if llm_filter == "has_rec":
                conditions.append("a.llm_recommendation IS NOT NULL AND a.llm_recommendation != ''")
            elif llm_filter == "no_rec":
                conditions.append("(a.llm_recommendation IS NULL OR a.llm_recommendation = '')")
            else:
                conditions.append(f"LOWER(a.llm_recommendation) LIKE {ph}")
                params.append(f"%{llm_filter.lower()}%")

        where = "WHERE " + " AND ".join(conditions) if conditions else ""

        # Count total matching
        count_join = "LEFT JOIN news_analysis a ON n.id = a.news_id" if need_join_for_count else ""
        cur.execute(f"SELECT COUNT(*) FROM news n {count_join} {where}", params[:])
        total_count = cur.fetchone()[0]

        query = f"""
            SELECT n.id, n.source, n.title, n.url, n.h1, n.description,
                   n.published_at, n.parsed_at, n.status,
                   a.bigrams, a.trigrams, a.trends_data, a.keyso_data,
                   a.llm_recommendation, a.llm_trend_forecast, a.sheets_row, a.processed_at,
                   a.viral_score, a.viral_level, a.viral_data,
                   a.sentiment_label, a.sentiment_score,
                   a.freshness_status, a.freshness_hours,
                   a.tags_data, a.momentum_score, a.headline_score, a.total_score
            FROM news n
            LEFT JOIN news_analysis a ON n.id = a.news_id
            {where}
            ORDER BY n.parsed_at DESC LIMIT {ph} OFFSET {ph}
        """
        params.append(limit)
        params.append(offset)
        cur.execute(query, params)

        if _is_postgres():
            columns = [desc[0] for desc in cur.description]
            rows = [dict(zip(columns, row)) for row in cur.fetchall()]
        else:
            rows = [dict(row) for row in cur.fetchall()]

        return {"news": rows, "total": total_count, "limit": limit, "offset": offset}

    def _get_final(self):
        """Финальная подборка: только publish_now с финальным скором."""
        conn = get_connection()
        cur = conn.cursor()
        try:
            return self._get_final_impl(cur)
        finally:
            cur.close()

    def _get_final_impl(self, cur):
        qs = parse_qs(urlparse(self.path).query)
        limit = int(qs.get("limit", [100])[0])
        offset = int(qs.get("offset", [0])[0])
        source_filter = qs.get("source", [None])[0]
        sort_field = qs.get("sort", ["final_score"])[0]
        sort_dir = qs.get("dir", ["desc"])[0]

        ph = "%s" if _is_postgres() else "?"
        conditions = [
            "n.status IN ('processed', 'ready')",
            "LOWER(a.llm_recommendation) = 'publish_now'",
        ]
        params = []

        if source_filter:
            conditions.append(f"n.source = {ph}")
            params.append(source_filter)

        where = "WHERE " + " AND ".join(conditions)

        cur.execute(f"SELECT COUNT(*) FROM news n JOIN news_analysis a ON n.id = a.news_id {where}", params[:])
        total_count = cur.fetchone()[0]

        # Финальный скор: внутренний (40%) + viral (15%) + keyso freq бонус (15%) + trends бонус (15%) + headline (15%)
        # keyso/trends бонусы вычисляются в JS, здесь берём raw данные
        allowed_sorts = {
            "final_score": "a.total_score",
            "total_score": "a.total_score",
            "viral_score": "a.viral_score",
            "freshness_hours": "a.freshness_hours",
            "source": "n.source",
            "parsed_at": "n.parsed_at",
        }
        order_col = allowed_sorts.get(sort_field, "a.total_score")
        order_dir = "ASC" if sort_dir == "asc" else "DESC"

        query = f"""
            SELECT n.id, n.source, n.title, n.url, n.h1,
                   n.published_at, n.parsed_at, n.status,
                   a.bigrams, a.trigrams, a.trends_data, a.keyso_data,
                   a.llm_recommendation, a.llm_trend_forecast,
                   a.viral_score, a.viral_level, a.viral_data,
                   a.sentiment_label, a.sentiment_score,
                   a.freshness_status, a.freshness_hours,
                   a.tags_data, a.momentum_score, a.headline_score,
                   a.total_score, a.quality_score, a.relevance_score,
                   a.entity_names, a.entity_best_tier, a.processed_at
            FROM news n
            JOIN news_analysis a ON n.id = a.news_id
            {where}
            ORDER BY {order_col} {order_dir} LIMIT {ph} OFFSET {ph}
        """
        params.append(limit)
        params.append(offset)
        cur.execute(query, params)

        if _is_postgres():
            columns = [desc[0] for desc in cur.description]
            rows = [dict(zip(columns, row)) for row in cur.fetchall()]
        else:
            rows = [dict(row) for row in cur.fetchall()]

        return {"news": rows, "total": total_count}

    def _get_editorial(self):
        """Единый endpoint для вкладки Редакция — все данные в одном запросе."""
        conn = get_connection()
        cur = conn.cursor()
        try:
            return self._get_editorial_impl(cur)
        finally:
            cur.close()

    def _get_editorial_impl(self, cur):
        qs = parse_qs(urlparse(self.path).query)
        limit = int(qs.get("limit", [100])[0])
        offset = int(qs.get("offset", [0])[0])
        status_filter = qs.get("status", [None])[0]
        source_filter = qs.get("source", [None])[0]
        min_score = int(qs.get("min_score", [0])[0])
        max_score = int(qs.get("max_score", [0])[0])
        score_filter = qs.get("score_filter", [None])[0]  # "zero" or "nonzero"
        viral_level = qs.get("viral_level", [None])[0]
        tier_filter = qs.get("tier", [None])[0]
        search = qs.get("q", [None])[0]

        ph = "%s" if _is_postgres() else "?"
        conditions = []
        params = []

        # Исключаем дубликаты и отклонённые по умолчанию (если не запрошены)
        if status_filter:
            conditions.append(f"n.status = {ph}")
            params.append(status_filter)
        else:
            conditions.append(f"n.status NOT IN ('duplicate', 'rejected')")

        if source_filter:
            conditions.append(f"n.source = {ph}")
            params.append(source_filter)
        if min_score > 0:
            conditions.append(f"COALESCE(a.total_score, 0) >= {ph}")
            params.append(min_score)
        if max_score > 0:
            conditions.append(f"COALESCE(a.total_score, 0) <= {ph}")
            params.append(max_score)
        if score_filter == "zero":
            conditions.append("COALESCE(a.total_score, 0) = 0")
        elif score_filter == "nonzero":
            conditions.append("COALESCE(a.total_score, 0) > 0")
        if viral_level:
            conditions.append(f"a.viral_level = {ph}")
            params.append(viral_level)
        if tier_filter:
            conditions.append(f"a.entity_best_tier = {ph}")
            params.append(tier_filter)
        if search:
            conditions.append(f"LOWER(n.title) LIKE {ph}")
            params.append(f"%{search.lower()}%")

        where = "WHERE " + " AND ".join(conditions) if conditions else ""

        # Count
        cur.execute(f"SELECT COUNT(*) FROM news n LEFT JOIN news_analysis a ON n.id = a.news_id {where}", params[:])
        total_count = cur.fetchone()[0]

        # Stats (counts by status)
        stat_params = []
        cur.execute("SELECT status, COUNT(*) FROM news GROUP BY status")
        status_counts = {}
        for row in cur.fetchall():
            if _is_postgres():
                status_counts[row[0]] = row[1]
            else:
                status_counts[row[0]] = row[1]

        query = f"""
            SELECT n.id, n.source, n.title, n.description, n.url, n.published_at, n.parsed_at, n.status,
                   COALESCE(a.total_score, 0) as total_score,
                   COALESCE(a.quality_score, 0) as quality_score,
                   COALESCE(a.relevance_score, 0) as relevance_score,
                   COALESCE(a.viral_score, 0) as viral_score,
                   COALESCE(a.viral_level, '') as viral_level,
                   COALESCE(a.viral_data, '[]') as viral_data,
                   COALESCE(a.sentiment_label, '') as sentiment_label,
                   COALESCE(a.sentiment_score, 0) as sentiment_score,
                   COALESCE(a.freshness_status, '') as freshness_status,
                   COALESCE(a.freshness_hours, -1) as freshness_hours,
                   COALESCE(a.tags_data, '[]') as tags_data,
                   COALESCE(a.momentum_score, 0) as momentum_score,
                   COALESCE(a.headline_score, 0) as headline_score,
                   COALESCE(a.all_checks_pass, 0) as all_checks_pass,
                   COALESCE(a.entity_names, '[]') as entity_names,
                   COALESCE(a.entity_best_tier, '') as entity_best_tier,
                   COALESCE(a.reviewed_at, '') as reviewed_at,
                   COALESCE(a.score_breakdown, '{{}}') as score_breakdown,
                   a.bigrams, a.llm_recommendation, a.llm_trend_forecast,
                   a.keyso_data, a.trends_data
            FROM news n
            LEFT JOIN news_analysis a ON n.id = a.news_id
            {where}
            ORDER BY n.parsed_at DESC
            LIMIT {ph} OFFSET {ph}
        """
        params.append(limit)
        params.append(offset)
        cur.execute(query, params)

        if _is_postgres():
            columns = [desc[0] for desc in cur.description]
            rows = [dict(zip(columns, row)) for row in cur.fetchall()]
        else:
            rows = [dict(row) for row in cur.fetchall()]

        return {
            "news": rows,
            "total": total_count,
            "limit": limit,
            "offset": offset,
            "stats": status_counts,
        }

    def _get_event_chain_by_id(self, news_id):
        """Возвращает цепочку событий для указанной новости (GET endpoint)."""
        try:
            conn = get_connection()
            cur = conn.cursor()
            ph = "%s" if _is_postgres() else "?"
            try:
                cur.execute(f"SELECT id, source, title, published_at, status FROM news WHERE id = {ph}", (news_id,))
                row = cur.fetchone()
                if not row:
                    return {"error": "news not found"}
                if _is_postgres():
                    columns = [desc[0] for desc in cur.description]
                    news = dict(zip(columns, row))
                else:
                    news = dict(row)
            finally:
                cur.close()
            from checks.temporal_clusters import get_event_chain
            return get_event_chain(news)
        except Exception as e:
            logger.error(f"Event chain error: {e}")
            return {"chain": [], "chain_length": 0, "days_span": 0, "phase": "single", "error": str(e)}

    def _get_sources(self):
        import config
        return config.SOURCES

    def _get_prompts(self):
        from apis.llm import PROMPT_TREND_FORECAST, PROMPT_MERGE_ANALYSIS, PROMPT_KEYSO_QUERIES
        return {
            "trend_forecast": PROMPT_TREND_FORECAST,
            "merge_analysis": PROMPT_MERGE_ANALYSIS,
            "keyso_queries": PROMPT_KEYSO_QUERIES,
        }

    def _get_settings(self):
        import config
        return {
            "llm_model": config.LLM_MODEL,
            "keyso_region": getattr(config, "KEYSO_REGION", "ru"),
            "regions": config.REGIONS,
            "sheets_id": config.GOOGLE_SHEETS_ID,
            "sheets_tab": config.SHEETS_TAB,
            "openai_key_set": bool(config.OPENAI_API_KEY),
            "keyso_key_set": bool(config.KEYSO_API_KEY),
            "google_sa_set": bool(config.GOOGLE_SERVICE_ACCOUNT_JSON),
            "auto_approve_threshold": getattr(config, "AUTO_APPROVE_THRESHOLD", 70),
            "auto_rewrite_on_publish_now": getattr(config, "AUTO_REWRITE_ON_PUBLISH_NOW", True),
            "auto_rewrite_style": getattr(config, "AUTO_REWRITE_STYLE", "news"),
        }

    # --- Actions ---
    def _run_process(self):
        try:
            from scheduler import process_news
            threading.Thread(target=process_news, daemon=True).start()
            self._json({"status": "ok", "message": "Processing started in background"})
        except Exception as e:
            self._json({"status": "error", "message": str(e)})

    def _process_one(self, body):
        news_id = body.get("news_id")
        if not news_id:
            self._json({"status": "error", "message": "news_id required"})
            return
        try:
            from scheduler import _process_single_news
            result = _process_single_news(news_id)
            self._json({"status": "ok", "result": result})
        except Exception as e:
            self._json({"status": "error", "message": str(e)})

    def _export_sheets(self, body):
        news_id = body.get("news_id")
        try:
            from storage.sheets import write_news_row
            conn = get_connection()
            cur = conn.cursor()
            try:
                ph = "%s" if _is_postgres() else "?"
                cur.execute(f"SELECT * FROM news WHERE id = {ph}", (news_id,))
                if _is_postgres():
                    columns = [desc[0] for desc in cur.description]
                    news = dict(zip(columns, cur.fetchone()))
                else:
                    news = dict(cur.fetchone())

                cur.execute(f"SELECT * FROM news_analysis WHERE news_id = {ph}", (news_id,))
                row = cur.fetchone()
                if row:
                    if _is_postgres():
                        columns = [desc[0] for desc in cur.description]
                        analysis = dict(zip(columns, row))
                    else:
                        analysis = dict(row)
                else:
                    analysis = {"bigrams": "[]", "trends_data": "{}", "keyso_data": "{}",
                               "llm_recommendation": "", "llm_trend_forecast": "", "llm_merged_with": ""}

                sheet_row = write_news_row(news, analysis)
                self._json({"status": "ok", "row": sheet_row})
            finally:
                cur.close()
        except Exception as e:
            self._json({"status": "error", "message": str(e)})

    def _add_source(self, body):
        import config
        source = {
            "name": body.get("name", ""),
            "type": body.get("type", "rss"),
            "url": body.get("url", ""),
            "interval": int(body.get("interval", 15)),
        }
        if body.get("selector"):
            source["selector"] = body["selector"]
        config.SOURCES.append(source)
        self._json({"status": "ok", "sources": config.SOURCES})

    def _edit_source(self, body):
        import config
        old_name = body.get("old_name", "")
        for s in config.SOURCES:
            if s["name"] == old_name:
                s["name"] = body.get("name", s["name"])
                s["type"] = body.get("type", s["type"])
                s["url"] = body.get("url", s["url"])
                s["interval"] = int(body.get("interval", s["interval"]))
                if body.get("selector"):
                    s["selector"] = body["selector"]
                elif "selector" in s and body.get("type") == "rss":
                    del s["selector"]
                break
        self._json({"status": "ok", "sources": config.SOURCES})

    def _delete_source(self, body):
        import config
        name = body.get("name")
        config.SOURCES[:] = [s for s in config.SOURCES if s["name"] != name]
        self._json({"status": "ok", "sources": config.SOURCES})

    def _save_prompts(self, body):
        import apis.llm as llm
        if "trend_forecast" in body:
            llm.PROMPT_TREND_FORECAST = body["trend_forecast"]
        if "merge_analysis" in body:
            llm.PROMPT_MERGE_ANALYSIS = body["merge_analysis"]
        if "keyso_queries" in body:
            llm.PROMPT_KEYSO_QUERIES = body["keyso_queries"]
        self._json({"status": "ok"})

    def _save_settings(self, body):
        if not self._require_perm("settings"):
            return
        import config
        user = self._get_session_user() or "admin"
        changes = []
        if "llm_model" in body and body["llm_model"] != config.LLM_MODEL:
            changes.append(("llm_model", config.LLM_MODEL, body["llm_model"]))
            config.LLM_MODEL = body["llm_model"]
        if "keyso_region" in body and body["keyso_region"] != config.KEYSO_REGION:
            changes.append(("keyso_region", config.KEYSO_REGION, body["keyso_region"]))
            config.KEYSO_REGION = body["keyso_region"]
        if "sheets_tab" in body and body["sheets_tab"] != config.SHEETS_TAB:
            changes.append(("sheets_tab", config.SHEETS_TAB, body["sheets_tab"]))
            config.SHEETS_TAB = body["sheets_tab"]
        if "auto_approve_threshold" in body:
            try:
                new_val = int(body["auto_approve_threshold"])
                if new_val != config.AUTO_APPROVE_THRESHOLD:
                    changes.append(("auto_approve_threshold", str(config.AUTO_APPROVE_THRESHOLD), str(new_val)))
                    config.AUTO_APPROVE_THRESHOLD = new_val
            except (ValueError, TypeError):
                pass
        if "auto_rewrite_on_publish_now" in body:
            new_val = bool(body["auto_rewrite_on_publish_now"])
            if new_val != config.AUTO_REWRITE_ON_PUBLISH_NOW:
                changes.append(("auto_rewrite_on_publish_now", str(config.AUTO_REWRITE_ON_PUBLISH_NOW), str(new_val)))
                config.AUTO_REWRITE_ON_PUBLISH_NOW = new_val
        if "auto_rewrite_style" in body and body["auto_rewrite_style"] != config.AUTO_REWRITE_STYLE:
            changes.append(("auto_rewrite_style", config.AUTO_REWRITE_STYLE, body["auto_rewrite_style"]))
            config.AUTO_REWRITE_STYLE = body["auto_rewrite_style"]

        # Audit log config changes
        for setting_name, old_val, new_val in changes:
            try:
                from core.observability import log_config_change
                log_config_change(setting_name, old_val, new_val, changed_by=user)
            except Exception:
                pass

        self._json({"status": "ok"})

    def _test_llm(self, body):
        try:
            import config
            from openai import OpenAI
            import json as _json
            prompt = body.get("prompt", "Ответь JSON: {\"test\": \"ok\"}")
            client = OpenAI(api_key=config.OPENAI_API_KEY, base_url=config.OPENAI_BASE_URL)
            response = client.chat.completions.create(
                model=config.LLM_MODEL,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.3,
            )
            text = response.choices[0].message.content
            # Try parse JSON
            cleaned = text.strip()
            if cleaned.startswith("```"):
                cleaned = "\n".join(cleaned.split("\n")[1:])
                if cleaned.endswith("```"):
                    cleaned = cleaned[:-3]
                cleaned = cleaned.strip()
            try:
                parsed = _json.loads(cleaned)
            except Exception:
                parsed = None
            self._json({"status": "ok", "model": config.LLM_MODEL, "base_url": config.OPENAI_BASE_URL, "raw": text, "result": parsed})
        except Exception as e:
            self._json({"status": "error", "message": str(e), "type": type(e).__name__})

    def _test_keyso(self, body):
        try:
            import config
            import requests as _req
            keyword = body.get("keyword", "gta 6")
            # Raw request for debugging
            url = f"{config.KEYSO_BASE_URL}/report/simple/keyword_dashboard"
            params = {"auth-token": config.KEYSO_API_KEY, "base": config.KEYSO_REGION, "keyword": keyword}
            resp = _req.get(url, params=params, timeout=15)
            raw = resp.json()
            self._json({"status": "ok", "http_code": resp.status_code, "raw_response": raw})
        except Exception as e:
            self._json({"status": "error", "message": str(e), "type": type(e).__name__})

    def _quick_tags(self, body):
        """Быстрый расчёт тегов по заголовкам (без полного review)."""
        news_ids = body.get("news_ids", [])
        if not news_ids:
            self._json({"status": "error", "message": "No news_ids"})
            return
        try:
            conn = get_connection()
            cur = conn.cursor()
            try:
                ph = "%s" if _is_postgres() else "?"
                placeholders = ",".join([ph] * len(news_ids))
                cur.execute(f"SELECT id, title, description FROM news WHERE id IN ({placeholders})", news_ids)
                if _is_postgres():
                    columns = [desc[0] for desc in cur.description]
                    rows = [dict(zip(columns, row)) for row in cur.fetchall()]
                else:
                    rows = [dict(row) for row in cur.fetchall()]

                from checks.tags import auto_tag
                from checks.deduplication import tfidf_similarity

                # Tags per news
                tags_map = {}
                for r in rows:
                    tags = auto_tag(r)
                    tags_map[r["id"]] = [{"id": t["id"], "label": t["label"], "hits": t["hits"]} for t in tags[:3]]

                # Similarity groups
                titles = [r.get("title", "") for r in rows]
                ids_ordered = [r["id"] for r in rows]
                pairs = tfidf_similarity(titles)

                # Build groups from pairs
                from collections import defaultdict
                graph = defaultdict(set)
                for i, j, score in pairs:
                    graph[i].add(j)
                    graph[j].add(i)
                visited = set()
                groups = []
                group_idx = 0
                id_to_group = {}
                for idx in range(len(rows)):
                    if idx in visited:
                        continue
                    cluster = set()
                    stack = [idx]
                    while stack:
                        node = stack.pop()
                        if node in visited:
                            continue
                        visited.add(node)
                        cluster.add(node)
                        stack.extend(graph[node] - visited)
                    if len(cluster) >= 2:
                        group_idx += 1
                        member_ids = [ids_ordered[i] for i in sorted(cluster)]
                        member_titles = [titles[i] for i in sorted(cluster)]
                        for mid in member_ids:
                            id_to_group[mid] = group_idx
                        groups.append({
                            "group": group_idx,
                            "count": len(member_ids),
                            "ids": member_ids,
                            "titles": member_titles,
                        })

                self._json({"status": "ok", "tags": tags_map, "groups": groups, "id_to_group": id_to_group})
            finally:
                cur.close()
        except Exception as e:
            self._json({"status": "error", "message": str(e), "type": type(e).__name__})

    def _dashboard_groups(self):
        """Возвращает теги и группы для новостей (учитывает фильтр статуса)."""
        try:
            conn = get_connection()
            cur = conn.cursor()
            try:
                ph = "%s" if _is_postgres() else "?"
                qs = parse_qs(urlparse(self.path).query)
                status_filter = qs.get("status", [None])[0]
                if status_filter:
                    cur.execute(f"SELECT id, title, description FROM news WHERE status = {ph} ORDER BY parsed_at DESC LIMIT 100", (status_filter,))
                else:
                    cur.execute("SELECT id, title, description FROM news ORDER BY parsed_at DESC LIMIT 100")
                if _is_postgres():
                    columns = [desc[0] for desc in cur.description]
                    rows = [dict(zip(columns, row)) for row in cur.fetchall()]
                else:
                    rows = [dict(row) for row in cur.fetchall()]

                if not rows:
                    self._json({"status": "ok", "tags": {}, "groups": [], "id_to_group": {}})
                    return

                from checks.tags import auto_tag
                from checks.deduplication import tfidf_similarity
                from collections import defaultdict

                tags_map = {}
                for r in rows:
                    tags = auto_tag(r)
                    tags_map[r["id"]] = [{"id": t["id"], "label": t["label"], "hits": t["hits"]} for t in tags[:3]]

                titles = [r.get("title", "") for r in rows]
                ids_ordered = [r["id"] for r in rows]
                pairs = tfidf_similarity(titles)

                graph = defaultdict(set)
                for i, j, score in pairs:
                    graph[i].add(j)
                    graph[j].add(i)
                visited = set()
                groups = []
                group_idx = 0
                id_to_group = {}
                for idx in range(len(rows)):
                    if idx in visited:
                        continue
                    cluster = set()
                    stack = [idx]
                    while stack:
                        node = stack.pop()
                        if node in visited:
                            continue
                        visited.add(node)
                        cluster.add(node)
                        stack.extend(graph[node] - visited)
                    if len(cluster) >= 2:
                        group_idx += 1
                        member_ids = [ids_ordered[i] for i in sorted(cluster)]
                        member_titles = [titles[i] for i in sorted(cluster)]
                        for mid in member_ids:
                            id_to_group[mid] = group_idx
                        groups.append({
                            "group": group_idx,
                            "count": len(member_ids),
                            "ids": member_ids,
                            "titles": member_titles,
                        })

                self._json({"status": "ok", "tags": tags_map, "groups": groups, "id_to_group": id_to_group})
            finally:
                cur.close()
        except Exception as e:
            self._json({"status": "error", "message": str(e), "type": type(e).__name__})

    def _run_review(self, body):
        """Запускает pipeline проверки для выбранных новостей."""
        news_ids = body.get("news_ids", [])
        if not news_ids:
            self._json({"status": "error", "message": "No news selected"})
            return
        try:
            conn = get_connection()
            cur = conn.cursor()
            try:
                ph = "%s" if _is_postgres() else "?"
                placeholders = ",".join([ph] * len(news_ids))
                cur.execute(f"SELECT id, source, url, title, h1, description, plain_text, published_at, parsed_at, status FROM news WHERE id IN ({placeholders})", news_ids)
                if _is_postgres():
                    columns = [desc[0] for desc in cur.description]
                    news_list = [dict(zip(columns, row)) for row in cur.fetchall()]
                else:
                    news_list = [dict(row) for row in cur.fetchall()]

                from checks.pipeline import run_review_pipeline
                result = run_review_pipeline(news_list)
                self._json({"status": "ok", **result})
            finally:
                cur.close()
        except Exception as e:
            self._json({"status": "error", "message": str(e), "type": type(e).__name__})

    def _review_batch(self, body):
        """Проверяет новости по статусу (batch, без изменения статуса)."""
        status = body.get("status", "new")
        limit = int(body.get("limit", 50))
        try:
            conn = get_connection()
            cur = conn.cursor()
            try:
                ph = "%s" if _is_postgres() else "?"
                if status:
                    cur.execute(f"SELECT id, source, url, title, h1, description, plain_text, published_at, parsed_at, status FROM news WHERE status = {ph} ORDER BY parsed_at DESC LIMIT {ph}", (status, limit))
                else:
                    cur.execute(f"SELECT id, source, url, title, h1, description, plain_text, published_at, parsed_at, status FROM news ORDER BY parsed_at DESC LIMIT {ph}", (limit,))
                if _is_postgres():
                    columns = [desc[0] for desc in cur.description]
                    news_list = [dict(zip(columns, row)) for row in cur.fetchall()]
                else:
                    news_list = [dict(row) for row in cur.fetchall()]

                if not news_list:
                    self._json({"status": "ok", "results": [], "groups": []})
                    return

                from checks.pipeline import run_review_pipeline
                # Прогоняем pipeline но НЕ меняем статусы
                from checks.deduplication import tfidf_similarity, build_groups
                from checks.quality import check_quality
                from checks.relevance import check_relevance
                from checks.freshness import check_freshness
                from checks.viral_score import viral_score
                from checks.tags import auto_tag
                from checks.sentiment import analyze_sentiment
                from checks.momentum import get_momentum

                results = []
                for news in news_list:
                    result = {
                        "id": news["id"],
                        "title": news.get("title", ""),
                        "source": news.get("source", ""),
                        "url": news.get("url", ""),
                        "published_at": news.get("published_at", ""),
                        "status": news.get("status", ""),
                        "checks": {},
                    }
                    result["checks"]["quality"] = check_quality(news)
                    result["checks"]["relevance"] = check_relevance(news)
                    result["checks"]["freshness"] = check_freshness(news)
                    result["checks"]["viral"] = viral_score(news)
                    result["tags"] = auto_tag(news)
                    result["sentiment"] = analyze_sentiment(news)
                    result["momentum"] = get_momentum(news)

                    all_pass = all(c["pass"] for c in result["checks"].values())
                    total_score = sum(c["score"] for c in result["checks"].values()) // 4
                    momentum_bonus = result["momentum"]["score"] // 5
                    total_score = min(100, total_score + momentum_bonus)
                    result["overall_pass"] = all_pass
                    result["total_score"] = total_score
                    results.append(result)

                # Dedup
                titles = [r["title"] for r in results]
                pairs = tfidf_similarity(titles)
                groups = build_groups(results, pairs)
                for group in groups:
                    for idx in group.get("duplicate_indices", []):
                        if idx < len(results):
                            results[idx]["overall_pass"] = False
                            results[idx]["is_duplicate"] = True
                    for member in group["members"]:
                        member["dedup_status"] = group["status"]

                self._json({"status": "ok", "results": results, "groups": groups})
            finally:
                cur.close()
        except Exception as e:
            self._json({"status": "error", "message": str(e), "type": type(e).__name__})

    def _run_auto_review(self, body):
        """Запускает авто-ревью батчами по 20 (сохраняет результаты в БД)."""
        try:
            conn = get_connection()
            cur = conn.cursor()
            try:
                ph = "%s" if _is_postgres() else "?"
                BATCH_SIZE = 20

                # Считаем сколько всего непроверенных (только new)
                cur.execute("SELECT COUNT(*) FROM news WHERE status = 'new'")
                total_pending = cur.fetchone()[0]

                if total_pending == 0:
                    self._json({"status": "ok", "reviewed": 0, "message": "Нет новых для проверки", "remaining": 0})
                    return

                # Берём один батч
                cur.execute(f"""
                    SELECT id, source, url, title, h1, description, plain_text, published_at, parsed_at, status
                    FROM news WHERE status = 'new'
                    ORDER BY parsed_at DESC LIMIT {ph}
                """, (BATCH_SIZE,))
                if _is_postgres():
                    columns = [desc[0] for desc in cur.description]
                    news_list = [dict(zip(columns, row)) for row in cur.fetchall()]
                else:
                    news_list = [dict(row) for row in cur.fetchall()]

                from checks.pipeline import run_review_pipeline
                result = run_review_pipeline(news_list, update_status=True)
                reviewed = len(result.get("results", []))
                dupes = sum(1 for r in result.get("results", []) if r.get("is_duplicate"))
                rejected = sum(1 for r in result.get("results", []) if r.get("auto_rejected"))
                remaining = total_pending - reviewed
                self._json({
                    "status": "ok",
                    "reviewed": reviewed,
                    "duplicates": dupes,
                    "auto_rejected": rejected,
                    "remaining": remaining,
                    "message": f"Проверено: {reviewed}, дубликатов: {dupes}, отклонено: {rejected}" +
                               (f". Осталось: {remaining}" if remaining > 0 else "")
                })
            finally:
                cur.close()
        except Exception as e:
            self._json({"status": "error", "message": str(e)})

    def _approve_news(self, body):
        """Одобряет новости и запускает обогащение в фоне."""
        news_ids = body.get("news_ids", [])
        if not news_ids:
            self._json({"status": "error", "message": "No news selected"})
            return
        try:
            from checks.pipeline import approve_for_enrichment
            from checks.feedback import record_decision
            approve_for_enrichment(news_ids)
            for nid in news_ids:
                try:
                    record_decision(nid, "approved")
                except Exception:
                    pass

            # Auto-enrich: запускаем обогащение в фоновом потоке
            import threading
            def _bg_enrich(ids):
                from scheduler import _process_single_news
                for nid in ids:
                    try:
                        _process_single_news(nid)
                    except Exception as e:
                        logger.warning("Background enrich failed for %s: %s", nid, e)
            threading.Thread(target=_bg_enrich, args=(list(news_ids),), daemon=True).start()

            self._json({"status": "ok", "approved": len(news_ids), "enriching": True})
        except Exception as e:
            self._json({"status": "error", "message": str(e)})

    def _reject_news(self, body):
        """Отклоняет новости (одну или массив)."""
        news_ids = body.get("news_ids", [])
        news_id = body.get("news_id")
        if news_id and not news_ids:
            news_ids = [news_id]
        if not news_ids:
            self._json({"status": "error", "message": "news_ids required"})
            return
        try:
            from checks.feedback import record_decision
            for nid in news_ids:
                update_news_status(nid, "rejected")
                try:
                    record_decision(nid, "rejected")
                except Exception:
                    pass
            self._json({"status": "ok", "rejected": len(news_ids)})
        except Exception as e:
            self._json({"status": "error", "message": str(e)})

    def _export_all_processed(self, body):
        """Экспортирует все обработанные новости (processed) в Sheets/Лист1."""
        conn = get_connection()
        cur = conn.cursor()
        try:
            cur.execute("SELECT id FROM news WHERE status IN ('processed', 'ready') ORDER BY parsed_at DESC")
            if _is_postgres():
                news_ids = [r[0] for r in cur.fetchall()]
            else:
                news_ids = [r["id"] for r in cur.fetchall()]
        finally:
            cur.close()

        if not news_ids:
            self._json({"status": "error", "message": "Нет обработанных новостей"})
            return

        # Queue sheets export for all
        from scheduler import _create_task
        task_ids = []
        cur2 = conn.cursor()
        ph = "%s" if _is_postgres() else "?"
        try:
            for nid in news_ids:
                try:
                    cur2.execute(f"SELECT title FROM news WHERE id = {ph}", (nid,))
                    row = cur2.fetchone()
                    title = (row[0] if _is_postgres() else row["title"]) if row else ""
                except Exception:
                    title = ""
                tid = _create_task("sheets", nid, title)
                task_ids.append(tid)
        finally:
            cur2.close()

        # Process in background with rate limiting
        import threading
        import time as _time
        def _bg_export(ids, tids):
            from storage.sheets import write_news_row
            from scheduler import _update_task, _fetch_news_by_id, _fetch_analysis_by_id
            ok_count = 0
            skip_count = 0
            err_count = 0
            for i, (nid, tid) in enumerate(zip(ids, tids)):
                try:
                    _update_task(tid, "running", {"stage": "exporting", "progress": f"{i+1}/{len(ids)}"})
                    news = _fetch_news_by_id(nid)
                    analysis = _fetch_analysis_by_id(nid)
                    if not news:
                        _update_task(tid, "error", {"error": "News not found"})
                        err_count += 1
                        continue
                    sheet_row = write_news_row(news, analysis or {})
                    if sheet_row and sheet_row > 0:
                        _update_task(tid, "done", {"sheet_row": sheet_row})
                        ok_count += 1
                    elif sheet_row == -1:
                        _update_task(tid, "skipped", {"reason": "duplicate in Sheets"})
                        skip_count += 1
                    else:
                        _update_task(tid, "error", {"error": "Sheets write returned None"})
                        err_count += 1
                except Exception as e:
                    _update_task(tid, "error", {"error": str(e)[:500]})
                    err_count += 1
                # Rate limit: ~1.5 sec between writes to stay under Google Sheets 60 req/min
                _time.sleep(1.5)
            logger.info("Mass export done: %d ok, %d skipped, %d errors out of %d", ok_count, skip_count, err_count, len(ids))

        threading.Thread(target=_bg_export, args=(list(news_ids), list(task_ids)), daemon=True).start()
        self._json({"status": "ok", "queued": len(news_ids), "task_ids": task_ids})

    def _export_ready_all(self, body):
        """Экспортирует ВСЕ переписанные статьи (articles) в Sheets/Ready."""
        conn = get_connection()
        cur = conn.cursor()
        try:
            # Get all articles with their news data
            cur.execute("""
                SELECT a.id, a.news_id, a.title, a.text, a.seo_title, a.seo_description,
                       a.tags, a.style, a.original_title, a.source_url, a.created_at,
                       n.source, n.parsed_at, n.url, n.title as news_title
                FROM articles a
                LEFT JOIN news n ON n.id = a.news_id
                ORDER BY a.created_at DESC
            """)
            if _is_postgres():
                columns = [d[0] for d in cur.description]
                articles = [dict(zip(columns, r)) for r in cur.fetchall()]
            else:
                articles = [dict(r) for r in cur.fetchall()]
        finally:
            cur.close()

        if not articles:
            self._json({"status": "error", "message": "Нет переписанных статей"})
            return

        import threading
        import time as _time

        def _bg_export_ready(arts):
            from storage.sheets import write_ready_row
            from scheduler import _fetch_analysis_by_id
            import json as _json
            ok = 0
            skip = 0
            err = 0
            for art in arts:
                try:
                    news_id = art.get("news_id", "")
                    analysis = _fetch_analysis_by_id(news_id) if news_id else None

                    # Build news dict
                    news = {
                        "parsed_at": art.get("parsed_at", art.get("created_at", "")),
                        "source": art.get("source", ""),
                        "title": art.get("original_title") or art.get("news_title", ""),
                        "url": art.get("source_url") or art.get("url", ""),
                    }

                    # Build rewrite dict
                    tags_raw = art.get("tags", "[]")
                    try:
                        tags_list = _json.loads(tags_raw) if isinstance(tags_raw, str) else (tags_raw if isinstance(tags_raw, list) else [])
                    except Exception:
                        tags_list = []

                    rewrite = {
                        "title": art.get("title", ""),
                        "text": art.get("text", ""),
                        "seo_title": art.get("seo_title", ""),
                        "seo_description": art.get("seo_description", ""),
                        "tags": tags_list,
                    }

                    row = write_ready_row(news, analysis, rewrite)
                    if row and row > 0:
                        ok += 1
                    elif row == -1:
                        skip += 1
                    else:
                        err += 1
                except Exception as e:
                    logger.error("Ready export error for article %s: %s", art.get("id"), e)
                    err += 1
                _time.sleep(1.5)  # Rate limit
            logger.info("Ready export done: %d ok, %d skipped, %d errors out of %d", ok, skip, err, len(arts))

        threading.Thread(target=_bg_export_ready, args=(list(articles),), daemon=True).start()
        self._json({"status": "ok", "queued": len(articles), "message": f"Экспорт {len(articles)} статей в Ready запущен"})

    # ─── Pipeline endpoints ───

    def _pipeline_full_auto(self, body):
        """Режим 1: Полный автомат — score → >70 на LLM → финальный скор → >60 на рерайт → Sheets/Ready."""
        if not self._require_perm("pipeline"):
            return
        news_ids = body.get("news_ids", [])
        select_all = body.get("all_new", False)

        if select_all and not news_ids:
            # Выбрать все new/in_review (исключая уже обработанные)
            conn = get_connection()
            cur = conn.cursor()
            try:
                cur.execute("SELECT id FROM news WHERE status IN ('new', 'in_review') AND status NOT IN ('duplicate', 'rejected', 'ready', 'processed') ORDER BY parsed_at DESC LIMIT 5000")
                if _is_postgres():
                    news_ids = [r[0] for r in cur.fetchall()]
                else:
                    news_ids = [r["id"] for r in cur.fetchall()]
            finally:
                cur.close()

        if not news_ids:
            self._json({"status": "error", "message": "Нет новостей для обработки"})
            return

        # Create tasks
        from scheduler import _create_task, run_full_auto_pipeline
        task_ids = []
        conn = get_connection()
        cur = conn.cursor()
        try:
            ph = "%s" if _is_postgres() else "?"
            for nid in news_ids:
                try:
                    cur.execute(f"SELECT title FROM news WHERE id = {ph}", (nid,))
                    row = cur.fetchone()
                    title = (row[0] if _is_postgres() else row["title"]) if row else ""
                except Exception:
                    title = ""
                tid = _create_task("full_auto", nid, title)
                task_ids.append(tid)
        finally:
            cur.close()

        # Run in background
        import threading
        threading.Thread(
            target=run_full_auto_pipeline,
            args=(list(news_ids), list(task_ids)),
            daemon=True
        ).start()

        self._json({"status": "ok", "queued": len(news_ids), "task_ids": task_ids})

    def _pipeline_no_llm(self, body):
        """Режим 2: Без LLM — score → Sheets/NotReady + Модерация."""
        if not self._require_perm("pipeline"):
            return
        news_ids = body.get("news_ids", [])
        select_all = body.get("all_new", False)

        if select_all and not news_ids:
            conn = get_connection()
            cur = conn.cursor()
            try:
                cur.execute("SELECT id FROM news WHERE status IN ('new', 'in_review') AND status NOT IN ('duplicate', 'rejected', 'ready', 'processed') ORDER BY parsed_at DESC LIMIT 5000")
                if _is_postgres():
                    news_ids = [r[0] for r in cur.fetchall()]
                else:
                    news_ids = [r["id"] for r in cur.fetchall()]
            finally:
                cur.close()

        if not news_ids:
            self._json({"status": "error", "message": "Нет новостей для обработки"})
            return

        from scheduler import _create_task, run_no_llm_pipeline
        task_ids = []
        conn = get_connection()
        cur = conn.cursor()
        try:
            ph = "%s" if _is_postgres() else "?"
            for nid in news_ids:
                try:
                    cur.execute(f"SELECT title FROM news WHERE id = {ph}", (nid,))
                    row = cur.fetchone()
                    title = (row[0] if _is_postgres() else row["title"]) if row else ""
                except Exception:
                    title = ""
                tid = _create_task("no_llm", nid, title)
                task_ids.append(tid)
        finally:
            cur.close()

        import threading
        threading.Thread(
            target=run_no_llm_pipeline,
            args=(list(news_ids), list(task_ids)),
            daemon=True
        ).start()

        self._json({"status": "ok", "queued": len(news_ids), "task_ids": task_ids})

    def _pipeline_stop(self, body):
        """Остановка текущего пайплайна + отмена pending задач в БД."""
        from scheduler import pipeline_stop
        pipeline_stop()
        # Cancel all pending/running pipeline tasks in DB
        conn = get_connection()
        cur = conn.cursor()
        try:
            from datetime import datetime, timezone
            now = datetime.now(timezone.utc).isoformat()
            if _is_postgres():
                cur.execute(
                    "UPDATE task_queue SET status = 'cancelled', result = %s, updated_at = %s "
                    "WHERE status IN ('pending', 'running') AND task_type IN ('full_auto', 'no_llm')",
                    ('{"reason":"Остановлено пользователем"}', now))
                cancelled = cur.rowcount
            else:
                cur.execute(
                    "UPDATE task_queue SET status = 'cancelled', result = ?, updated_at = ? "
                    "WHERE status IN ('pending', 'running') AND task_type IN ('full_auto', 'no_llm')",
                    ('{"reason":"Остановлено пользователем"}', now))
                cancelled = cur.rowcount
                conn.commit()
        finally:
            cur.close()
        self._json({"status": "ok", "message": f"Остановлено, отменено задач: {cancelled}"})

    def _get_pipeline_status(self):
        """Возвращает текущий статус пайплайна (running tasks)."""
        conn = get_connection()
        cur = conn.cursor()
        try:
            cur.execute("""
                SELECT task_type, status, COUNT(*) as cnt
                FROM task_queue
                WHERE status IN ('pending', 'running')
                GROUP BY task_type, status
            """)
            if _is_postgres():
                rows = [{"type": r[0], "status": r[1], "count": r[2]} for r in cur.fetchall()]
            else:
                rows = [{"type": r[0], "status": r[1], "count": r[2]} for r in cur.fetchall()]

            # Any running pipeline?
            running = any(r["status"] == "running" for r in rows)
            pending = sum(r["count"] for r in rows if r["status"] == "pending")
            running_count = sum(r["count"] for r in rows if r["status"] == "running")
            # Determine active type
            active_type = ""
            for r in rows:
                if r["status"] == "running" and r["type"] in ("full_auto", "no_llm"):
                    active_type = r["type"]
                    break
            if not active_type:
                for r in rows:
                    if r["status"] == "pending" and r["type"] in ("full_auto", "no_llm"):
                        active_type = r["type"]
                        break

            # Count done tasks and find earliest running start
            total_done = 0
            started_at = ""
            try:
                cur.execute("""
                    SELECT COUNT(*) FROM task_queue
                    WHERE status = 'done' AND task_type IN ('full_auto', 'no_llm')
                """)
                total_done = cur.fetchone()[0]
                cur.execute("""
                    SELECT MIN(created_at) FROM task_queue
                    WHERE status IN ('pending', 'running') AND task_type IN ('full_auto', 'no_llm')
                """)
                row = cur.fetchone()
                started_at = str(row[0]) if row and row[0] else ""
            except Exception:
                pass

            return {
                "running": running or pending > 0,
                "active_type": active_type,
                "running_count": running_count,
                "pending_count": pending,
                "total_done": total_done,
                "started_at": started_at,
                "details": rows,
            }
        finally:
            cur.close()

    def _get_moderation_list(self):
        """Возвращает новости со статусом moderation (с данными локального анализа)."""
        from urllib.parse import parse_qs
        qs = parse_qs(urlparse(self.path).query)
        limit = int(qs.get("limit", ["100"])[0])
        offset = int(qs.get("offset", ["0"])[0])
        source = qs.get("source", [""])[0]
        min_score = int(qs.get("min_score", ["0"])[0])
        q = qs.get("q", [""])[0]

        conn = get_connection()
        cur = conn.cursor()
        ph = "%s" if _is_postgres() else "?"

        conditions = ["n.status = 'moderation'"]
        params = []

        if source:
            conditions.append(f"n.source = {ph}")
            params.append(source)
        if min_score > 0:
            conditions.append(f"COALESCE(na.total_score, 0) >= {ph}")
            params.append(min_score)
        if q:
            conditions.append(f"LOWER(n.title) LIKE {ph}")
            params.append(f"%{q.lower()}%")

        where = " AND ".join(conditions)

        try:
            cur.execute(f"""
                SELECT n.id, n.source, n.title, n.url, n.published_at, n.parsed_at, n.status,
                       n.description,
                       na.total_score, na.quality_score, na.relevance_score,
                       na.freshness_hours, na.viral_score, na.viral_data,
                       na.sentiment_label,
                       na.tags_data as tags, na.entity_names as entities, na.headline_score, na.momentum_score
                FROM news n
                LEFT JOIN news_analysis na ON na.news_id = n.id
                WHERE {where}
                ORDER BY n.parsed_at DESC
                LIMIT {ph} OFFSET {ph}
            """, (*params, limit, offset))

            if _is_postgres():
                columns = [d[0] for d in cur.description]
                rows = [dict(zip(columns, r)) for r in cur.fetchall()]
            else:
                rows = [dict(r) for r in cur.fetchall()]

            # Count total + avg score
            cur.execute(f"SELECT COUNT(*), COALESCE(AVG(na.total_score), 0) FROM news n LEFT JOIN news_analysis na ON na.news_id = n.id WHERE {where}", tuple(params))
            row = cur.fetchone()
            total = row[0]
            avg_score = round(row[1], 1) if row[1] else 0
        finally:
            cur.close()

        return {"status": "ok", "news": rows, "total": total, "avg_score": avg_score}

    def _get_moderation(self, body):
        """POST версия для совместимости."""
        self._json(self._get_moderation_list())

    def _seo_check(self, body):
        """SEO-анализ статьи."""
        from checks.seo_check import analyze_seo
        title = body.get("title", "")
        seo_title = body.get("seo_title", "")
        seo_description = body.get("seo_description", "")
        text = body.get("text", "")
        tags = body.get("tags", [])
        result = analyze_seo(title, seo_title, seo_description, text, tags)
        self._json({"status": "ok", **result})

    def _moderation_rewrite(self, body):
        """Отправляет новости из Модерации на рерайт (без API анализа, только LLM рерайт)."""
        news_ids = body.get("news_ids", [])
        style = body.get("style", "news")
        if not news_ids:
            self._json({"status": "error", "message": "news_ids required"})
            return

        from scheduler import _create_task
        task_ids = []
        conn = get_connection()
        cur = conn.cursor()
        try:
            ph = "%s" if _is_postgres() else "?"
            for nid in news_ids:
                try:
                    cur.execute(f"SELECT title FROM news WHERE id = {ph}", (nid,))
                    row = cur.fetchone()
                    title = (row[0] if _is_postgres() else row["title"]) if row else ""
                except Exception:
                    title = ""
                tid = _create_task("mod_rewrite", nid, title, style)
                task_ids.append(tid)
        finally:
            cur.close()

        # Process rewrites in background (no enrichment, just LLM rewrite)
        import threading
        def _bg_mod_rewrite(ids, tids, rewrite_style):
            from apis.llm import rewrite_news
            from scheduler import _update_task, _fetch_news_by_id, _fetch_analysis_by_id
            from storage.sheets import write_ready_row
            import uuid as _uuid
            import json as _json2
            for nid, tid in zip(ids, tids):
                try:
                    news = _fetch_news_by_id(nid)
                    if not news:
                        _update_task(tid, "error", {"error": "News not found"})
                        continue
                    _update_task(tid, "running", {"stage": "rewriting"})
                    result = rewrite_news(
                        title=news.get("title", ""),
                        text=news.get("plain_text", ""),
                        style=rewrite_style,
                        language="русский",
                    )
                    if result:
                        # Save article to DB
                        conn2 = get_connection()
                        cur2 = conn2.cursor()
                        ph2 = "%s" if _is_postgres() else "?"
                        _now = __import__('datetime').datetime.now(__import__('datetime').timezone.utc).isoformat()
                        aid = str(_uuid.uuid4())[:12]
                        tags_j = _json2.dumps(result.get("tags", []), ensure_ascii=False)
                        try:
                            cur2.execute(f"""INSERT INTO articles (id, news_id, title, text, seo_title, seo_description, tags,
                                style, language, original_title, original_text, source_url, status, created_at, updated_at)
                                VALUES ({','.join([ph2]*15)})""",
                                (aid, nid, result.get("title", ""), result.get("text", ""),
                                 result.get("seo_title", ""), result.get("seo_description", ""), tags_j,
                                 rewrite_style, "русский", news.get("title", ""), (news.get("plain_text", "") or "")[:5000],
                                 news.get("url", ""), "draft", _now, _now))
                            if not _is_postgres():
                                conn2.commit()
                        finally:
                            cur2.close()

                        # Export to Sheets/Ready
                        try:
                            analysis = _fetch_analysis_by_id(nid)
                            write_ready_row(news, analysis, result)
                        except Exception as se:
                            logger.warning("Sheets Ready export failed for %s: %s", nid, se)

                        _update_task(tid, "done", {
                            "stage": "complete",
                            "rewrite_title": result.get("title", "")[:100],
                            "article_id": aid,
                        })
                        update_news_status(nid, "processed")
                    else:
                        _update_task(tid, "error", {"stage": "rewriting", "error": "Rewrite returned None"})
                except Exception as e:
                    _update_task(tid, "error", {"error": str(e)[:500]})

        threading.Thread(
            target=_bg_mod_rewrite,
            args=(list(news_ids), list(task_ids), style),
            daemon=True
        ).start()

        self._json({"status": "ok", "queued": len(news_ids), "task_ids": task_ids})

    def _test_sheets(self, body):
        try:
            import config
            from storage.sheets import _get_client
            client = _get_client()
            if not client:
                self._json({"status": "error", "message": "Google client init failed. Check GOOGLE_SERVICE_ACCOUNT_JSON"})
                return
            sheet = client.open_by_key(config.GOOGLE_SHEETS_ID)
            worksheets = [ws.title for ws in sheet.worksheets()]
            tab = sheet.worksheet(config.SHEETS_TAB)
            rows = len(tab.get_all_values())
            self._json({"status": "ok", "sheets_id": config.GOOGLE_SHEETS_ID, "tabs": worksheets, "active_tab": config.SHEETS_TAB, "rows": rows})
        except Exception as e:
            self._json({"status": "error", "message": str(e), "type": type(e).__name__})

    def _reparse_source(self, body):
        name = body.get("name")
        import config
        source = next((s for s in config.SOURCES if s["name"] == name), None)
        if not source:
            self._json({"status": "error", "message": "Source not found"})
            return
        try:
            if source["type"] == "rss":
                from parsers.rss_parser import parse_rss_source
                count = parse_rss_source(source)
            elif source["type"] == "sitemap":
                from parsers.html_parser import parse_sitemap_source
                count = parse_sitemap_source(source)
            else:
                from parsers.html_parser import parse_html_source
                count = parse_html_source(source)
            self._json({"status": "ok", "new_articles": count})
        except Exception as e:
            self._json({"status": "error", "message": str(e)})

    def _heal_source(self, body):
        """Диагностика и автоматическое лечение проблемного источника."""
        name = body.get("name")
        if not name:
            self._json({"status": "error", "message": "name required"})
            return
        import config
        source = next((s for s in config.SOURCES if s["name"] == name), None)
        if not source:
            self._json({"status": "error", "message": "Source not found"})
            return

        steps = []
        healed = False
        new_articles = 0

        # Step 1: Reset circuit breaker for this domain
        from parsers.proxy import _circuit_breaker, _get_domain
        domain = _get_domain(source["url"])
        if domain in _circuit_breaker:
            del _circuit_breaker[domain]
            steps.append({"action": "circuit_breaker_reset", "status": "ok", "detail": f"Сброшен circuit breaker для {domain}"})
        else:
            steps.append({"action": "circuit_breaker_check", "status": "skip", "detail": "Circuit breaker не активен"})

        # Step 2: Test URL accessibility with fresh UA
        from parsers.proxy import _get_random_ua
        import requests
        test_ok = False
        test_status = 0
        test_error = ""
        try:
            headers = {"User-Agent": _get_random_ua()}
            resp = requests.get(source["url"], headers=headers, timeout=15, allow_redirects=True)
            test_status = resp.status_code
            test_ok = resp.status_code == 200
            content_len = len(resp.text)
            steps.append({"action": "url_test", "status": "ok" if test_ok else "fail",
                          "detail": f"HTTP {test_status}, {content_len} байт" + (", Cloudflare?" if resp.status_code == 403 else "")})
        except Exception as e:
            test_error = str(e)
            steps.append({"action": "url_test", "status": "fail", "detail": f"Ошибка: {test_error[:200]}"})

        # Step 3: Try alternative strategies based on source type
        if test_ok:
            # Step 3a: Try reparse with current config
            try:
                if source["type"] == "rss":
                    from parsers.rss_parser import parse_rss_source
                    new_articles = parse_rss_source(source)
                elif source["type"] == "sitemap":
                    from parsers.html_parser import parse_sitemap_source
                    new_articles = parse_sitemap_source(source)
                else:
                    from parsers.html_parser import parse_html_source
                    new_articles = parse_html_source(source)
                steps.append({"action": "reparse", "status": "ok", "detail": f"Получено {new_articles} новых статей"})
                if new_articles > 0:
                    healed = True
            except Exception as e:
                steps.append({"action": "reparse", "status": "fail", "detail": str(e)[:200]})

            # Step 3b: If HTML source failed — try alternative selectors
            if not healed and source["type"] in ("html", "dtf"):
                alt_selectors = ["article", "div.article", ".news-item", ".post", "a[href*='news']", "a[href*='article']",
                                 ".card", ".feed-item", "h2 a", "h3 a"]
                original_sel = source.get("selector", "article")
                for alt_sel in alt_selectors:
                    if alt_sel == original_sel:
                        continue
                    try:
                        from bs4 import BeautifulSoup
                        soup = BeautifulSoup(resp.text, "lxml")
                        found = soup.select(alt_sel)
                        links_found = sum(1 for el in found if el.find("a", href=True) or el.name == "a")
                        if links_found >= 3:
                            steps.append({"action": "alt_selector", "status": "found",
                                          "detail": f"Селектор '{alt_sel}' нашёл {links_found} элементов (текущий: '{original_sel}')",
                                          "selector": alt_sel, "links_count": links_found})
                            break
                    except Exception:
                        pass
                else:
                    steps.append({"action": "alt_selector", "status": "skip", "detail": "Альтернативные селекторы не помогли"})

            # Step 3c: If RSS source got 0 articles — check feed validity
            if not healed and source["type"] == "rss":
                try:
                    import feedparser
                    feed = feedparser.parse(resp.content)
                    n_entries = len(feed.entries)
                    is_bozo = feed.bozo
                    if n_entries > 0:
                        steps.append({"action": "feed_check", "status": "ok",
                                      "detail": f"RSS валидный: {n_entries} записей" + (" (bozo)" if is_bozo else "")})
                        # All entries might be existing (already parsed)
                        if new_articles == 0:
                            steps.append({"action": "feed_check", "status": "info",
                                          "detail": "Все записи уже в БД — источник работает, новых нет"})
                            healed = True
                    else:
                        steps.append({"action": "feed_check", "status": "fail",
                                      "detail": f"RSS пустой или невалидный" + (f": {feed.bozo_exception}" if is_bozo else "")})
                        # Try common alternative RSS URLs
                        alt_urls = []
                        base = source["url"].rstrip("/")
                        if "/feed/" not in base:
                            alt_urls.append(base + "/feed/")
                        if "/rss" not in base:
                            alt_urls.append(base.rsplit("/", 1)[0] + "/rss/")
                            alt_urls.append(base.rsplit("/", 1)[0] + "/feed.xml")
                        for alt_url in alt_urls:
                            try:
                                alt_resp = requests.get(alt_url, headers=headers, timeout=10)
                                if alt_resp.status_code == 200:
                                    alt_feed = feedparser.parse(alt_resp.content)
                                    if len(alt_feed.entries) > 0:
                                        steps.append({"action": "alt_rss_url", "status": "found",
                                                      "detail": f"Найден рабочий RSS: {alt_url} ({len(alt_feed.entries)} записей)",
                                                      "url": alt_url, "entries": len(alt_feed.entries)})
                                        break
                            except Exception:
                                pass
                except Exception as e:
                    steps.append({"action": "feed_check", "status": "fail", "detail": str(e)[:200]})
        else:
            # URL not accessible — suggest solutions
            if test_status == 403:
                steps.append({"action": "diagnosis", "status": "info", "detail": "403 Forbidden — вероятно Cloudflare/WAF. Рекомендация: прокси или рендер-сервис"})
            elif test_status == 404:
                steps.append({"action": "diagnosis", "status": "info", "detail": "404 Not Found — URL изменился. Проверьте актуальный адрес RSS/страницы"})
            elif test_status == 429:
                steps.append({"action": "diagnosis", "status": "info", "detail": "429 Too Many Requests — увеличьте интервал парсинга"})
            elif test_status >= 500:
                steps.append({"action": "diagnosis", "status": "info", "detail": f"Сервер {test_status} — временная проблема, повторите позже"})
            elif test_error:
                if "timeout" in test_error.lower():
                    steps.append({"action": "diagnosis", "status": "info", "detail": "Таймаут — сервер не отвечает. Попробуйте прокси"})
                elif "ssl" in test_error.lower():
                    steps.append({"action": "diagnosis", "status": "info", "detail": "Ошибка SSL — проблема с сертификатом"})
                else:
                    steps.append({"action": "diagnosis", "status": "info", "detail": f"Сетевая ошибка: {test_error[:150]}"})

        # Step 4: Recommendation
        recommendations = []
        for step in steps:
            if step["status"] == "found" and step["action"] == "alt_selector":
                recommendations.append(f"Сменить селектор на '{step['selector']}' ({step['links_count']} элементов)")
            if step["status"] == "found" and step["action"] == "alt_rss_url":
                recommendations.append(f"Сменить URL на {step['url']}")
        if test_status == 403:
            recommendations.append("Включить прокси-ротацию (PROXY_LIST)")
            recommendations.append("Увеличить интервал парсинга")
        if test_status == 429:
            recommendations.append("Увеличить интервал парсинга до 30+ мин")
        if not healed and not recommendations:
            recommendations.append("Попробуйте парсинг вручную позже")

        self._json({
            "status": "ok",
            "healed": healed,
            "new_articles": new_articles,
            "steps": steps,
            "recommendations": recommendations,
            "source": name,
        })

    def _change_password(self, body):
        username = body.get("username", "")
        password = body.get("password", "")
        if not username or not password:
            self._json({"status": "error", "message": "Username and password required"})
            return
        if username not in USERS:
            self._json({"status": "error", "message": "User not found"})
            return
        USERS[username] = hashlib.sha256(password.encode()).hexdigest()
        self._json({"status": "ok"})

    def _bulk_status(self, body):
        news_ids = body.get("news_ids", [])
        new_status = body.get("status", "")
        if not news_ids or not new_status:
            self._json({"status": "error", "message": "news_ids and status required"})
            return
        for nid in news_ids:
            update_news_status(nid, new_status)
        self._json({"status": "ok", "updated": len(news_ids)})

    def _delete_news(self, body):
        if not self._require_perm("delete"):
            return
        news_ids = body.get("news_ids", [])
        if not news_ids:
            self._json({"status": "error", "message": "news_ids required"})
            return
        conn = get_connection()
        cur = conn.cursor()
        try:
            if _is_postgres():
                placeholders = ",".join(["%s"] * len(news_ids))
            else:
                placeholders = ",".join(["?"] * len(news_ids))
            cur.execute(f"DELETE FROM news_analysis WHERE news_id IN ({placeholders})", tuple(news_ids))
            cur.execute(f"DELETE FROM news WHERE id IN ({placeholders})", tuple(news_ids))
            conn.commit()
            self._json({"status": "ok", "deleted": len(news_ids)})

        finally:
            cur.close()
    def _test_parse(self, body):
        url = body.get("url", "")
        if not url:
            self._json({"status": "error", "message": "URL required"})
            return
        try:
            from parsers.html_parser import _fetch_article
            h1, description, plain_text = _fetch_article(url)
            self._json({
                "status": "ok",
                "h1": h1,
                "description": description[:500],
                "plain_text": plain_text[:1000],
                "text_length": len(plain_text),
            })
        except Exception as e:
            self._json({"status": "error", "message": str(e)})

    def _setup_headers(self, body):
        try:
            from storage.sheets import setup_headers
            setup_headers()
            self._json({"status": "ok"})
        except Exception as e:
            self._json({"status": "error", "message": str(e)})

    def _reparse_all(self, body):
        try:
            import config
            from parsers.rss_parser import parse_rss_source
            from parsers.html_parser import parse_html_source, parse_sitemap_source
            total = 0
            for source in config.SOURCES:
                try:
                    if source["type"] == "rss":
                        total += parse_rss_source(source)
                    elif source["type"] == "sitemap":
                        total += parse_sitemap_source(source)
                    else:
                        total += parse_html_source(source)
                except Exception as e:
                    logger.error("Reparse %s error: %s", source["name"], e)
            self._json({"status": "ok", "new_articles": total})
        except Exception as e:
            self._json({"status": "error", "message": str(e)})

    def _get_sources_stats(self):
        conn = get_connection()
        cur = conn.cursor()
        try:
            cur.execute("SELECT source, COUNT(*) as cnt, MAX(parsed_at) as last_parsed FROM news GROUP BY source ORDER BY cnt DESC")
            if _is_postgres():
                columns = [desc[0] for desc in cur.description]
                return [dict(zip(columns, row)) for row in cur.fetchall()]
            else:
                return [dict(row) for row in cur.fetchall()]

        finally:
            cur.close()
    def _get_db_info(self):
        conn = get_connection()
        cur = conn.cursor()
        try:
            info = {"type": "PostgreSQL" if _is_postgres() else "SQLite"}
            cur.execute("SELECT COUNT(*) FROM news")
            info["total_news"] = cur.fetchone()[0]
            cur.execute("SELECT COUNT(*) FROM news_analysis")
            info["total_analyzed"] = cur.fetchone()[0]
            for status in ["new", "in_review", "approved", "processed", "rejected", "duplicate"]:
                ph = "%s" if _is_postgres() else "?"
                cur.execute(f"SELECT COUNT(*) FROM news WHERE status = {ph}", (status,))
                info[f"status_{status}"] = cur.fetchone()[0]
            cur.execute("SELECT MIN(parsed_at), MAX(parsed_at) FROM news")
            row = cur.fetchone()
            info["oldest"] = str(row[0]) if row[0] else "-"
            info["newest"] = str(row[1]) if row[1] else "-"
            return info

        finally:
            cur.close()
    def _export_sheets_bulk(self, body):
        news_ids = body.get("news_ids", [])
        if not news_ids:
            self._json({"status": "error", "message": "news_ids required"})
            return
        try:
            from storage.sheets import write_news_row
            conn = get_connection()
            cur = conn.cursor()
            try:
                ph = "%s" if _is_postgres() else "?"
                exported = 0
                skipped = 0
                errors = 0
                for nid in news_ids:
                    try:
                        cur.execute(f"SELECT * FROM news WHERE id = {ph}", (nid,))
                        row = cur.fetchone()
                        if not row:
                            continue
                        if _is_postgres():
                            columns = [desc[0] for desc in cur.description]
                            news = dict(zip(columns, row))
                        else:
                            news = dict(row)
                        cur.execute(f"SELECT * FROM news_analysis WHERE news_id = {ph}", (nid,))
                        arow = cur.fetchone()
                        if arow:
                            if _is_postgres():
                                columns = [desc[0] for desc in cur.description]
                                analysis = dict(zip(columns, arow))
                            else:
                                analysis = dict(arow)
                        else:
                            analysis = {"bigrams": "[]", "trends_data": "{}", "keyso_data": "{}",
                                       "llm_recommendation": "", "llm_trend_forecast": "", "llm_merged_with": ""}
                        sheet_row = write_news_row(news, analysis)
                        if sheet_row and sheet_row > 0:
                            exported += 1
                        elif sheet_row == -1:
                            skipped += 1
                    except Exception as e:
                        logger.warning("Bulk export error for %s: %s", nid, e)
                        errors += 1
                self._json({"status": "ok", "exported": exported, "skipped": skipped, "errors": errors})
            finally:
                cur.close()
        except Exception as e:
            self._json({"status": "error", "message": str(e)})

    def _rewrite_news(self, body):
        news_id = body.get("news_id")
        style = body.get("style", "news")
        language = body.get("language", "русский")
        try:
            conn = get_connection()
            cur = conn.cursor()
            try:
                ph = "%s" if _is_postgres() else "?"
                cur.execute(f"SELECT title, plain_text, description FROM news WHERE id = {ph}", (news_id,))
                row = cur.fetchone()
                if not row:
                    self._json({"status": "error", "message": "News not found"})
                    return
                if _is_postgres():
                    columns = [desc[0] for desc in cur.description]
                    news = dict(zip(columns, row))
                else:
                    news = dict(row)
                title = news.get("title", "")
                text = news.get("plain_text", "") or news.get("description", "")
                from apis.llm import rewrite_news
                result = rewrite_news(title, text, style, language)
                if result:
                    self._json({"status": "ok", "result": result, "original_title": title})
                else:
                    self._json({"status": "error", "message": "LLM returned no result"})
            finally:
                cur.close()
        except Exception as e:
            self._json({"status": "error", "message": str(e)})

    def _merge_news(self, body):
        news_ids = body.get("news_ids", [])
        if len(news_ids) < 2:
            self._json({"status": "error", "message": "Need at least 2 news to merge"})
            return
        try:
            conn = get_connection()
            cur = conn.cursor()
            try:
                ph = "%s" if _is_postgres() else "?"
                placeholders = ",".join([ph] * len(news_ids))
                cur.execute(f"SELECT id, source, title, plain_text FROM news WHERE id IN ({placeholders})", news_ids)
                if _is_postgres():
                    columns = [desc[0] for desc in cur.description]
                    news_list = [dict(zip(columns, row)) for row in cur.fetchall()]
                else:
                    news_list = [dict(row) for row in cur.fetchall()]
                from apis.llm import merge_news
                result = merge_news(news_list)
                if result:
                    self._json({"status": "ok", "result": result, "sources": [n["source"] for n in news_list]})
                else:
                    self._json({"status": "error", "message": "LLM returned no result"})
            finally:
                cur.close()
        except Exception as e:
            self._json({"status": "error", "message": str(e)})

    def _news_detail(self, body):
        news_id = body.get("news_id")
        if not news_id:
            self._json({"status": "error", "message": "news_id required"})
            return
        conn = get_connection()
        cur = conn.cursor()
        try:
            ph = "%s" if _is_postgres() else "?"
            cur.execute(f"SELECT * FROM news WHERE id = {ph}", (news_id,))
            row = cur.fetchone()
            if not row:
                self._json({"status": "error", "message": "Not found"})
                return
            if _is_postgres():
                columns = [desc[0] for desc in cur.description]
                news = dict(zip(columns, row))
            else:
                news = dict(row)
            # Get analysis too
            cur.execute(f"SELECT * FROM news_analysis WHERE news_id = {ph}", (news_id,))
            arow = cur.fetchone()
            analysis = None
            if arow:
                if _is_postgres():
                    columns = [desc[0] for desc in cur.description]
                    analysis = dict(zip(columns, arow))
                else:
                    analysis = dict(arow)
            self._json({"status": "ok", "news": news, "analysis": analysis})

        finally:
            cur.close()
    def _analyze_news(self, body):
        """Полный анализ одной новости: viral, freshness, quality, relevance, sentiment, tags, trends, keyso."""
        news_id = body.get("news_id")
        if not news_id:
            self._json({"status": "error", "message": "news_id required"})
            return
        conn = get_connection()
        cur = conn.cursor()
        try:
            ph = "%s" if _is_postgres() else "?"
            cur.execute(f"SELECT * FROM news WHERE id = {ph}", (news_id,))
            row = cur.fetchone()
            if not row:
                self._json({"status": "error", "message": "Not found"})
                return
            if _is_postgres():
                columns = [desc[0] for desc in cur.description]
                news = dict(zip(columns, row))
            else:
                news = dict(row)

            from checks.viral_score import viral_score
            from checks.freshness import check_freshness
            from checks.quality import check_quality
            from checks.relevance import check_relevance
            from checks.sentiment import analyze_sentiment
            from checks.tags import auto_tag
            from checks.momentum import get_momentum

            result = {
                "viral": viral_score(news),
                "freshness": check_freshness(news),
                "quality": check_quality(news),
                "relevance": check_relevance(news),
                "sentiment": analyze_sentiment(news),
                "tags": auto_tag(news),
                "momentum": get_momentum(news),
            }

            # Get existing analysis data (trends, keyso, llm)
            cur.execute(f"SELECT * FROM news_analysis WHERE news_id = {ph}", (news_id,))
            arow = cur.fetchone()
            if arow:
                if _is_postgres():
                    columns = [desc[0] for desc in cur.description]
                    analysis = dict(zip(columns, arow))
                else:
                    analysis = dict(arow)
                import json
                try:
                    result["trends_data"] = json.loads(analysis.get("trends_data", "{}"))
                except Exception:
                    result["trends_data"] = {}
                try:
                    result["keyso_data"] = json.loads(analysis.get("keyso_data", "{}"))
                except Exception:
                    result["keyso_data"] = {}
                result["llm_recommendation"] = analysis.get("llm_recommendation", "")
                result["llm_trend_forecast"] = analysis.get("llm_trend_forecast", "")
                try:
                    result["bigrams"] = json.loads(analysis.get("bigrams", "[]"))
                except Exception:
                    result["bigrams"] = []
            else:
                result["trends_data"] = {}
                result["keyso_data"] = {}
                result["llm_recommendation"] = ""
                result["llm_trend_forecast"] = ""
                result["bigrams"] = []

            total = sum(result[k]["score"] for k in ("viral", "freshness", "quality", "relevance")) // 4
            result["total_score"] = min(100, total + result["momentum"]["score"] // 5)
            self._json({"status": "ok", "analysis": result})

        finally:
            cur.close()
    def _batch_rewrite(self, body):
        """Батч-переписка новостей: создаёт статьи из списка news_ids."""
        news_ids = body.get("news_ids", [])
        style = body.get("style", "news")
        language = body.get("language", "русский")
        if not news_ids:
            self._json({"status": "error", "message": "news_ids required"})
            return
        conn = get_connection()
        cur = conn.cursor()
        try:
            ph = "%s" if _is_postgres() else "?"
            from apis.llm import rewrite_news
            import uuid, json
            from datetime import datetime, timezone

            results = []
            for nid in news_ids:
                try:
                    cur.execute(f"SELECT id, title, plain_text, description, url, source FROM news WHERE id = {ph}", (nid,))
                    row = cur.fetchone()
                    if not row:
                        results.append({"news_id": nid, "ok": False, "error": "not found"})
                        continue
                    if _is_postgres():
                        columns = [desc[0] for desc in cur.description]
                        news = dict(zip(columns, row))
                    else:
                        news = dict(row)
                    title = news.get("title", "")
                    text = news.get("plain_text", "") or news.get("description", "")
                    result = rewrite_news(title, text, style, language)
                    if not result:
                        results.append({"news_id": nid, "ok": False, "error": "LLM failed"})
                        continue
                    # Save as article
                    aid = str(uuid.uuid4())[:12]
                    now = datetime.now(timezone.utc).isoformat()
                    tags = json.dumps(result.get("tags", []), ensure_ascii=False)
                    cur.execute(f"""INSERT INTO articles (id, news_id, title, text, seo_title, seo_description, tags,
                        style, language, original_title, original_text, source_url, status, created_at, updated_at)
                        VALUES ({','.join([ph]*15)})""",
                        (aid, nid, result.get("title", ""), result.get("text", ""),
                         result.get("seo_title", ""), result.get("seo_description", ""), tags,
                         style, language, title, text[:5000],
                         news.get("url", ""), "draft", now, now))
                    if not _is_postgres():
                        conn.commit()
                    results.append({"news_id": nid, "ok": True, "article_id": aid, "title": result.get("title", "")})
                except Exception as e:
                    logger.warning("Batch rewrite error for %s: %s", nid, e)
                    results.append({"news_id": nid, "ok": False, "error": str(e)})

            ok_count = sum(1 for r in results if r.get("ok"))
            self._json({"status": "ok", "total": len(news_ids), "success": ok_count,
                         "failed": len(news_ids) - ok_count, "results": results})

        finally:
            cur.close()
    # ---- Analytics methods ----

    def _get_analytics(self):
        from api.analytics import get_analytics
        return get_analytics()
    def _get_prompt_versions(self):
        conn = get_connection()
        cur = conn.cursor()
        try:
            cur.execute("SELECT * FROM prompt_versions ORDER BY prompt_name, version DESC")
            if _is_postgres():
                columns = [desc[0] for desc in cur.description]
                rows = [dict(zip(columns, row)) for row in cur.fetchall()]
            else:
                rows = [dict(row) for row in cur.fetchall()]
            return {"status": "ok", "versions": rows}

        finally:
            cur.close()
    def _save_prompt_version(self, body):
        import uuid
        from datetime import datetime, timezone
        conn = get_connection()
        cur = conn.cursor()
        try:
            ph = "%s" if _is_postgres() else "?"
            name = body.get("prompt_name", "")
            content = body.get("content", "")
            notes = body.get("notes", "")
            if not name or not content:
                self._json({"status": "error", "message": "name and content required"})
                return
            # Get next version
            cur.execute(f"SELECT MAX(version) as mv FROM prompt_versions WHERE prompt_name = {ph}", (name,))
            row = cur.fetchone()
            if _is_postgres():
                max_v = row[0] if row and row[0] else 0
            else:
                max_v = row["mv"] if row and row["mv"] else 0
            if max_v is None:
                max_v = 0
            version = max_v + 1
            vid = str(uuid.uuid4())[:12]
            now = datetime.now(timezone.utc).isoformat()
            cur.execute(f"""INSERT INTO prompt_versions (id, prompt_name, version, content, is_active, created_at, notes)
                VALUES ({','.join([ph]*7)})""", (vid, name, version, content, 0, now, notes))
            if not _is_postgres():
                conn.commit()
            self._json({"status": "ok", "id": vid, "version": version})

        finally:
            cur.close()
    def _activate_prompt_version(self, body):
        conn = get_connection()
        cur = conn.cursor()
        try:
            ph = "%s" if _is_postgres() else "?"
            vid = body.get("id", "")
            if not vid:
                self._json({"status": "error", "message": "id required"})
                return
            # Get prompt name and content
            cur.execute(f"SELECT prompt_name, content FROM prompt_versions WHERE id = {ph}", (vid,))
            row = cur.fetchone()
            if not row:
                self._json({"status": "error", "message": "not found"})
                return
            if _is_postgres():
                name, content = row[0], row[1]
            else:
                name, content = row["prompt_name"], row["content"]
            # Deactivate all for this name
            cur.execute(f"UPDATE prompt_versions SET is_active = 0 WHERE prompt_name = {ph}", (name,))
            # Activate this one
            cur.execute(f"UPDATE prompt_versions SET is_active = 1 WHERE id = {ph}", (vid,))
            if not _is_postgres():
                conn.commit()
            # Apply to live prompts
            import apis.llm as llm
            prompt_map = {
                "trend_forecast": "PROMPT_TREND_FORECAST",
                "merge_analysis": "PROMPT_MERGE_ANALYSIS",
                "keyso_queries": "PROMPT_KEYSO_QUERIES",
                "rewrite": "PROMPT_REWRITE",
            }
            attr = prompt_map.get(name)
            if attr and hasattr(llm, attr):
                setattr(llm, attr, content)
                logger.info("Activated prompt version %s for %s", vid, name)
            self._json({"status": "ok", "prompt_name": name, "applied": bool(attr)})

        finally:
            cur.close()
    def _generate_digest(self, body):
        """Генерирует дайджест за указанный период."""
        period = body.get("period", "today")  # today, week
        conn = get_connection()
        cur = conn.cursor()
        try:
            ph = "%s" if _is_postgres() else "?"
            if period == "week":
                interval = "7 days"
            else:
                interval = "1 day"
            if _is_postgres():
                cur.execute(f"SELECT id, title, source, url FROM news WHERE status IN ('approved', 'processed') AND parsed_at::timestamptz > (NOW() - INTERVAL '{interval}') ORDER BY parsed_at DESC LIMIT 30")
                columns = [desc[0] for desc in cur.description]
                news_list = [dict(zip(columns, row)) for row in cur.fetchall()]
            else:
                cur.execute(f"SELECT id, title, source, url FROM news WHERE status IN ('approved', 'processed') AND parsed_at > datetime('now', '-{interval}') ORDER BY parsed_at DESC LIMIT 30")
                news_list = [dict(row) for row in cur.fetchall()]

            if not news_list:
                self._json({"status": "ok", "digest": {"title": "Нет данных", "summary": "Нет одобренных новостей за выбранный период.", "top_news": [], "trends": []}, "news_count": 0})
                return

            from apis.llm import _call_llm
            news_text = "\n".join(f"- [{n['source']}] {n['title']}" for n in news_list)
            period_label = 'неделю' if period == 'week' else 'день'
            prompt = f"""Ты — главный редактор крупного игрового портала. Составь профессиональный дайджест «Главное за {period_label}» из новостей ниже.
    
    ## Новости ({len(news_list)} шт.):
    {news_text}
    
    ## Правила:
    1. title — яркий заголовок дайджеста (напр. «Игровой дайджест: GTA 6, новый патч Elden Ring и скандал вокруг Ubisoft»)
    2. summary — связный текст на 4-6 предложений, охватывающий самые значимые события, не простое перечисление
    3. top_news — 3-5 самых важных новостей, одной фразой каждая (не копируй заголовки дословно, перефразируй)
    4. trends — 2-3 тенденции, которые прослеживаются в потоке новостей (напр. «Рост интереса к ретро-играм», «Волна переносов релизов»)
    5. Язык: русский
    
    Ответь строго JSON без markdown:
    {{
      "title": "Заголовок дайджеста",
      "summary": "Связный обзорный текст",
      "top_news": ["Ключевая новость 1", "Ключевая новость 2", "Ключевая новость 3"],
      "trends": ["Тенденция 1", "Тенденция 2"]
    }}"""
            result = _call_llm(prompt)
            if result:
                self._json({"status": "ok", "digest": result, "news_count": len(news_list)})
            else:
                self._json({"status": "error", "message": "LLM failed"})

        finally:
            cur.close()

    def _get_digests(self):
        """Возвращает последние 10 сохранённых дайджестов."""
        from storage.database import get_digests
        digests = get_digests(limit=10)
        return {"status": "ok", "digests": digests}

    def _generate_and_save_digest(self, body):
        """Ручная генерация дайджеста с сохранением в БД."""
        style = body.get("style", "brief")
        conn = get_connection()
        cur = conn.cursor()
        try:
            ph = "%s" if _is_postgres() else "?"
            if _is_postgres():
                cur.execute("""
                    SELECT n.id, n.title, n.source, n.url,
                           COALESCE(a.total_score, 0) as total_score
                    FROM news n
                    LEFT JOIN news_analysis a ON a.news_id = n.id
                    WHERE n.status IN ('approved', 'processed', 'in_review', 'ready')
                      AND n.parsed_at::timestamptz > (NOW() - INTERVAL '24 hours')
                    ORDER BY COALESCE(a.total_score, 0) DESC
                    LIMIT 20
                """)
                columns = [desc[0] for desc in cur.description]
                news_list = [dict(zip(columns, row)) for row in cur.fetchall()]
            else:
                cur.execute("""
                    SELECT n.id, n.title, n.source, n.url,
                           COALESCE(a.total_score, 0) as total_score
                    FROM news n
                    LEFT JOIN news_analysis a ON a.news_id = n.id
                    WHERE n.status IN ('approved', 'processed', 'in_review', 'ready')
                      AND n.parsed_at > datetime('now', '-1 day')
                    ORDER BY COALESCE(a.total_score, 0) DESC
                    LIMIT 20
                """)
                news_list = [dict(row) for row in cur.fetchall()]

            if not news_list:
                self._json({"status": "ok", "digest": {"title": "Нет данных", "text": "Нет новостей за последние 24 часа.", "news_count": 0}})
                return

            from apis.digest import generate_daily_digest
            result = generate_daily_digest(news_list, style=style)

            # Save to DB
            import uuid
            from datetime import datetime, timezone
            from storage.database import save_digest
            digest_id = str(uuid.uuid4())[:12]
            digest_date = datetime.now(timezone.utc).strftime("%Y-%m-%d")
            save_digest(
                digest_id=digest_id,
                digest_date=digest_date,
                style=style,
                title=result.get("title", ""),
                text=result.get("text", ""),
                news_count=result.get("news_count", 0),
            )

            self._json({"status": "ok", "digest": result})

        except Exception as e:
            self._json({"status": "error", "message": str(e)[:500]})
        finally:
            cur.close()

    def _get_event_chain(self, body):
        news_id = body.get("news_id", "")
        if not news_id:
            self._json({"status": "error", "message": "news_id required"})
            return
        conn = get_connection()
        cur = conn.cursor()
        try:
            ph = "%s" if _is_postgres() else "?"
            cur.execute(f"SELECT * FROM news WHERE id = {ph}", (news_id,))
            if _is_postgres():
                columns = [desc[0] for desc in cur.description]
                row = cur.fetchone()
                if not row:
                    self._json({"status": "error", "message": "not found"})
                    return
                news = dict(zip(columns, row))
            else:
                row = cur.fetchone()
                if not row:
                    self._json({"status": "error", "message": "not found"})
                    return
                news = dict(row)
            from checks.temporal_clusters import get_event_chain
            chain = get_event_chain(news)
            self._json({"status": "ok", **chain})

        finally:
            cur.close()
    # ---- Queue methods ----

    def _get_queue(self):
        conn = get_connection()
        cur = conn.cursor()
        try:
            q = "SELECT * FROM task_queue ORDER BY created_at DESC LIMIT 200"
            cur.execute(q)
            if _is_postgres():
                columns = [desc[0] for desc in cur.description]
                rows = [dict(zip(columns, row)) for row in cur.fetchall()]
            else:
                rows = [dict(row) for row in cur.fetchall()]
            return {"status": "ok", "tasks": rows}

        finally:
            cur.close()
    def _cancel_queue_task(self, body):
        task_id = body.get("task_id")
        if not task_id:
            self._json({"status": "error", "message": "task_id required"})
            return
        conn = get_connection()
        cur = conn.cursor()
        try:
            ph = "%s" if _is_postgres() else "?"
            now = __import__('datetime').datetime.now(__import__('datetime').timezone.utc).isoformat()
            cur.execute(f"UPDATE task_queue SET status = 'cancelled', updated_at = {ph} WHERE id = {ph} AND status = 'pending'", (now, task_id))
            if not _is_postgres():
                conn.commit()
            self._json({"status": "ok"})

        finally:
            cur.close()
    def _cancel_all_queue(self, body):
        task_type = body.get("task_type", "")
        conn = get_connection()
        cur = conn.cursor()
        try:
            ph = "%s" if _is_postgres() else "?"
            now = __import__('datetime').datetime.now(__import__('datetime').timezone.utc).isoformat()
            if task_type:
                cur.execute(f"UPDATE task_queue SET status = 'cancelled', updated_at = {ph} WHERE status = 'pending' AND task_type = {ph}", (now, task_type))
            else:
                cur.execute(f"UPDATE task_queue SET status = 'cancelled', updated_at = {ph} WHERE status = 'pending'", (now,))
            if not _is_postgres():
                conn.commit()
            self._json({"status": "ok"})

        finally:
            cur.close()
    def _clear_done_queue(self, body):
        conn = get_connection()
        cur = conn.cursor()
        try:
            cur.execute("DELETE FROM task_queue WHERE status IN ('done', 'cancelled', 'skipped', 'error')")
            if not _is_postgres():
                conn.commit()
            self._json({"status": "ok"})

        finally:
            cur.close()
    def _retry_queue_tasks(self, body):
        """Повторяет выбранные задачи из очереди (error → pending)."""
        task_ids = body.get("task_ids", [])
        if not task_ids:
            self._json({"status": "error", "message": "task_ids required"})
            return
        conn = get_connection()
        cur = conn.cursor()
        try:
            ph = "%s" if _is_postgres() else "?"
            from datetime import datetime, timezone
            now = datetime.now(timezone.utc).isoformat()
            count = 0
            for tid in task_ids:
                cur.execute(f"UPDATE task_queue SET status = 'pending', result = '', updated_at = {ph} WHERE id = {ph} AND status IN ('error', 'cancelled', 'skipped', 'done')", (now, tid))
                count += cur.rowcount if hasattr(cur, 'rowcount') else 1
            if not _is_postgres():
                conn.commit()

            # Re-run pending tasks in background
            cur.execute("SELECT id, task_type, news_id, news_title, style FROM task_queue WHERE status = 'pending' ORDER BY created_at")
            if _is_postgres():
                cols = [d[0] for d in cur.description]
                pending = [dict(zip(cols, r)) for r in cur.fetchall()]
            else:
                pending = [dict(r) for r in cur.fetchall()]

            if pending:
                import threading
                no_llm_tasks = [t for t in pending if t["task_type"] == "no_llm"]
                full_auto_tasks = [t for t in pending if t["task_type"] == "full_auto"]

                if no_llm_tasks:
                    nids = [t["news_id"] for t in no_llm_tasks]
                    tids = [t["id"] for t in no_llm_tasks]
                    from scheduler import run_no_llm_pipeline
                    threading.Thread(target=run_no_llm_pipeline, args=(nids, tids), daemon=True).start()
                if full_auto_tasks:
                    nids = [t["news_id"] for t in full_auto_tasks]
                    tids = [t["id"] for t in full_auto_tasks]
                    from scheduler import run_full_auto_pipeline
                    threading.Thread(target=run_full_auto_pipeline, args=(nids, tids), daemon=True).start()

            self._json({"status": "ok", "retried": count})
        finally:
            cur.close()

    def _get_viral_triggers(self):
        """Возвращает все триггеры виральности (дефолтные + кастомные из БД)."""
        from checks.viral_score import VIRAL_TRIGGERS
        conn = get_connection()
        cur = conn.cursor()
        try:
            # Load DB overrides
            db_triggers = {}
            try:
                cur.execute("SELECT trigger_id, label, weight, keywords, is_active, is_custom FROM viral_triggers_config")
                for row in cur.fetchall():
                    if _is_postgres():
                        tid, label, weight, kw_json, active, custom = row
                    else:
                        tid, label, weight, kw_json, active, custom = row["trigger_id"], row["label"], row["weight"], row["keywords"], row["is_active"], row["is_custom"]
                    import json as _j
                    kws = _j.loads(kw_json) if isinstance(kw_json, str) else (kw_json or [])
                    db_triggers[tid] = {"label": label, "weight": weight, "keywords": kws, "is_active": bool(active), "is_custom": bool(custom)}
            except Exception:
                pass

            result = []
            # Default triggers
            for tid, tdata in VIRAL_TRIGGERS.items():
                if tid in db_triggers:
                    dt = db_triggers[tid]
                    result.append({"id": tid, "label": dt["label"], "weight": dt["weight"], "keywords": dt["keywords"], "is_active": dt["is_active"], "is_custom": False, "modified": True})
                else:
                    result.append({"id": tid, "label": tdata["label"], "weight": tdata["weight"], "keywords": tdata["keywords"], "is_active": True, "is_custom": False, "modified": False})

            # Custom-only triggers (not in defaults)
            for tid, dt in db_triggers.items():
                if tid not in VIRAL_TRIGGERS:
                    result.append({"id": tid, "label": dt["label"], "weight": dt["weight"], "keywords": dt["keywords"], "is_active": dt["is_active"], "is_custom": True, "modified": False})

            result.sort(key=lambda x: (-x["weight"], x["label"]))
            return {"triggers": result, "total": len(result)}
        finally:
            cur.close()

    def _save_viral_trigger(self, body):
        """Сохраняет/обновляет триггер виральности."""
        trigger_id = body.get("trigger_id", "").strip()
        label = body.get("label", "").strip()
        weight = int(body.get("weight", 0))
        keywords = body.get("keywords", [])
        is_active = bool(body.get("is_active", True))
        is_custom = bool(body.get("is_custom", False))

        if not trigger_id or not label:
            self._json({"status": "error", "message": "trigger_id and label required"})
            return

        if isinstance(keywords, str):
            keywords = [k.strip() for k in keywords.split(",") if k.strip()]

        import json as _j
        from datetime import datetime, timezone
        conn = get_connection()
        cur = conn.cursor()
        try:
            ph = "%s" if _is_postgres() else "?"
            now = datetime.now(timezone.utc).isoformat()
            kw_json = _j.dumps(keywords, ensure_ascii=False)

            if _is_postgres():
                cur.execute(f"""
                    INSERT INTO viral_triggers_config (trigger_id, label, weight, keywords, is_active, is_custom, updated_at)
                    VALUES ({','.join([ph]*7)})
                    ON CONFLICT (trigger_id) DO UPDATE SET label={ph}, weight={ph}, keywords={ph}, is_active={ph}, updated_at={ph}
                """, (trigger_id, label, weight, kw_json, 1 if is_active else 0, 1 if is_custom else 0, now,
                      label, weight, kw_json, 1 if is_active else 0, now))
            else:
                cur.execute(f"INSERT OR REPLACE INTO viral_triggers_config (trigger_id, label, weight, keywords, is_active, is_custom, updated_at) VALUES ({','.join([ph]*7)})",
                            (trigger_id, label, weight, kw_json, 1 if is_active else 0, 1 if is_custom else 0, now))
                conn.commit()

            # Rebuild index
            from checks.viral_score import reload_viral_triggers
            reload_viral_triggers()

            self._json({"status": "ok", "trigger_id": trigger_id})
        finally:
            cur.close()

    def _delete_viral_trigger(self, body):
        """Удаляет кастомный триггер или сбрасывает изменённый дефолтный."""
        trigger_id = body.get("trigger_id", "")
        if not trigger_id:
            self._json({"status": "error", "message": "trigger_id required"})
            return
        conn = get_connection()
        cur = conn.cursor()
        try:
            ph = "%s" if _is_postgres() else "?"
            cur.execute(f"DELETE FROM viral_triggers_config WHERE trigger_id = {ph}", (trigger_id,))
            if not _is_postgres():
                conn.commit()
            from checks.viral_score import reload_viral_triggers
            reload_viral_triggers()
            self._json({"status": "ok"})
        finally:
            cur.close()

    def _queue_batch_rewrite(self, body):
        """Ставит новости в очередь на переписку и запускает обработку в фоне."""
        news_ids = body.get("news_ids", [])
        style = body.get("style", "news")
        language = body.get("language", "русский")
        if not news_ids:
            self._json({"status": "error", "message": "news_ids required"})
            return

        import uuid
        from datetime import datetime, timezone
        conn = get_connection()
        cur = conn.cursor()
        try:
            ph = "%s" if _is_postgres() else "?"
            now = datetime.now(timezone.utc).isoformat()
            created = []

            for nid in news_ids:
                cur.execute(f"SELECT title FROM news WHERE id = {ph}", (nid,))
                row = cur.fetchone()
                title = ""
                if row:
                    title = row[0] if _is_postgres() else row["title"]
                tid = str(uuid.uuid4())[:12]
                cur.execute(f"""INSERT INTO task_queue (id, task_type, news_id, news_title, style, status, created_at, updated_at)
                    VALUES ({','.join([ph]*8)})""",
                    (tid, "rewrite", nid, title[:200], style, "pending", now, now))
                created.append(tid)

            if not _is_postgres():
                conn.commit()

            # Process in background thread
            def _process_rewrite_queue():
                import json as _json
                from apis.llm import rewrite_news
                conn2 = get_connection()
                cur2 = conn2.cursor()
                try:
                    for tid in created:
                        cur2.execute(f"SELECT * FROM task_queue WHERE id = {ph}", (tid,))
                        if _is_postgres():
                            cols = [d[0] for d in cur2.description]
                            task = dict(zip(cols, cur2.fetchone()))
                        else:
                            task = dict(cur2.fetchone())
                        if task["status"] != "pending":
                            continue
                        nid = task["news_id"]
                        _now = datetime.now(timezone.utc).isoformat()
                        cur2.execute(f"UPDATE task_queue SET status = 'processing', updated_at = {ph} WHERE id = {ph}", (_now, tid))
                        if not _is_postgres():
                            conn2.commit()
                        try:
                            cur2.execute(f"SELECT id, title, plain_text, description, url, source FROM news WHERE id = {ph}", (nid,))
                            row = cur2.fetchone()
                            if not row:
                                raise Exception("news not found")
                            if _is_postgres():
                                cols = [d[0] for d in cur2.description]
                                news = dict(zip(cols, row))
                            else:
                                news = dict(row)
                            ntitle = news.get("title", "")
                            ntext = news.get("plain_text", "") or news.get("description", "")
                            result = rewrite_news(ntitle, ntext, style, language)
                            if not result:
                                raise Exception("LLM failed")
                            aid = str(uuid.uuid4())[:12]
                            tags = _json.dumps(result.get("tags", []), ensure_ascii=False)
                            cur2.execute(f"""INSERT INTO articles (id, news_id, title, text, seo_title, seo_description, tags,
                                style, language, original_title, original_text, source_url, status, created_at, updated_at)
                                VALUES ({','.join([ph]*15)})""",
                                (aid, nid, result.get("title", ""), result.get("text", ""),
                                 result.get("seo_title", ""), result.get("seo_description", ""), tags,
                                 style, language, ntitle, ntext[:5000],
                                 news.get("url", ""), "draft", _now, _now))
                            if not _is_postgres():
                                conn2.commit()

                            # Export to Sheets/Ready
                            try:
                                from storage.sheets import write_ready_row
                                # Fetch full news + analysis for Sheets
                                cur2.execute(f"SELECT * FROM news WHERE id = {ph}", (nid,))
                                if _is_postgres():
                                    nc = [d[0] for d in cur2.description]
                                    full_news = dict(zip(nc, cur2.fetchone()))
                                else:
                                    full_news = dict(cur2.fetchone())
                                cur2.execute(f"SELECT * FROM news_analysis WHERE news_id = {ph}", (nid,))
                                arow = cur2.fetchone()
                                analysis = None
                                if arow:
                                    if _is_postgres():
                                        ac = [d[0] for d in cur2.description]
                                        analysis = dict(zip(ac, arow))
                                    else:
                                        analysis = dict(arow)
                                write_ready_row(full_news, analysis, result)
                            except Exception as sheets_err:
                                logger.warning("Sheets Ready export failed for %s: %s", nid, sheets_err)

                            _now2 = datetime.now(timezone.utc).isoformat()
                            res_data = _json.dumps({"article_id": aid, "title": result.get("title", "")}, ensure_ascii=False)
                            cur2.execute(f"UPDATE task_queue SET status = 'done', result = {ph}, updated_at = {ph} WHERE id = {ph}", (res_data, _now2, tid))
                            if not _is_postgres():
                                conn2.commit()
                        except Exception as e:
                            logger.warning("Queue rewrite error %s: %s", tid, e)
                            _now2 = datetime.now(timezone.utc).isoformat()
                            cur2.execute(f"UPDATE task_queue SET status = 'error', result = {ph}, updated_at = {ph} WHERE id = {ph}", (str(e), _now2, tid))
                            if not _is_postgres():
                                conn2.commit()

                finally:
                    cur2.close()
            t = threading.Thread(target=_process_rewrite_queue, daemon=True)
            t.start()
            self._json({"status": "ok", "queued": len(created), "task_ids": created})

        finally:
            cur.close()
    def _queue_sheets_export(self, body):
        """Ставит новости в очередь на экспорт в Sheets и запускает обработку в фоне."""
        news_ids = body.get("news_ids", [])
        if not news_ids:
            self._json({"status": "error", "message": "news_ids required"})
            return

        import uuid
        from datetime import datetime, timezone
        conn = get_connection()
        cur = conn.cursor()
        try:
            ph = "%s" if _is_postgres() else "?"
            now = datetime.now(timezone.utc).isoformat()
            created = []

            for nid in news_ids:
                cur.execute(f"SELECT title FROM news WHERE id = {ph}", (nid,))
                row = cur.fetchone()
                title = ""
                if row:
                    title = row[0] if _is_postgres() else row["title"]
                tid = str(uuid.uuid4())[:12]
                cur.execute(f"""INSERT INTO task_queue (id, task_type, news_id, news_title, style, status, created_at, updated_at)
                    VALUES ({','.join([ph]*8)})""",
                    (tid, "sheets", nid, title[:200], "", "pending", now, now))
                created.append(tid)

            if not _is_postgres():
                conn.commit()

            # Process in background
            def _process_sheets_queue():
                import json as _json
                from storage.sheets import write_news_row
                conn2 = get_connection()
                cur2 = conn2.cursor()
                try:
                    for tid in created:
                        cur2.execute(f"SELECT * FROM task_queue WHERE id = {ph}", (tid,))
                        if _is_postgres():
                            cols = [d[0] for d in cur2.description]
                            task = dict(zip(cols, cur2.fetchone()))
                        else:
                            task = dict(cur2.fetchone())
                        if task["status"] != "pending":
                            continue
                        nid = task["news_id"]
                        _now = datetime.now(timezone.utc).isoformat()
                        cur2.execute(f"UPDATE task_queue SET status = 'processing', updated_at = {ph} WHERE id = {ph}", (_now, tid))
                        if not _is_postgres():
                            conn2.commit()
                        try:
                            cur2.execute(f"SELECT * FROM news WHERE id = {ph}", (nid,))
                            row = cur2.fetchone()
                            if not row:
                                raise Exception("news not found")
                            if _is_postgres():
                                cols = [d[0] for d in cur2.description]
                                news = dict(zip(cols, row))
                            else:
                                news = dict(row)
                            cur2.execute(f"SELECT * FROM news_analysis WHERE news_id = {ph}", (nid,))
                            arow = cur2.fetchone()
                            if arow:
                                if _is_postgres():
                                    cols = [d[0] for d in cur2.description]
                                    analysis = dict(zip(cols, arow))
                                else:
                                    analysis = dict(arow)
                            else:
                                analysis = {"bigrams": "[]", "trends_data": "{}", "keyso_data": "{}",
                                           "llm_recommendation": "", "llm_trend_forecast": "", "llm_merged_with": ""}
                            sheet_row = write_news_row(news, analysis)
                            _now2 = datetime.now(timezone.utc).isoformat()
                            if sheet_row and sheet_row > 0:
                                res_data = _json.dumps({"row": sheet_row}, ensure_ascii=False)
                                cur2.execute(f"UPDATE task_queue SET status = 'done', result = {ph}, updated_at = {ph} WHERE id = {ph}", (res_data, _now2, tid))
                            elif sheet_row == -1:
                                cur2.execute(f"UPDATE task_queue SET status = 'skipped', result = 'duplicate', updated_at = {ph} WHERE id = {ph}", (_now2, tid))
                            else:
                                cur2.execute(f"UPDATE task_queue SET status = 'error', result = 'no row', updated_at = {ph} WHERE id = {ph}", (_now2, tid))
                            if not _is_postgres():
                                conn2.commit()
                        except Exception as e:
                            logger.warning("Queue sheets error %s: %s", tid, e)
                            _now2 = datetime.now(timezone.utc).isoformat()
                            cur2.execute(f"UPDATE task_queue SET status = 'error', result = {ph}, updated_at = {ph} WHERE id = {ph}", (str(e), _now2, tid))
                            if not _is_postgres():
                                conn2.commit()

                finally:
                    cur2.close()
            t = threading.Thread(target=_process_sheets_queue, daemon=True)
            t.start()
            self._json({"status": "ok", "queued": len(created), "task_ids": created})

        finally:
            cur.close()
    def _serve_docx_bulk(self, article_ids):
        """Генерирует ZIP с несколькими DOCX файлами."""
        import io
        import json
        import zipfile
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        conn = get_connection()
        cur = conn.cursor()
        try:
            ph = "%s" if _is_postgres() else "?"

            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
                for aid in article_ids:
                    cur.execute(f"SELECT * FROM articles WHERE id = {ph}", (aid,))
                    row = cur.fetchone()
                    if not row:
                        continue
                    if _is_postgres():
                        columns = [desc[0] for desc in cur.description]
                        article = dict(zip(columns, row))
                    else:
                        article = dict(row)

                    doc = Document()
                    style = doc.styles['Normal']
                    style.font.name = 'Calibri'
                    style.font.size = Pt(11)

                    doc.add_heading(article.get("title", ""), level=1)

                    meta_p = doc.add_paragraph()
                    run = meta_p.add_run(f"Стиль: {article.get('style', '')} | Язык: {article.get('language', '')}")
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(128, 128, 128)
                    if article.get("source_url"):
                        run2 = meta_p.add_run(f"\nИсточник: {article['source_url']}")
                        run2.font.size = Pt(9)
                        run2.font.color.rgb = RGBColor(128, 128, 128)

                    if article.get("seo_title") or article.get("seo_description"):
                        doc.add_heading("SEO", level=2)
                        if article.get("seo_title"):
                            p = doc.add_paragraph()
                            p.add_run("Title: ").bold = True
                            p.add_run(article["seo_title"])
                        if article.get("seo_description"):
                            p = doc.add_paragraph()
                            p.add_run("Description: ").bold = True
                            p.add_run(article["seo_description"])

                    tags = []
                    try:
                        tags = json.loads(article.get("tags", "[]"))
                    except Exception:
                        pass
                    if tags:
                        p = doc.add_paragraph()
                        p.add_run("Теги: ").bold = True
                        p.add_run(", ".join(tags))

                    doc.add_paragraph("")
                    doc.add_heading("Текст статьи", level=2)
                    text = article.get("text", "")
                    for paragraph in text.split("\n"):
                        paragraph = paragraph.strip()
                        if paragraph:
                            if paragraph.startswith("## "):
                                doc.add_heading(paragraph[3:], level=3)
                            elif paragraph.startswith("# "):
                                doc.add_heading(paragraph[2:], level=2)
                            else:
                                doc.add_paragraph(paragraph)

                    doc_buffer = io.BytesIO()
                    doc.save(doc_buffer)
                    safe_title = "".join(c for c in article.get("title", "article")[:40] if c.isalnum() or c in " _-").strip() or "article"
                    zf.writestr(f"{safe_title}.docx", doc_buffer.getvalue())

            data = zip_buffer.getvalue()
            self.send_response(200)
            self.send_header("Content-Type", "application/zip")
            self.send_header("Content-Disposition", 'attachment; filename="articles.zip"')
            self.send_header("Content-Length", str(len(data)))
            self.end_headers()
            self.wfile.write(data)

        finally:
            cur.close()
    # --- Articles ---
    def _get_articles(self):
        conn = get_connection()
        cur = conn.cursor()
        try:
            cur.execute("SELECT * FROM articles ORDER BY updated_at DESC")
            if _is_postgres():
                columns = [desc[0] for desc in cur.description]
                return [dict(zip(columns, row)) for row in cur.fetchall()]
            return [dict(row) for row in cur.fetchall()]

        finally:
            cur.close()
    def _save_article(self, body):
        import uuid
        from datetime import datetime, timezone
        aid = str(uuid.uuid4())[:12]
        now = datetime.now(timezone.utc).isoformat()
        conn = get_connection()
        cur = conn.cursor()
        try:
            ph = "%s" if _is_postgres() else "?"
            import json
            tags = json.dumps(body.get("tags", []), ensure_ascii=False)
            cur.execute(f"""INSERT INTO articles (id, news_id, title, text, seo_title, seo_description, tags,
                style, language, original_title, original_text, source_url, status, created_at, updated_at)
                VALUES ({','.join([ph]*15)})""",
                (aid, body.get("news_id", ""), body.get("title", ""), body.get("text", ""),
                 body.get("seo_title", ""), body.get("seo_description", ""), tags,
                 body.get("style", ""), body.get("language", "русский"),
                 body.get("original_title", ""), body.get("original_text", ""),
                 body.get("source_url", ""), "draft", now, now))
            if not _is_postgres():
                conn.commit()
            self._json({"status": "ok", "id": aid})

        finally:
            cur.close()
    def _update_article(self, body):
        from datetime import datetime, timezone
        aid = body.get("id")
        if not aid:
            self._json({"status": "error", "message": "id required"})
            return

        # Save version snapshot before update (Phase 2)
        try:
            self._save_article_version(aid, body, change_type="manual",
                                        changed_by=self._get_session_user() or "admin")
        except Exception:
            pass

        now = datetime.now(timezone.utc).isoformat()
        conn = get_connection()
        cur = conn.cursor()
        try:
            ph = "%s" if _is_postgres() else "?"
            import json
            tags = json.dumps(body.get("tags", []), ensure_ascii=False)
            cur.execute(f"""UPDATE articles SET title={ph}, text={ph}, seo_title={ph},
                seo_description={ph}, tags={ph}, status={ph}, updated_at={ph} WHERE id={ph}""",
                (body.get("title", ""), body.get("text", ""), body.get("seo_title", ""),
                 body.get("seo_description", ""), tags, body.get("status", "draft"), now, aid))
            if not _is_postgres():
                conn.commit()
            self._json({"status": "ok"})

        finally:
            cur.close()
    def _delete_article(self, body):
        aid = body.get("id")
        if not aid:
            self._json({"status": "error", "message": "id required"})
            return
        conn = get_connection()
        cur = conn.cursor()
        try:
            ph = "%s" if _is_postgres() else "?"
            cur.execute(f"DELETE FROM articles WHERE id = {ph}", (aid,))
            if not _is_postgres():
                conn.commit()
            self._json({"status": "ok"})

        finally:
            cur.close()
    def _article_detail(self, body):
        aid = body.get("id")
        if not aid:
            self._json({"status": "error", "message": "id required"})
            return
        conn = get_connection()
        cur = conn.cursor()
        try:
            ph = "%s" if _is_postgres() else "?"
            cur.execute(f"SELECT * FROM articles WHERE id = {ph}", (aid,))
            row = cur.fetchone()
            if not row:
                self._json({"status": "error", "message": "Not found"})
                return
            if _is_postgres():
                columns = [desc[0] for desc in cur.description]
                article = dict(zip(columns, row))
            else:
                article = dict(row)
            self._json({"status": "ok", "article": article})

        finally:
            cur.close()
    def _schedule_article(self, body):
        """Запланировать публикацию статьи на указанное время."""
        from datetime import datetime, timezone
        aid = body.get("article_id") or body.get("id")
        scheduled_at = body.get("scheduled_at")
        if not aid:
            self._json({"status": "error", "message": "article_id required"})
            return
        if not scheduled_at:
            self._json({"status": "error", "message": "scheduled_at required"})
            return
        now = datetime.now(timezone.utc).isoformat()
        conn = get_connection()
        cur = conn.cursor()
        try:
            ph = "%s" if _is_postgres() else "?"
            cur.execute(f"UPDATE articles SET scheduled_at={ph}, status='scheduled', updated_at={ph} WHERE id={ph}",
                        (scheduled_at, now, aid))
            if not _is_postgres():
                conn.commit()
            self._json({"status": "ok", "scheduled_at": scheduled_at})
        finally:
            cur.close()
    def _rewrite_article(self, body):
        """Переписать существующую статью в другом стиле."""
        aid = body.get("id")
        style = body.get("style", "news")
        language = body.get("language", "русский")
        conn = get_connection()
        cur = conn.cursor()
        try:
            ph = "%s" if _is_postgres() else "?"
            cur.execute(f"SELECT title, text, original_title, original_text FROM articles WHERE id = {ph}", (aid,))
            row = cur.fetchone()
            if not row:
                self._json({"status": "error", "message": "Article not found"})
                return
            if _is_postgres():
                columns = [desc[0] for desc in cur.description]
                article = dict(zip(columns, row))
            else:
                article = dict(row)
            # Use original text for rewriting to avoid degradation
            src_title = article.get("original_title") or article.get("title", "")
            src_text = article.get("original_text") or article.get("text", "")
            from apis.llm import rewrite_news
            result = rewrite_news(src_title, src_text, style, language)
            if result:
                self._json({"status": "ok", "result": result})
            else:
                self._json({"status": "error", "message": "LLM returned no result"})

        finally:
            cur.close()
    def _improve_article(self, body):
        """Улучшить текст статьи через LLM (грамматика, стиль, SEO)."""
        aid = body.get("id")
        action = body.get("action", "improve")  # improve, expand, shorten, fix_grammar, add_seo
        conn = get_connection()
        cur = conn.cursor()
        try:
            ph = "%s" if _is_postgres() else "?"
            cur.execute(f"SELECT title, text FROM articles WHERE id = {ph}", (aid,))
            row = cur.fetchone()
            if not row:
                self._json({"status": "error", "message": "Article not found"})
                return
            if _is_postgres():
                columns = [desc[0] for desc in cur.description]
                article = dict(zip(columns, row))
            else:
                article = dict(row)

            actions_map = {
                "improve": "Улучши текст: исправь стилистические ошибки, сделай более профессиональным, сохрани факты.",
                "expand": "Расширь текст: добавь подробностей, контекста, аналитики. Увеличь объём в 1.5-2 раза, не добавляя вымышленных фактов.",
                "shorten": "Сократи текст в 2 раза, оставив только ключевые факты. Убери воду и повторы.",
                "fix_grammar": "Исправь все грамматические, пунктуационные и стилистические ошибки. Не меняй смысл и структуру.",
                "add_seo": "Добавь SEO-оптимизацию: включи ключевые слова естественно, добавь подзаголовки (## H2), улучши мета-описание.",
                "make_engaging": "Сделай текст более вовлекающим: добавь интригу, живые примеры, вопросы к читателю. Сохрани факты.",
            }
            instruction = actions_map.get(action, actions_map["improve"])

            from apis.llm import _call_llm
            prompt = f"""Ты — профессиональный редактор игровых новостей.
    
    Задача: {instruction}
    
    Заголовок: {article['title']}
    Текст: {article['text'][:4000]}
    
    Верни строго JSON (без markdown):
    {{
      "title": "обновлённый заголовок",
      "text": "обновлённый текст",
      "seo_title": "SEO title до 60 символов",
      "seo_description": "meta description до 155 символов",
      "changes_summary": "что было изменено (1-2 предложения)"
    }}"""
            result = _call_llm(prompt)
            if result:
                self._json({"status": "ok", "result": result})
            else:
                self._json({"status": "error", "message": "LLM returned no result"})

        finally:
            cur.close()
    def _serve_docx(self, article_id):
        """Генерация и отдача DOCX файла."""
        conn = get_connection()
        cur = conn.cursor()
        try:
            ph = "%s" if _is_postgres() else "?"
            cur.execute(f"SELECT * FROM articles WHERE id = {ph}", (article_id,))
            row = cur.fetchone()
            if not row:
                self.send_response(404)
                self.end_headers()
                return
            if _is_postgres():
                columns = [desc[0] for desc in cur.description]
                article = dict(zip(columns, row))
            else:
                article = dict(row)

            import io
            import json
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()
            style = doc.styles['Normal']
            style.font.name = 'Calibri'
            style.font.size = Pt(11)

            # Title
            title_p = doc.add_heading(article.get("title", ""), level=1)
            title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Meta info block
            meta_p = doc.add_paragraph()
            meta_p.paragraph_format.space_after = Pt(6)
            run = meta_p.add_run(f"Стиль: {article.get('style', '')} | Язык: {article.get('language', '')}")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(128, 128, 128)
            if article.get("source_url"):
                run2 = meta_p.add_run(f"\nИсточник: {article['source_url']}")
                run2.font.size = Pt(9)
                run2.font.color.rgb = RGBColor(128, 128, 128)

            # SEO block
            if article.get("seo_title") or article.get("seo_description"):
                doc.add_heading("SEO", level=2)
                if article.get("seo_title"):
                    p = doc.add_paragraph()
                    p.add_run("Title: ").bold = True
                    p.add_run(article["seo_title"])
                if article.get("seo_description"):
                    p = doc.add_paragraph()
                    p.add_run("Description: ").bold = True
                    p.add_run(article["seo_description"])

            # Tags
            tags = []
            try:
                tags = json.loads(article.get("tags", "[]"))
            except Exception:
                pass
            if tags:
                p = doc.add_paragraph()
                p.add_run("Теги: ").bold = True
                p.add_run(", ".join(tags))

            doc.add_paragraph("")  # spacer

            # Article text
            doc.add_heading("Текст статьи", level=2)
            text = article.get("text", "")
            for paragraph in text.split("\n"):
                paragraph = paragraph.strip()
                if paragraph:
                    if paragraph.startswith("## "):
                        doc.add_heading(paragraph[3:], level=3)
                    elif paragraph.startswith("# "):
                        doc.add_heading(paragraph[2:], level=2)
                    else:
                        doc.add_paragraph(paragraph)

            # Original text if exists
            if article.get("original_text"):
                doc.add_page_break()
                doc.add_heading("Оригинал", level=2)
                orig_p = doc.add_paragraph()
                if article.get("original_title"):
                    run = orig_p.add_run(article["original_title"] + "\n\n")
                    run.bold = True
                for line in article["original_text"][:3000].split("\n"):
                    line = line.strip()
                    if line:
                        p = doc.add_paragraph(line)
                        for run in p.runs:
                            run.font.color.rgb = RGBColor(128, 128, 128)
                            run.font.size = Pt(10)

            buffer = io.BytesIO()
            doc.save(buffer)
            data = buffer.getvalue()

            safe_title = "".join(c for c in article.get("title", "article")[:40] if c.isalnum() or c in " _-").strip() or "article"
            filename = f"{safe_title}.docx"

            self.send_response(200)
            self.send_header("Content-Type", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            self.send_header("Content-Disposition", f'attachment; filename="{filename}"')
            self.send_header("Content-Length", str(len(data)))
            self.end_headers()
            self.wfile.write(data)

        finally:
            cur.close()
    # --- Logs, Cache, Rate, Translate, AI ---

    def _get_viral(self):
        """Анализ виральности: прогоняет все новости через viral_score + sentiment + momentum."""
        from checks.viral_score import viral_score, VIRAL_TRIGGERS, get_calendar_boost
        from checks.sentiment import analyze_sentiment
        from checks.tags import auto_tag
        from apis.cache import cache_get, cache_set, cache_key

        qs = parse_qs(urlparse(self.path).query)
        limit = int(qs.get("limit", [200])[0])
        level_filter = qs.get("level", [None])[0]
        category_filter = qs.get("category", [None])[0]
        sentiment_filter = qs.get("sentiment", [None])[0]
        source_filter = qs.get("source", [None])[0]
        date_from = qs.get("date_from", [None])[0]
        date_to = qs.get("date_to", [None])[0]
        trigger_filter = qs.get("trigger", [None])[0]
        min_score = int(qs.get("min_score", [0])[0])

        conn = get_connection()
        cur = conn.cursor()
        try:
            ph = "%s" if _is_postgres() else "?"

            conditions = []
            params = []
            if source_filter:
                conditions.append(f"n.source = {ph}")
                params.append(source_filter)
            if date_from:
                conditions.append(f"n.parsed_at >= {ph}")
                params.append(date_from)
            if date_to:
                conditions.append(f"n.parsed_at <= {ph}")
                params.append(date_to + "T23:59:59")
            where = "WHERE " + " AND ".join(conditions) if conditions else ""

            cur.execute(f"""
                SELECT n.id, n.source, n.title, n.url, n.description, n.plain_text,
                       n.published_at, n.parsed_at, n.status
                FROM news n {where}
                ORDER BY n.parsed_at DESC LIMIT {ph}
            """, params + [limit])

            if _is_postgres():
                columns = [desc[0] for desc in cur.description]
                rows = [dict(zip(columns, row)) for row in cur.fetchall()]
            else:
                rows = [dict(row) for row in cur.fetchall()]

            items = []
            stats = {"total": 0, "high": 0, "medium": 0, "low": 0, "none": 0}
            trigger_counts = {}
            category_counts = {}
            sentiment_counts = {"positive": 0, "negative": 0, "neutral": 0}
            source_scores = {}

            # Category mapping from trigger_id prefix
            CATEGORY_MAP = {
                "scandal": "Скандалы", "leak": "Утечки", "shadow": "Shadow Drops",
                "bad": "Плохие релизы", "ai": "AI", "major_event": "Ивенты",
                "event": "Ивенты", "money": "Деньги", "culture": "Культура",
                "person": "Персоны", "speed": "Скорость",
                "sequel": "Базовые", "free_content": "Базовые", "delay": "Базовые",
                "canceled": "Базовые", "award": "Базовые", "next_gen": "Базовые",
                "big_update": "Базовые", "release_date": "Базовые",
                "trailer": "Базовые", "record": "Базовые", "digest": "Базовые",
            }

            for row in rows:
                ck = cache_key("viral_tab", row["id"])
                cached = cache_get(ck)
                if cached:
                    vr = cached["viral"]
                    sent = cached["sentiment"]
                    tags = cached["tags"]
                else:
                    vr = viral_score(row)
                    sent = analyze_sentiment(row)
                    tags = auto_tag(row)
                    cache_set(ck, {"viral": vr, "sentiment": sent, "tags": tags}, ttl=3600)

                # Determine categories of triggers
                trigger_categories = set()
                for t in vr["triggers"]:
                    tid = t["id"]
                    prefix = tid.split("_")[0]
                    cat = CATEGORY_MAP.get(tid, CATEGORY_MAP.get(prefix, "Прочее"))
                    trigger_categories.add(cat)

                # Apply filters
                if level_filter and vr["level"] != level_filter:
                    continue
                if min_score and vr["score"] < min_score:
                    continue
                if sentiment_filter and sent["label"] != sentiment_filter:
                    continue
                if trigger_filter:
                    if not any(t["id"] == trigger_filter for t in vr["triggers"]):
                        continue
                if category_filter:
                    if category_filter not in trigger_categories:
                        continue

                item = {
                    "id": row["id"],
                    "source": row["source"],
                    "title": row["title"],
                    "url": row["url"],
                    "published_at": row["published_at"],
                    "parsed_at": row["parsed_at"],
                    "status": row["status"],
                    "viral_score": vr["score"],
                    "viral_level": vr["level"],
                    "triggers": vr["triggers"],
                    "sentiment": sent["label"],
                    "sentiment_score": sent["score"],
                    "tags": [{"id": t["id"], "label": t["label"]} for t in tags[:3]],
                }
                items.append(item)

                # Aggregate stats
                stats["total"] += 1
                stats[vr["level"]] = stats.get(vr["level"], 0) + 1
                sentiment_counts[sent["label"]] = sentiment_counts.get(sent["label"], 0) + 1
                for t in vr["triggers"]:
                    trigger_counts[t["label"]] = trigger_counts.get(t["label"], 0) + 1
                    prefix = t["id"].split("_")[0]
                    cat = CATEGORY_MAP.get(t["id"], CATEGORY_MAP.get(prefix, "Прочее"))
                    category_counts[cat] = category_counts.get(cat, 0) + 1
                src = row["source"]
                if src not in source_scores:
                    source_scores[src] = {"total": 0, "sum": 0}
                source_scores[src]["total"] += 1
                source_scores[src]["sum"] += vr["score"]

            # Sort by viral_score desc
            items.sort(key=lambda x: x["viral_score"], reverse=True)

            # Top triggers sorted
            top_triggers = sorted(trigger_counts.items(), key=lambda x: x[1], reverse=True)[:20]

            # Top categories sorted
            top_categories = sorted(category_counts.items(), key=lambda x: x[1], reverse=True)

            # Source avg scores
            source_avg = []
            for src, data in source_scores.items():
                source_avg.append({"source": src, "avg": round(data["sum"] / data["total"], 1), "count": data["total"]})
            source_avg.sort(key=lambda x: x["avg"], reverse=True)

            # Calendar event
            cal_boost, cal_event = get_calendar_boost()

            # Available triggers for filter
            all_triggers = [{"id": k, "label": v["label"], "category": CATEGORY_MAP.get(k, CATEGORY_MAP.get(k.split("_")[0], "Прочее"))} for k, v in VIRAL_TRIGGERS.items()]

            return {
                "items": items,
                "stats": stats,
                "sentiment": sentiment_counts,
                "top_triggers": top_triggers,
                "top_categories": top_categories,
                "source_avg": source_avg[:15],
                "calendar": {"boost": cal_boost, "event": cal_event},
                "all_triggers": all_triggers,
            }

        finally:
            cur.close()
    def _get_logs(self):
        qs = parse_qs(urlparse(self.path).query)
        limit = int(qs.get("limit", [100])[0])
        level = qs.get("level", [""])[0]
        from apis.cache import get_logs
        return {"logs": get_logs(limit=limit, level=level)}

    def _get_rate_stats(self):
        from apis.cache import get_rate_stats
        return get_rate_stats()

    def _get_cache_stats(self):
        from apis.cache import get_cache_stats
        return get_cache_stats()

    def _clear_cache(self, body):
        from apis.cache import clear_cache
        clear_cache()
        self._json({"status": "ok"})

    def _translate_title(self, body):
        news_id = body.get("news_id", "")
        if not news_id:
            self._json({"status": "error", "message": "news_id required"})
            return
        conn = get_connection()
        cur = conn.cursor()
        try:
            ph = "%s" if _is_postgres() else "?"
            cur.execute(f"SELECT title FROM news WHERE id = {ph}", (news_id,))
            row = cur.fetchone()
            if not row:
                self._json({"status": "error", "message": "not found"})
                return
            title = row[0] if _is_postgres() else row["title"]
            try:
                from apis.llm import translate_title
                result = translate_title(title)
                if result:
                    # Save translated title as h1 if not Russian
                    if not result.get("is_russian") and result.get("translated"):
                        cur.execute(f"UPDATE news SET h1 = {ph} WHERE id = {ph}", (result["translated"], news_id))
                        if not _is_postgres():
                            conn.commit()
                    self._json({"status": "ok", **result})
                else:
                    self._json({"status": "error", "message": "LLM not responding. Check API keys and rate limits."})
            except Exception as e:
                logger.error("Translate error: %s", e)
                self._json({"status": "error", "message": str(e)})

        finally:
            cur.close()
    def _ai_recommend(self, body):
        news_id = body.get("news_id", "")
        if not news_id:
            self._json({"status": "error", "message": "news_id required"})
            return
        conn = get_connection()
        cur = conn.cursor()
        try:
            ph = "%s" if _is_postgres() else "?"
            cur.execute(f"SELECT * FROM news WHERE id = {ph}", (news_id,))
            row = cur.fetchone()
            if not row:
                self._json({"status": "error", "message": "not found"})
                return
            if _is_postgres():
                columns = [desc[0] for desc in cur.description]
                news = dict(zip(columns, row))
            else:
                news = dict(row)

            # Run checks for context
            from checks.quality import check_quality
            from checks.relevance import check_relevance
            from checks.freshness import check_freshness
            from checks.viral_score import viral_score
            checks = {
                "quality": check_quality(news),
                "relevance": check_relevance(news),
                "freshness": check_freshness(news),
                "viral": viral_score(news),
            }
            from apis.llm import ai_recommendation
            result = ai_recommendation(
                title=news.get("title", ""),
                text=news.get("plain_text", "") or news.get("description", ""),
                source=news.get("source", ""),
                checks=checks,
            )
            if result:
                self._json({"status": "ok", "recommendation": result, "checks": {k: v.get("score", 0) for k, v in checks.items()}})
            else:
                self._json({"status": "error", "message": "AI recommendation failed"})

        finally:
            cur.close()
    # --- Phase 3: Analytics Funnel & Source Intelligence ---

    def _get_funnel_analytics(self):
        from api.analytics import get_funnel_analytics
        return get_funnel_analytics()

    def _get_cost_by_source(self):
        from api.analytics import get_cost_by_source
        return get_cost_by_source()

    def _get_prompt_insights(self):
        from api.analytics import get_prompt_insights
        return get_prompt_insights()

    # --- Phase 2: Content Versioning & Multi-Output ---

    def _get_article_versions(self, body):
        """Get version history for an article."""
        article_id = body.get("article_id", "")
        if not article_id:
            self._json({"error": "article_id required"}, 400)
            return
        conn = get_connection()
        cur = conn.cursor()
        ph = "%s" if _is_postgres() else "?"
        try:
            cur.execute(f"""
                SELECT * FROM article_versions
                WHERE article_id = {ph}
                ORDER BY version DESC
            """, (article_id,))
            if _is_postgres():
                columns = [desc[0] for desc in cur.description]
                rows = [dict(zip(columns, r)) for r in cur.fetchall()]
            else:
                rows = [dict(r) for r in cur.fetchall()]
            self._json({"versions": rows})
        except Exception as e:
            self._json({"versions": [], "error": str(e)})
        finally:
            cur.close()

    def _save_article_version(self, article_id, article_data, change_type="manual", changed_by="system"):
        """Save a version snapshot before modification (internal helper)."""
        try:
            from core.feature_flags import is_enabled
            if not is_enabled("content_versions_v1"):
                return
        except Exception:
            return

        import uuid
        from datetime import datetime, timezone
        conn = get_connection()
        cur = conn.cursor()
        ph = "%s" if _is_postgres() else "?"
        try:
            # Get current max version
            cur.execute(f"SELECT COALESCE(MAX(version), 0) FROM article_versions WHERE article_id = {ph}", (article_id,))
            max_ver = cur.fetchone()[0]
            now = datetime.now(timezone.utc).isoformat()
            ver_id = uuid.uuid4().hex[:12]

            import json
            cur.execute(f"""
                INSERT INTO article_versions (id, article_id, version, title, text, seo_title, seo_description, tags, change_type, changed_by, created_at)
                VALUES ({','.join([ph]*11)})
            """, (ver_id, article_id, max_ver + 1,
                  article_data.get("title", "")[:500],
                  article_data.get("text", "")[:5000],
                  article_data.get("seo_title", "")[:500],
                  article_data.get("seo_description", "")[:1000],
                  json.dumps(article_data.get("tags", []), ensure_ascii=False),
                  change_type, changed_by, now))
            if not _is_postgres():
                conn.commit()
        except Exception as e:
            logger.debug("Failed to save article version: %s", e)
        finally:
            cur.close()

    def _generate_multi_output(self, body):
        """Generate multiple output formats from one article."""
        article_id = body.get("article_id", "")
        formats = body.get("formats", ["social", "short"])
        if not article_id:
            self._json({"error": "article_id required"}, 400)
            return

        conn = get_connection()
        cur = conn.cursor()
        ph = "%s" if _is_postgres() else "?"
        try:
            cur.execute(f"SELECT title, text FROM articles WHERE id = {ph}", (article_id,))
            row = cur.fetchone()
            if not row:
                self._json({"error": "article not found"}, 404)
                return
            if _is_postgres():
                title, text = row[0], row[1]
            else:
                title, text = row["title"], row["text"]

            from apis.llm import rewrite_news
            results = {}
            for fmt in formats[:3]:  # Max 3 formats at once
                result = rewrite_news(title=title, text=text, style=fmt, language="русский")
                if result:
                    results[fmt] = result

            self._json({"article_id": article_id, "outputs": results})
        except Exception as e:
            self._json({"error": str(e)}, 500)
        finally:
            cur.close()

    def _regenerate_field(self, body):
        """Regenerate a single field (title, seo_title, seo_description, tags) via LLM."""
        article_id = body.get("article_id", "")
        field = body.get("field", "")
        if not article_id or field not in ("title", "seo_title", "seo_description", "tags"):
            self._json({"error": "article_id and valid field required"}, 400)
            return

        conn = get_connection()
        cur = conn.cursor()
        ph = "%s" if _is_postgres() else "?"
        try:
            cur.execute(f"SELECT title, text FROM articles WHERE id = {ph}", (article_id,))
            row = cur.fetchone()
            if not row:
                self._json({"error": "article not found"}, 404)
                return
            if _is_postgres():
                title, text = row[0], row[1]
            else:
                title, text = row["title"], row["text"]

            from apis.llm import _call_llm
            prompts = {
                "title": f"Придумай новый заголовок для игровой новости. Текст:\n{text[:1500]}\n\nОтветь JSON: {{\"title\": \"новый заголовок\"}}",
                "seo_title": f"Придумай SEO-заголовок до 60 символов для:\n{title}\n\nОтветь JSON: {{\"seo_title\": \"SEO заголовок\"}}",
                "seo_description": f"Придумай meta description до 155 символов для:\n{title}\n{text[:500]}\n\nОтветь JSON: {{\"seo_description\": \"описание\"}}",
                "tags": f"Предложи 5 тегов для игровой новости:\n{title}\n\nОтветь JSON: {{\"tags\": [\"тег1\", \"тег2\", \"тег3\", \"тег4\", \"тег5\"]}}",
            }

            result = _call_llm(prompts[field])
            if result and field in result:
                self._json({"field": field, "value": result[field]})
            else:
                self._json({"error": "LLM returned no result"}, 500)
        except Exception as e:
            self._json({"error": str(e)}, 500)
        finally:
            cur.close()

    # --- Phase 0: Feature Flags & Observability API ---

    def _get_feature_flags(self):
        try:
            from core.feature_flags import get_all_flags
            return {"flags": get_all_flags()}
        except Exception as e:
            return {"flags": [], "error": str(e)}

    def _toggle_feature_flag(self, body):
        if not self._require_perm("flags"):
            return
        flag_id = body.get("flag_id", "")
        enabled = body.get("enabled", False)
        if not flag_id:
            self._json({"error": "flag_id required"}, 400)
            return
        try:
            from core.feature_flags import set_flag
            user = self._get_session_user() or "admin"
            set_flag(flag_id, bool(enabled), updated_by=user)
            self._json({"status": "ok", "flag_id": flag_id, "enabled": enabled})
        except Exception as e:
            self._json({"error": str(e)}, 500)

    def _get_cost_summary(self):
        from api.analytics import get_cost_summary
        return get_cost_summary()

    def _get_config_audit(self):
        try:
            from core.observability import get_config_audit
            return {"audit": get_config_audit(limit=50)}
        except Exception as e:
            return {"audit": [], "error": str(e)}

    def _get_decision_trace(self, body):
        news_id = body.get("news_id", "")
        if not news_id:
            self._json({"error": "news_id required"}, 400)
            return
        try:
            from core.observability import get_decision_trace
            trace = get_decision_trace(news_id)
            self._json({"news_id": news_id, "trace": trace})
        except Exception as e:
            self._json({"error": str(e)}, 500)

    def _get_ops_dashboard(self):
        from api.dashboard import get_ops_dashboard
        return get_ops_dashboard()

    # --- Phase 4: Storylines, Source Health Plus, Threshold Simulator ---

    def _get_storylines(self):
        from api.dashboard import get_storylines
        return get_storylines()

    def _get_source_health_plus(self):
        from api.dashboard import get_source_health_plus
        return get_source_health_plus()

    def _simulate_thresholds(self, body):
        from api.dashboard import simulate_thresholds
        self._json(simulate_thresholds(body))

    def _rescore_news(self, body):
        """Re-run scoring pipeline for specific news or all with score=0."""
        if not self._require_perm("pipeline"):
            return
        news_ids = body.get("news_ids", [])
        rescore_zero = body.get("rescore_zero", False)
        conn = get_connection()
        cur = conn.cursor()
        ph = "%s" if _is_postgres() else "?"
        try:
            if rescore_zero:
                # Find all with score=0, missing analysis, or rejected due to low quality
                cur.execute(f"""
                    SELECT n.* FROM news n
                    LEFT JOIN news_analysis a ON n.id = a.news_id
                    WHERE n.status IN ('in_review', 'new', 'rejected')
                    AND (a.total_score IS NULL OR a.total_score = 0 OR a.news_id IS NULL)
                    ORDER BY n.parsed_at DESC
                    LIMIT 500
                """)
            elif news_ids:
                placeholders = ",".join([ph] * len(news_ids))
                cur.execute(f"SELECT * FROM news WHERE id IN ({placeholders})", tuple(news_ids))
            else:
                self._json({"status": "error", "message": "news_ids or rescore_zero required"})
                return

            if _is_postgres():
                columns = [desc[0] for desc in cur.description]
                news_list = [dict(zip(columns, r)) for r in cur.fetchall()]
            else:
                news_list = [dict(r) for r in cur.fetchall()]
        finally:
            cur.close()

        if not news_list:
            self._json({"status": "ok", "rescored": 0, "message": "Нет новостей для пересчёта"})
            return

        from checks.pipeline import run_review_pipeline
        result = run_review_pipeline(news_list, update_status=False)
        scored = len(result.get("results", []))
        self._json({"status": "ok", "rescored": scored})

    # --- Dashboard HTML ---
    def _serve_dashboard(self):
        global _DASHBOARD_HTML_GZIP, _DASHBOARD_HTML_BYTES
        # Pre-compress dashboard HTML once (saves ~400KB per request)
        if _DASHBOARD_HTML_BYTES is None:
            _DASHBOARD_HTML_BYTES = DASHBOARD_HTML.encode()
            _DASHBOARD_HTML_GZIP = gzip.compress(_DASHBOARD_HTML_BYTES, compresslevel=9)
            logger.info("Dashboard HTML compressed: %dKB -> %dKB (%.0f%% saved)",
                        len(_DASHBOARD_HTML_BYTES) // 1024, len(_DASHBOARD_HTML_GZIP) // 1024,
                        (1 - len(_DASHBOARD_HTML_GZIP) / len(_DASHBOARD_HTML_BYTES)) * 100)

        self.send_response(200)
        self.send_header("Content-Type", "text/html; charset=utf-8")
        # Allow browser to cache for 5 min, revalidate after
        self.send_header("Cache-Control", "public, max-age=300, must-revalidate")
        if self._accepts_gzip():
            body = _DASHBOARD_HTML_GZIP
            self.send_header("Content-Encoding", "gzip")
        else:
            body = _DASHBOARD_HTML_BYTES
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def log_message(self, format, *args):
        pass


# Dashboard HTML loaded from static file
def _load_dashboard_html():
    import os
    html_path = os.path.join(os.path.dirname(__file__), "static", "dashboard.html")
    with open(html_path, "r", encoding="utf-8") as f:
        return f.read()

DASHBOARD_HTML = _load_dashboard_html()



def _load_active_prompts():
    """Загружает активные версии промптов из БД при старте."""
    try:
        conn = get_connection()
        cur = conn.cursor()
        try:
            cur.execute("SELECT prompt_name, content FROM prompt_versions WHERE is_active = 1")
            import apis.llm as llm
            prompt_map = {
                "trend_forecast": "PROMPT_TREND_FORECAST",
                "merge_analysis": "PROMPT_MERGE_ANALYSIS",
                "keyso_queries": "PROMPT_KEYSO_QUERIES",
                "rewrite": "PROMPT_REWRITE",
            }
            if _is_postgres():
                for row in cur.fetchall():
                    attr = prompt_map.get(row[0])
                    if attr and hasattr(llm, attr):
                        setattr(llm, attr, row[1])
                        logger.info("Loaded active prompt: %s", row[0])
            else:
                for row in cur.fetchall():
                    r = dict(row)
                    attr = prompt_map.get(r["prompt_name"])
                    if attr and hasattr(llm, attr):
                        setattr(llm, attr, r["content"])
                        logger.info("Loaded active prompt: %s", r["prompt_name"])
        finally:
            cur.close()
    except Exception as e:
        logger.debug("Could not load active prompts: %s", e)


def start_web():
    _load_active_prompts()
    server = HTTPServer(("0.0.0.0", PORT), AdminHandler)
    thread = threading.Thread(target=server.serve_forever, daemon=True)
    thread.start()
    logger.info("Admin panel running on port %d", PORT)
    return server
