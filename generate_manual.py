"""Generate IgroNews User Manual v3.0 as DOCX — comprehensive edition."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import os

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Styles ──
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
    if level == 1:
        h.font.size = Pt(22)
        h.paragraph_format.space_before = Pt(24)
        h.paragraph_format.space_after = Pt(12)
    elif level == 2:
        h.font.size = Pt(16)
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(8)
    else:
        h.font.size = Pt(13)
        h.font.color.rgb = RGBColor(0x42, 0x85, 0xF4)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)


def add_styled_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = 'Calibri'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = cell._element.get_or_add_tcPr()
        shd = shading.makeelement(qn('w:shd'), {qn('w:fill'): '1A73E8', qn('w:val'): 'clear'})
        shading.append(shd)
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(text))
            run.font.size = Pt(9.5)
            run.font.name = 'Calibri'
            if r_idx % 2 == 1:
                shading = cell._element.get_or_add_tcPr()
                shd = shading.makeelement(qn('w:shd'), {qn('w:fill'): 'F0F4FF', qn('w:val'): 'clear'})
                shading.append(shd)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph('')
    return table


def add_bullets(items, bold_prefix=False):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        if bold_prefix and ':' in item:
            parts = item.split(':', 1)
            run = p.add_run(parts[0] + ':')
            run.bold = True
            run.font.size = Pt(10.5)
            p.add_run(parts[1]).font.size = Pt(10.5)
        else:
            run = p.add_run(item)
            run.font.size = Pt(10.5)


def add_info_box(text, color='blue'):
    p = doc.add_paragraph()
    colors = {
        'blue': ('E8F0FE', '1A73E8'),
        'green': ('E6F4EA', '137333'),
        'orange': ('FEF7E0', 'E37400'),
        'red': ('FCE8E6', 'C5221F'),
        'purple': ('F3E8FD', '7B1FA2'),
    }
    bg, fg = colors.get(color, colors['blue'])
    run = p.add_run(f'  {text}')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(int(fg[:2], 16), int(fg[2:4], 16), int(fg[4:], 16))
    run.font.name = 'Calibri'
    pPr = p._element.get_or_add_pPr()
    shd = pPr.makeelement(qn('w:shd'), {qn('w:fill'): bg, qn('w:val'): 'clear'})
    pPr.append(shd)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)


def add_numbered(items):
    for item in items:
        p = doc.add_paragraph(style='List Number')
        if ':' in item:
            parts = item.split(':', 1)
            run = p.add_run(parts[0] + ':')
            run.bold = True
            run.font.size = Pt(10.5)
            p.add_run(parts[1]).font.size = Pt(10.5)
        else:
            p.add_run(item).font.size = Pt(10.5)


def add_formula(label, formula, note=''):
    p = doc.add_paragraph()
    run = p.add_run(f'{label}: ')
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
    run2 = p.add_run(formula)
    run2.font.size = Pt(11)
    run2.font.name = 'Consolas'
    if note:
        p2 = doc.add_paragraph()
        r = p2.add_run(note)
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        r.italic = True


# ══════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════

for _ in range(6):
    doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('IgroNews')
run.font.size = Pt(42)
run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
run.bold = True
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Агрегатор игровых новостей с NLP-аналитикой')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Полное руководство пользователя')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x42, 0x85, 0xF4)

for _ in range(8):
    doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Версия 3.0  •  Март 2026')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ══════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════

doc.add_heading('Содержание', level=1)

toc_items = [
    '1. Обзор системы',
    '2. Быстрый старт',
    '3. Рабочий процесс (пайплайн)',
    '4. Полный автомат — автоматический конвейер',
    '5. Вкладка «Редакция» — основная работа',
    '6. Вкладка «Обогащённые» — данные API',
    '7. Вкладка «Финал» — лучшие новости',
    '8. Вкладка «Контент» — статьи и рерайт',
    '9. Вкладка «Виральность» — тренды и триггеры',
    '10. Вкладка «Аналитика» — статистика и графики',
    '11. Вкладка «Очередь» — управление задачами',
    '12. Вкладка «Здоровье» — мониторинг источников',
    '13. Вкладка «Настройки» — конфигурация',
    '14. Источники новостей (15 сайтов)',
    '15. Система проверок — 12 критериев',
    '16. Формулы скоринга — все расчёты',
    '17. Финальный скор — композитная формула',
    '18. API-интеграции — подробно',
    '19. SEO-анализ статей',
    '20. Планировщик публикаций',
    '21. Уведомления браузера',
    '22. Умный polling и оптимизации',
    '23. Прокси-ротация для парсеров',
    '24. Система сущностей (NER)',
    '25. Дедупликация — как работает',
    '26. Feedback loop — обучаемость',
    '27. Google Sheets — экспорт',
    '28. Часто задаваемые вопросы',
]
for item in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(item)
    run.font.size = Pt(11.5)
    run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
    p.paragraph_format.space_after = Pt(3)

doc.add_page_break()

# ══════════════════════════════════════════════
# 1. OVERVIEW
# ══════════════════════════════════════════════

doc.add_heading('1. Обзор системы', level=1)

doc.add_paragraph(
    'IgroNews — автоматизированный агрегатор игровых новостей с NLP-аналитикой. '
    'Система парсит 15 источников (IGN, PCGamer, DTF, StopGame и др.), '
    'автоматически оценивает каждую новость по 12 критериям, '
    'обогащает данными из внешних API (Keys.so, Google Trends, LLM) '
    'и готовит контент к публикации через рерайт и экспорт в Google Sheets.'
)

doc.add_heading('Ключевые возможности', level=3)
add_bullets([
    'Автоматический парсинг: 15 источников, RSS + HTML + JSON, каждые 5–30 минут',
    'Прокси-ротация: round-robin, circuit breaker, ротация User-Agent',
    'Умная оценка: 12 проверок (качество, виральность, свежесть, релевантность, моментум, заголовок и др.)',
    'Авто-дедупликация: MD5 по URL + TF-IDF cosine similarity + entity overlap',
    'Обогащение API: Keys.so (поисковые частоты), Google Trends (4 региона), LLM (прогноз трендов)',
    'Двойной скоринг: внутренний total_score (бесплатный) + финальный скор (с API-данными)',
    'Полный автомат: скор → >70 на LLM → финальный скор → >60 на рерайт → Sheets',
    'Рерайт через LLM: 6 стилей (новость, SEO, обзор, кликбейт, короткий, соцсети)',
    'SEO-анализ: 9 критериев оценки статьи',
    'Планировщик публикаций: отложенная публикация по расписанию',
    'Браузерные уведомления: оповещения о топ-новостях',
    'Экспорт в Google Sheets: автоматический и ручной',
    'Обучаемость: feedback loop, система учится на решениях редактора',
    'Веб-панель: 9 вкладок для полного управления',
], bold_prefix=True)

doc.add_heading('Структура вкладок', level=3)

add_styled_table(
    ['Вкладка', 'Назначение'],
    [
        ['Редакция', 'Основная работа: ревью, одобрение, отклонение, полный автомат'],
        ['Обогащённые', 'Новости с API-данными (Keys.so, Trends, LLM)'],
        ['Финал', 'Лучшие новости (publish_now) с финальным композитным скором'],
        ['Контент', 'Статьи + рерайт (подвкладки Рерайт/Статьи)'],
        ['Виральность', 'Анализ виральности, триггеры, отправка в контент'],
        ['Аналитика', 'Графики, статистика, дайджесты, версии промптов'],
        ['Очередь', 'Управление задачами: полный автомат, рерайт, экспорт'],
        ['Здоровье', 'Мониторинг источников, веса, статистика'],
        ['⚙ Настройки', 'Общие, источники, промпты, инструменты, логи, пользователи'],
    ],
    col_widths=[3.5, 11.5]
)

doc.add_heading('Технологии', level=3)
add_bullets([
    'Backend: Python 3.11, HTTP-сервер (без фреймворка), APScheduler',
    'NLP: scikit-learn (TF-IDF с дисковым кэшем), NLTK, кастомный NER (80+ игр, 30+ студий)',
    'LLM: OpenRouter (GPT-4o-mini по умолчанию, настраивается в панели)',
    'База данных: PostgreSQL (Railway) / SQLite (локально)',
    'Парсинг: feedparser (RSS), BeautifulSoup (HTML), прокси-ротация',
    'Деплой: Railway + Docker, автодеплой при push в master',
])

doc.add_page_break()

# ══════════════════════════════════════════════
# 2. QUICK START
# ══════════════════════════════════════════════

doc.add_heading('2. Быстрый старт', level=1)

doc.add_heading('Первый вход', level=2)
doc.add_paragraph('Откройте панель управления в браузере:')
add_info_box('URL: https://igronews-production.up.railway.app/')
add_info_box('Логин: admin  |  Пароль: admin123', 'orange')

doc.add_heading('Что происходит автоматически', level=2)
add_numbered([
    'Парсинг: каждые 5–30 минут система парсит все 15 источников через RSS/HTML/JSON',
    'Авто-ревью: после каждого парсинга новости оцениваются по 12 критериям (бесплатно, без API)',
    'Дедупликация: дубликаты обнаруживаются по URL + TF-IDF cosine + entity overlap',
    'Авто-отклонение: новости с total_score < 15 отклоняются автоматически',
    'Дайджест: ежедневно в 23:00 МСК генерируется LLM-дайджест топ-20 новостей',
    'Публикация: каждую минуту проверяются запланированные статьи',
])

doc.add_heading('Ежедневный цикл работы редактора', level=2)
add_numbered([
    'Откройте вкладку «Редакция» — просмотрите новости «на проверке»',
    'Нажмите «Проверить новые» если есть непроверенные (new) — батч 20 штук',
    'Одобрите лучшие вручную или «Авто скор>70» для массового одобрения',
    'Или нажмите «Полный автомат» — система сделает всё сама (скор → LLM → рерайт → Sheets)',
    'Откройте «Обогащённые» — проверьте результаты Keys.so, Trends, LLM',
    'Откройте «Финал» — здесь только publish_now с финальным скором',
    'Отправьте лучшие на рерайт или экспортируйте в Sheets',
    'Откройте «Контент» — отредактируйте и опубликуйте статьи',
])

add_info_box('Парсинг и авто-ревью полностью бесплатны (локальные). API вызываются только при одобрении или полном автомате.', 'green')

doc.add_page_break()

# ══════════════════════════════════════════════
# 3. PIPELINE
# ══════════════════════════════════════════════

doc.add_heading('3. Рабочий процесс (пайплайн)', level=1)

doc.add_paragraph(
    'Каждая новость проходит чёткий жизненный цикл. '
    'Статус определяет, на каком этапе она находится и в какой вкладке отображается.'
)

doc.add_heading('Статусы новостей', level=2)

add_styled_table(
    ['Статус', 'Описание', 'Где видна', 'Следующий шаг'],
    [
        ['new', 'Только спарсена, не проверена', 'Редакция', 'Авто-ревью → in_review'],
        ['in_review', 'Оценена по 12 критериям, ждёт решения', 'Редакция', 'Одобрить / отклонить'],
        ['duplicate', 'Обнаружен дубликат (терминальный)', 'Редакция', '—'],
        ['rejected', 'Отклонена редактором или авто (score<15)', 'Редакция', '—'],
        ['approved', 'Одобрена, идёт фоновое обогащение API', 'Обогащённые', 'Авто → processed'],
        ['processed', 'Обогащена (Keys.so, Trends, LLM готовы)', 'Обогащённые, Финал', 'В контент / Sheets'],
        ['moderation', 'Отправлена через «Без LLM» (ручной обзор)', 'Редакция', 'Одобрить / отклонить'],
        ['ready', 'Экспортирована, готова к публикации', 'Обогащённые, Финал', 'Опубликовать'],
    ],
    col_widths=[2.5, 4.5, 3, 3.5]
)

doc.add_heading('Схема пайплайна', level=2)

pipeline_steps = [
    ('① Парсинг', 'Автоматически каждые 5–30 мин. RSS/HTML/JSON → таблица news. Статус: new'),
    ('② Авто-ревью', 'Бесплатно, 12 локальных проверок (качество, виральность, свежесть, дедупликация). Статус: in_review / duplicate / rejected'),
    ('③ Решение', 'Ручное одобрение (кнопка ✓), «Авто скор>70», или «Полный автомат». Статус: approved'),
    ('④ Обогащение', 'Keys.so (частоты) + Google Trends (4 региона) + LLM (прогноз тренда). Статус: processed'),
    ('⑤ Финал', 'Рассчитывается финальный композитный скор. Только publish_now попадают в вкладку «Финал»'),
    ('⑥ Рерайт', 'LLM переписывает новость в выбранном стиле. Создаётся статья в таблице articles'),
    ('⑦ Экспорт', 'В Google Sheets (вкладка Ready) + статус: ready'),
    ('⑧ Публикация', 'Ручная или по расписанию (планировщик)'),
]
for step_title, step_desc in pipeline_steps:
    p = doc.add_paragraph()
    run = p.add_run(step_title + '  ')
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
    run2 = p.add_run(step_desc)
    run2.font.size = Pt(10.5)

add_info_box('API-вызовы (Keys.so, Trends, LLM) происходят ТОЛЬКО после одобрения. Парсинг и ревью — бесплатны.', 'orange')

doc.add_page_break()

# ══════════════════════════════════════════════
# 4. FULL AUTO PIPELINE
# ══════════════════════════════════════════════

doc.add_heading('4. Полный автомат — автоматический конвейер', level=1)

doc.add_paragraph(
    'Кнопка «Полный автомат» запускает полностью автоматическую обработку. '
    'Система берёт выбранные или все новые новости и проводит их через 6 стадий '
    'без участия редактора.'
)

add_info_box('Полный автомат использует платные API (Keys.so, Google Trends, LLM). Стоимость: ~$0.01–0.05 за новость.', 'orange')

doc.add_heading('6 стадий полного автомата', level=2)

add_styled_table(
    ['Стадия', 'Название', 'Описание', 'Условие перехода'],
    [
        ['1', 'Скоринг', '12 локальных проверок (бесплатно). Дубли и авто-отклонения отсеиваются.', 'total_score > 0, не дубль'],
        ['2', 'Фильтр скора', 'Проверка: достаточно ли высокий внутренний скор для API-анализа?', 'total_score ≥ 70'],
        ['3', 'Обогащение', 'Keys.so + Google Trends + LLM forecast_trend. Статус → processed.', 'Успешный ответ API'],
        ['4', 'Финальный скор', 'Расчёт композитного финального скора по формуле (5 компонентов).', 'final_score ≥ 60'],
        ['5', 'Рерайт', 'LLM переписывает новость. Создаётся статья в таблице articles.', 'Успешный рерайт'],
        ['6', 'Экспорт', 'Запись в Google Sheets (вкладка Ready). Статус → ready.', '—'],
    ],
    col_widths=[1, 2.5, 6.5, 4]
)

doc.add_heading('Пороги фильтрации', level=2)

add_styled_table(
    ['Порог', 'Значение', 'Назначение', 'Что происходит при непрохождении'],
    [
        ['Авто-отклонение', 'total_score < 15', 'Отсев мусора и некачественных', 'Статус → rejected, задача skipped'],
        ['Фильтр на LLM', 'total_score < 70', 'Экономия API: слабые не отправляются', 'Задача skipped с причиной «Скор N < 70»'],
        ['Финальный фильтр', 'final_score < 60', 'Только лучшие идут на рерайт', 'Задача done с причиной «Финальный скор N < 60»'],
    ],
    col_widths=[3, 3, 4.5, 4.5]
)

doc.add_heading('Что видно в очереди задач', level=2)
doc.add_paragraph('Каждая задача показывает:')
add_bullets([
    'Текущая стадия (Скоринг → Фильтр скора → Обогащение → Финальный скор → Рерайт → Экспорт → Готово)',
    'Внутренний скор (скор: N)',
    'Финальный скор (финал: N) — после обогащения',
    'Причина пропуска, если задача отфильтрована',
    'Номер строки в Sheets, если экспорт завершён',
    'Название рерайта, если рерайт завершён',
])

doc.add_heading('Пример прохождения', level=2)
doc.add_paragraph(
    'Новость «Sony announces PS5 Pro» с 200 новостями в пакете:\n'
    '• 120 отсеяны как дубли или rejected (score<15)\n'
    '• 50 пропущены: внутренний скор < 70\n'
    '• 30 отправлены на LLM-обогащение\n'
    '• 12 отфильтрованы: финальный скор < 60\n'
    '• 18 прошли рерайт → экспорт в Sheets Ready'
)

add_info_box('Кнопка «Стоп» останавливает пайплайн. Оставшиеся задачи помечаются как cancelled.', 'red')

doc.add_page_break()

# ══════════════════════════════════════════════
# 5. EDITORIAL TAB
# ══════════════════════════════════════════════

doc.add_heading('5. Вкладка «Редакция»', level=1)
doc.add_paragraph('Главная рабочая вкладка. Здесь происходит основная работа с новостями.')

doc.add_heading('Карточки статистики', level=2)
doc.add_paragraph('В верхней части — кликабельные карточки. Клик фильтрует таблицу по статусу.')
add_styled_table(
    ['Карточка', 'Что показывает', 'Клик'],
    [
        ['Всего', 'Общее число новостей в базе', 'Сброс фильтров'],
        ['Новые', 'Непроверенные (new)', 'Фильтр status=new'],
        ['На ревью', 'Оценены, ждут решения (in_review)', 'Фильтр status=in_review'],
        ['Дубли', 'Обнаруженные дубликаты', 'Фильтр status=duplicate'],
        ['Одобрено', 'Отправлены на обогащение', 'Фильтр status=approved'],
        ['Обработано', 'Обогащены API-данными', 'Фильтр status=processed'],
        ['Отклонено', 'Отклонённые', 'Фильтр status=rejected'],
        ['Готовые', 'Готовы к публикации', 'Фильтр status=ready'],
    ],
    col_widths=[2.5, 5, 6]
)

doc.add_heading('Фильтры', level=2)
add_styled_table(
    ['Фильтр', 'Тип', 'Описание'],
    [
        ['Статус', 'Выпадающий список', 'new, in_review, approved, processed, rejected, duplicate, ready, moderation'],
        ['Источник', 'Выпадающий список', 'IGN, DTF, StopGame и др. (все 15 источников)'],
        ['Виральность', 'Выпадающий список', 'LOW / MEDIUM / HIGH'],
        ['Тир сущности', 'Выпадающий список', 'S (AAA) / A (Major) / B (Notable) / C (Niche)'],
        ['Мин. скор', 'Ползунок', 'Минимальный total_score (0–100)'],
        ['Поиск', 'Текстовое поле', 'Поиск по заголовку, URL, описанию'],
    ],
    col_widths=[3, 3, 9]
)

doc.add_heading('Кнопки действий', level=2)

add_styled_table(
    ['Кнопка', 'Что делает', 'Когда использовать'],
    [
        ['Проверить новые', 'Авто-ревью для 20 новых (бесплатно, 12 проверок)', 'Есть непроверенные new'],
        ['Полный автомат', 'Скор→>70→LLM→финал→>60→рерайт→Sheets', 'Автоматическая обработка (платно)'],
        ['Без LLM', 'Скор→Sheets/NotReady, статус moderation', 'Экономный режим без API'],
        ['Авто скор>70', 'Массово одобряет все in_review с score≥70', 'Быстрое одобрение лучших'],
        ['✓ Одобрить', 'Одобряет → фоновое обогащение API', 'Ручное одобрение конкретной новости'],
        ['✗ Отклонить', 'Отклоняет новость (терминально)', 'Неподходящая новость'],
        ['✎ В контент', 'Отправляет в очередь рерайта', 'Хотите переписать через LLM'],
        ['Перевод', 'Переводит заголовок на русский (LLM)', 'Английский заголовок'],
        ['В Sheets', 'Экспорт в Google Sheets (Лист1)', 'Ручной экспорт'],
        ['Удалить', 'Удаляет выбранные (необратимо)', 'Очистка мусора'],
        ['Стоп', 'Останавливает пайплайн', 'Нужно прервать автомат'],
    ],
    col_widths=[3.5, 5.5, 5.5]
)

doc.add_heading('Раскрываемые детали', level=2)
doc.add_paragraph(
    'Клик по строке раскрывает детали новости:'
)
add_bullets([
    'Полный текст новости (до 500 символов)',
    'Все баллы проверок: quality, relevance, freshness, viral, momentum, headline',
    'Виральные триггеры с иконками',
    'Тональность: positive / neutral / negative',
    'Обнаруженные сущности: игры, студии, платформы (с тиром)',
    'Авто-теги: release, update, esports, industry и др.',
    'Кнопки: одобрить, отклонить, в контент, перевести, в Sheets',
    'Временные кластеры: связанные новости (если обнаружены)',
])

doc.add_page_break()

# ══════════════════════════════════════════════
# 6. ENRICHED TAB
# ══════════════════════════════════════════════

doc.add_heading('6. Вкладка «Обогащённые»', level=1)
doc.add_paragraph(
    'Показывает новости после обогащения API-данными (статусы: approved, processed, ready). '
    'Здесь видны результаты Keys.so, Google Trends и LLM-анализа.'
)

doc.add_heading('Колонки таблицы', level=2)
add_styled_table(
    ['Колонка', 'Описание', 'Источник данных'],
    [
        ['Биграммы', 'Ключевые словосочетания (топ-5 пар слов)', 'TF-IDF анализ текста (бесплатно)'],
        ['Keys.so (freq)', 'Частота поискового запроса (0–999999)', 'Keys.so API (платно)'],
        ['Похожие', 'Количество похожих запросов', 'Keys.so API similar keywords'],
        ['Trends RU/US', 'Индекс Google Trends (0–100) по регионам', 'Google Trends API (бесплатно)'],
        ['LLM рекомендация', 'publish_now / publish_later / rewrite / skip', 'LLM forecast_trend (платно)'],
        ['LLM Score', 'Оценка трендового потенциала (0–100)', 'LLM forecast_trend'],
        ['Скор', 'Внутренний total_score из 12 проверок', 'Локальный scoring'],
    ],
    col_widths=[3, 5, 6.5]
)

doc.add_heading('Фильтры', level=2)
add_bullets([
    'Все обогащённые — approved + processed + ready (по умолчанию)',
    'Одобренные — ждут обогащения (approved)',
    'Обогащённые — есть API-данные (processed)',
    'Готовые — экспортированы (ready)',
    'LLM фильтр: publish_now / publish_later / rewrite / skip / есть / нет рекомендации',
    'Источник: фильтр по сайту',
])

doc.add_heading('Автоматический выбор региона Keys.so', level=2)
doc.add_paragraph(
    'Система автоматически определяет регион для Keys.so по источнику:'
)
add_styled_table(
    ['Регион', 'Источники'],
    [
        ['RU (Россия)', 'DTF, StopGame, Playground.ru, iXBT.games, VGTimes'],
        ['US (США)', 'IGN, GameSpot, PCGamer, GameRant, Eurogamer, Kotaku, GamesRadar, Polygon, Destructoid, RockPaperShotgun'],
    ],
    col_widths=[3, 12]
)

doc.add_page_break()

# ══════════════════════════════════════════════
# 7. FINAL TAB
# ══════════════════════════════════════════════

doc.add_heading('7. Вкладка «Финал»', level=1)
doc.add_paragraph(
    'Финальная подборка лучших новостей. Попадают ТОЛЬКО новости с LLM-рекомендацией '
    'publish_now и статусом processed или ready. Здесь все данные на одном экране '
    'и вычисленный финальный балл.'
)

add_info_box('Это главная вкладка для принятия решения «что публиковать». Финальный скор — итоговая метрика.', 'green')

doc.add_heading('14 колонок', level=2)
add_styled_table(
    ['#', 'Колонка', 'Описание'],
    [
        ['1', 'Чекбокс', 'Выбор для массовых действий'],
        ['2', 'Источник', 'Сайт-источник (IGN, DTF и др.)'],
        ['3', 'Заголовок', 'Кликабельная ссылка на оригинал'],
        ['4', 'Скор', 'Внутренний total_score (0–100)'],
        ['5', 'Вирал', 'Виральный балл (0–100) с цветом и тултипом триггеров'],
        ['6', 'Свеж.', 'Время с публикации. Зелёный <6ч, жёлтый <24ч, красный >24ч'],
        ['7', 'Тон', 'Тональность: positive / neutral / negative (иконка)'],
        ['8', 'Теги', 'Авто-теги (release, update, esports и др.)'],
        ['9', 'Биграммы', 'Ключевые словосочетания из TF-IDF'],
        ['10', 'Keys.so', 'Частота поискового запроса'],
        ['11', 'Похож.', 'Количество похожих запросов'],
        ['12', 'Trends', 'Google Trends по регионам (RU:50 US:80 и т.д.)'],
        ['13', 'Финал', 'Финальный композитный скор (0–100) с цветом'],
        ['14', 'Действия', 'Кнопки: ✎ (в контент), ☰ (в Sheets)'],
    ],
    col_widths=[0.8, 2.5, 11]
)

doc.add_heading('Карточки статистики', level=2)
add_bullets([
    'Всего publish_now — количество новостей в финале',
    'Финал ≥ 60 — количество с высоким финальным скором (зелёный)',
    'Средний финал — среднее значение финального скора',
])

doc.add_heading('Кнопки', level=2)
add_styled_table(
    ['Кнопка', 'Что делает'],
    [
        ['Отметить >60', 'Автоматически выделяет все новости с финальным скором > 60'],
        ['✎ В контент', 'Отправляет выбранные на рерайт через очередь задач'],
        ['☰ В Sheets', 'Экспортирует выбранные в Google Sheets'],
    ],
    col_widths=[4, 11]
)

doc.add_heading('Тултип виральности', level=2)
doc.add_paragraph(
    'При наведении на виральный скор показывается тултип со списком активных триггеров. '
    'Например: «Скандал (+15), AAA-игра (+30), Breaking (+10)».'
)

doc.add_heading('Серверная пагинация', level=2)
doc.add_paragraph(
    'Вкладка «Финал» использует серверную пагинацию: 50 новостей на страницу. '
    'Навигация внизу таблицы (◀ Назад / Вперёд ▶). Это обеспечивает быструю загрузку '
    'даже при тысячах обработанных новостей.'
)

doc.add_page_break()

# ══════════════════════════════════════════════
# 8. CONTENT TAB
# ══════════════════════════════════════════════

doc.add_heading('8. Вкладка «Контент»', level=1)
doc.add_paragraph(
    'Единая вкладка для работы со статьями и рерайтом. '
    'Подвкладки: «Статьи» (сохранённые) и «Новости» (для рерайта).'
)

doc.add_heading('Подвкладка «Статьи»', level=2)
add_bullets([
    'Список всех сохранённых статей с поиском',
    'Фильтр по статусу: Черновики / Готовые / Опубликованные / Запланированные',
    'Массовое скачивание в DOCX (ZIP-архив)',
    'Массовое удаление',
    'Клик по статье → открывается в редакторе справа',
])

doc.add_heading('Подвкладка «Новости»', level=2)
add_bullets([
    'Список одобренных и обработанных новостей',
    'Фильтр по источнику',
    'Выбор стиля рерайта (6 стилей)',
    'Массовый рерайт выбранных',
    'Быстрый предпросмотр: клик → детали справа',
])

doc.add_heading('Редактор статьи (правая панель)', level=2)
add_styled_table(
    ['Поле', 'Описание', 'Ограничения'],
    [
        ['Заголовок', 'Основной заголовок статьи', 'Счётчик символов'],
        ['SEO Title', 'Заголовок для поисковиков', 'Оптимально до 60 символов'],
        ['SEO Description', 'Мета-описание', 'Оптимально до 155 символов'],
        ['Теги', 'Через запятую', 'Авто-генерация при рерайте'],
        ['Текст статьи', 'Полное редактирование', 'Счётчик слов и символов'],
        ['Статус', 'Черновик → Готово → Опубликовано', 'Выпадающий список'],
        ['Планировщик', 'Дата и время отложенной публикации', 'ISO datetime picker'],
    ],
    col_widths=[3, 6.5, 4.5]
)

doc.add_heading('Стили рерайта', level=2)
add_styled_table(
    ['Стиль', 'Описание', 'Длина'],
    [
        ['news', 'Классическая новостная статья', '3–5 абзацев'],
        ['seo', 'SEO-оптимизированный текст с ключевиками', '4–6 абзацев'],
        ['review', 'Обзорный стиль с мнением автора', '4–6 абзацев'],
        ['clickbait', 'Цепляющий заголовок и интригующее начало', '3–4 абзаца'],
        ['short', 'Краткая заметка', '2–3 абзаца'],
        ['social', 'Пост для соцсетей (Twitter, Telegram)', '1–2 абзаца'],
    ],
    col_widths=[3, 7, 4]
)

doc.add_heading('SEO-анализ (кнопка в редакторе)', level=2)
doc.add_paragraph(
    'Кнопка «SEO-проверка» анализирует текущую статью по 9 критериям и выдаёт итоговый балл.'
)

doc.add_page_break()

# ══════════════════════════════════════════════
# 9. VIRAL TAB
# ══════════════════════════════════════════════

doc.add_heading('9. Вкладка «Виральность»', level=1)
doc.add_paragraph('Анализ виральности новостей. Показывает новости с наибольшим виральным потенциалом.')

doc.add_heading('Уровни виральности', level=2)
add_styled_table(
    ['Уровень', 'Цвет', 'Диапазон viral_score', 'Описание'],
    [
        ['HIGH', 'Красный', '≥ 60', 'Высокий виральный потенциал, срочно'],
        ['MEDIUM', 'Оранжевый', '30–59', 'Средний потенциал, стоит рассмотреть'],
        ['LOW', 'Серый', '< 30', 'Низкий потенциал'],
    ],
    col_widths=[2, 2, 3.5, 7]
)

doc.add_heading('Триггеры виральности', level=2)
doc.add_paragraph('50+ встроенных триггеров, сгруппированных по категориям:')
add_styled_table(
    ['Категория', 'Примеры', 'Вес'],
    [
        ['Скандалы', 'Увольнения студий, контроверсии, разочарования', '10–15'],
        ['Утечки', 'Неанонсированные игры, инсайды, датамайнинг', '10–15'],
        ['Релизы', 'Shadow drop, неожиданные релизы, анонсы', '10–12'],
        ['Провалы', 'Баги, краши, негативные отзывы, даунгрейд', '10–15'],
        ['AI', 'Использование ИИ в играх, генерация контента', '8–12'],
        ['События', 'E3, gamescom, TGA, State of Play, Summer Game Fest', '10–12'],
        ['M&A', 'Слияния: Microsoft-Activision, Sony-Bungie', '12–15'],
        ['Киберспорт', 'Турниры, трансферы, призовые', '8–10'],
        ['Железо', 'Новые консоли, видеокарты, VR-устройства', '10–12'],
        ['Рекорды', 'Продажи, онлайн, рейтинги, достижения', '8–12'],
    ],
    col_widths=[2.5, 6.5, 1.5]
)

doc.add_heading('Entity Boost (бонус за сущности)', level=2)
add_styled_table(
    ['Тир', 'Бонус к viral_score', 'Примеры'],
    [
        ['S (AAA)', '+30', 'GTA VI, Elder Scrolls VI, Nintendo, Rockstar, Sony PlayStation'],
        ['A (Major)', '+25', 'Cyberpunk 2077, Elden Ring, Ubisoft, EA, Valve'],
        ['B (Notable)', '+15', "Baldur's Gate 3, Hollow Knight: Silksong, Capcom"],
        ['C (Niche)', '+5', 'Indie-игры, малые студии'],
    ],
    col_widths=[2.5, 3, 9]
)

doc.add_heading('Действия', level=2)
add_styled_table(
    ['Кнопка', 'Что делает'],
    [
        ['В редактор', 'Отправляет одну новость в Контент на рерайт'],
        ['Отправить HIGH', 'Массово отправляет все high-виральные'],
        ['Отправить MEDIUM+', 'Отправляет medium + high'],
        ['Отправить все', 'Отправляет все виральные новости'],
    ],
    col_widths=[4, 11]
)

doc.add_page_break()

# ══════════════════════════════════════════════
# 10. ANALYTICS TAB
# ══════════════════════════════════════════════

doc.add_heading('10. Вкладка «Аналитика»', level=1)

doc.add_heading('Графики (Canvas)', level=2)
add_styled_table(
    ['График', 'Описание', 'Данные'],
    [
        ['Тренд среднего скора', 'Линейный график среднего total_score по дням за 30 дней', 'AVG(total_score) GROUP BY date'],
        ['Конверсия по дням', 'Столбчатый: одобренные vs отклонённые за 30 дней', 'COUNT по статусам GROUP BY date'],
    ],
    col_widths=[4, 5.5, 5]
)

doc.add_heading('Статистика', level=2)
add_bullets([
    'Распределение по статусам: количество новостей в каждом статусе',
    'Топ источников: по количеству одобренных / отклонённых',
    'Процент одобрений по источникам за 30 дней',
    'Среднее время обработки',
])

doc.add_heading('Дайджесты', level=2)
doc.add_paragraph(
    'Автоматическая генерация ежедневного дайджеста в 23:00 МСК. '
    'Топ-20 новостей за сутки по total_score, создаётся через LLM.'
)
add_bullets([
    'Стили: brief (краткий), standard (стандартный), detailed (подробный)',
    'История дайджестов: просмотр всех сгенерированных с датами',
    'Ручная генерация: кнопка «Сгенерировать дайджест»',
])

doc.add_heading('Версии промптов', level=2)
doc.add_paragraph(
    'Версионирование LLM-промптов для A/B-тестирования. '
    'Каждая версия отслеживает средний скор и количество использований.'
)

doc.add_page_break()

# ══════════════════════════════════════════════
# 11. QUEUE TAB
# ══════════════════════════════════════════════

doc.add_heading('11. Вкладка «Очередь»', level=1)
doc.add_paragraph('Управление фоновыми задачами. Все пайплайны создают задачи в очереди.')

doc.add_heading('Кликабельные карточки статусов', level=2)
add_styled_table(
    ['Карточка', 'Статус задачи', 'Цвет', 'Описание'],
    [
        ['Ожидает', 'pending', 'Жёлтый', 'В очереди, ещё не начата'],
        ['Работает', 'running', 'Синий', 'В процессе выполнения'],
        ['Готово', 'done', 'Зелёный', 'Успешно завершена'],
        ['Пропущено', 'skipped', 'Серый', 'Пропущена (дубль, низкий скор)'],
        ['Ошибка', 'error', 'Красный', 'Ошибка при выполнении'],
        ['Отменено', 'cancelled', 'Серый', 'Отменена пользователем (Стоп)'],
    ],
    col_widths=[2.5, 2, 2, 7.5]
)

doc.add_paragraph('Клик по карточке фильтрует таблицу задач по этому статусу. Повторный клик сбрасывает фильтр.')

doc.add_heading('Типы задач', level=2)
add_styled_table(
    ['Тип', 'Иконка', 'Описание'],
    [
        ['full_auto', '🚀', 'Полный автомат: скор → LLM → финал → рерайт → Sheets'],
        ['no_llm', '📋', 'Без LLM: скор → Sheets/NotReady → moderation'],
        ['rewrite', '✎', 'Рерайт одной новости через LLM'],
        ['sheets', '📄', 'Экспорт в Google Sheets'],
        ['mod_rewrite', '✎', 'Рерайт из модерации'],
    ],
    col_widths=[3, 1.5, 10]
)

doc.add_heading('Информация в задаче', level=2)
doc.add_paragraph('Каждая задача показывает развёрнутый результат:')
add_bullets([
    'Стадия: Скоринг → Фильтр скора → Обогащение → Финальный скор → Рерайт → Экспорт → Готово',
    'Внутренний скор (скор: N)',
    'Финальный скор (финал: N) — после стадии обогащения',
    'Причина пропуска/фильтрации',
    'Номер строки в Sheets (при экспорте)',
    'Название рерайта (при рерайте)',
    'Текст ошибки (при ошибке)',
])

doc.add_page_break()

# ══════════════════════════════════════════════
# 12. HEALTH TAB
# ══════════════════════════════════════════════

doc.add_heading('12. Вкладка «Здоровье»', level=1)
doc.add_paragraph('Мониторинг работоспособности источников новостей за последние 30 дней.')

add_styled_table(
    ['Метрика', 'Описание', 'Как считается'],
    [
        ['Всего статей', 'Количество спарсенных за 30 дней', 'COUNT WHERE parsed_at > 30d'],
        ['Одобрено', 'Одобренные редактором', 'COUNT WHERE status IN (approved, processed, ready)'],
        ['Отклонено', 'Отклонённые', 'COUNT WHERE status = rejected'],
        ['% одобрений', 'Процент одобрений от решений', 'approved / (approved + rejected) × 100'],
        ['Вес', 'Текущий вес источника (0.5–2.0)', 'Базовый 1.0 ± корректировки'],
        ['Новые', 'Количество необработанных', 'COUNT WHERE status = new'],
    ],
    col_widths=[3, 5, 6.5]
)

doc.add_heading('Автокоррекция весов', level=2)
doc.add_paragraph(
    'Вес источника корректируется автоматически на основе истории решений редактора:'
)
add_styled_table(
    ['Условие', 'Действие', 'Диапазон'],
    [
        ['> 80% одобрений за 30 дней', '+0.2 к весу', 'Макс. 2.0'],
        ['< 30% одобрений за 30 дней', '-0.2 от веса', 'Мин. 0.5'],
        ['30–80% одобрений', 'Без изменений', '—'],
    ],
    col_widths=[5, 4, 4]
)

doc.add_paragraph('Вес влияет на total_score: чем выше вес источника, тем выше балл его новостей.')

doc.add_page_break()

# ══════════════════════════════════════════════
# 13. SETTINGS TAB
# ══════════════════════════════════════════════

doc.add_heading('13. Вкладка «Настройки» (⚙)', level=1)
doc.add_paragraph('Конфигурация системы. 6 подвкладок:')

add_styled_table(
    ['Подвкладка', 'Содержимое', 'Ключевые настройки'],
    [
        ['Общие', 'Модель LLM, порог авто-одобрения, очистка', 'LLM model, AUTO_APPROVE_THRESHOLD, clear cache/DB'],
        ['Источники', '15 источников, добавление/удаление', 'Имя, тип (rss/html), URL, интервал, CSS-селектор'],
        ['Промпты', 'LLM-промпты с версионированием', 'forecast, rewrite, translate, merge'],
        ['Инструменты', 'Тест API-подключений', 'Тест LLM, Keys.so, Sheets, перевод'],
        ['Логи', 'Последние 50 записей', 'ERROR / WARNING / INFO с фильтрацией'],
        ['Пользователи', 'Управление аккаунтами', 'Логин, пароль, роль'],
    ],
    col_widths=[2.5, 5, 7]
)

doc.add_page_break()

# ══════════════════════════════════════════════
# 14. SOURCES
# ══════════════════════════════════════════════

doc.add_heading('14. Источники новостей (15 сайтов)', level=1)

add_styled_table(
    ['#', 'Источник', 'Тип', 'Интервал', 'Язык', 'Keys.so регион'],
    [
        ['1', 'IGN', 'RSS', '5 мин', 'EN', 'US'],
        ['2', 'GameSpot', 'RSS', '10 мин', 'EN', 'US'],
        ['3', 'PCGamer', 'RSS', '10 мин', 'EN', 'US'],
        ['4', 'GameRant', 'RSS', '10 мин', 'EN', 'US'],
        ['5', 'DTF', 'JSON (API)', '10 мин', 'RU', 'RU'],
        ['6', 'StopGame', 'HTML', '10 мин', 'RU', 'RU'],
        ['7', 'Eurogamer', 'RSS', '15 мин', 'EN', 'US'],
        ['8', 'Kotaku', 'RSS', '15 мин', 'EN', 'US'],
        ['9', 'GamesRadar', 'RSS', '15 мин', 'EN', 'US'],
        ['10', 'Polygon', 'RSS', '15 мин', 'EN', 'US'],
        ['11', 'Destructoid', 'RSS', '15 мин', 'EN', 'US'],
        ['12', 'Playground.ru', 'RSS', '15 мин', 'RU', 'RU'],
        ['13', 'RockPaperShotgun', 'RSS', '30 мин', 'EN', 'US'],
        ['14', 'iXBT.games', 'HTML', '30 мин', 'RU', 'RU'],
        ['15', 'VGTimes', 'HTML', '30 мин', 'RU', 'RU'],
    ],
    col_widths=[1, 4, 2, 2, 1.5, 3]
)

doc.add_heading('Типы парсеров', level=2)
add_styled_table(
    ['Тип', 'Технология', 'Описание'],
    [
        ['RSS', 'feedparser + fetch_with_retry', 'Стандартный RSS/Atom фид. Самый надёжный.'],
        ['HTML', 'BeautifulSoup + CSS-селекторы', 'Парсинг HTML-страниц. Нужен CSS-селектор.'],
        ['JSON (API)', 'requests + JSON парсинг', 'Для сайтов с API (DTF).'],
        ['Sitemap', 'XML парсинг', 'Для сайтов без RSS/API.'],
    ],
    col_widths=[3, 4.5, 7]
)

doc.add_page_break()

# ══════════════════════════════════════════════
# 15. SCORING SYSTEM
# ══════════════════════════════════════════════

doc.add_heading('15. Система проверок — 12 критериев', level=1)
doc.add_paragraph(
    'Каждая новость проходит 12 автоматических проверок при авто-ревью. '
    'Все проверки бесплатны (без API). Результаты формируют total_score.'
)

add_styled_table(
    ['#', 'Проверка', 'Диапазон', 'Что оценивает', 'Как работает'],
    [
        ['1', 'Quality', '0–100', 'Качество текста', 'Длина текста (≥150 симв.), заголовка (≥20), наличие кликбейта, мусора'],
        ['2', 'Relevance', '0–100', 'Релевантность геймингу', 'Количество гейминг-ключевиков vs общий шум'],
        ['3', 'Freshness', '0–100', 'Свежесть', '<2ч→100, <6ч→80, <12ч→65, <24ч→50, <48ч→30, <72ч→25, >72ч→10'],
        ['4', 'Viral Score', '0–100', 'Виральный потенциал', '50+ триггеров + entity boost (S/A/B/C) + комбо-бонус'],
        ['5', 'Momentum', '0–100', 'Моментум темы', 'Сколько источников пишут о том же за 1ч/6ч/24ч (word overlap)'],
        ['6', 'Headline', '0–100', 'Сила заголовка', 'Числа (+10), вопрос (+8), breaking (+15), эксклюзив (+12), длина'],
        ['7', 'Source Weight', '0.5–2.0', 'Надёжность источника', 'Базовый 1.0 + авто-коррекция по решениям за 30 дней'],
        ['8', 'Sentiment', '-1..+1', 'Тональность', 'Positive/neutral/negative анализ текста'],
        ['9', 'Tags', 'список', 'Авто-теги', '9 тегов: industry, release, update, esports, hardware, mobile, review, deal, opinion'],
        ['10', 'Dedup', 'bool', 'Дедупликация', 'TF-IDF cosine + entity overlap, порог 0.7'],
        ['11', 'NER', 'dict', 'Сущности', '80+ игр, 30+ студий, 10+ платформ с тирами S/A/B/C'],
        ['12', 'Feedback', '±10', 'Обучаемость', 'Корректировка на основе истории решений (source + tag)'],
    ],
    col_widths=[0.7, 2.5, 1.5, 3, 7]
)

doc.add_heading('Early Exit (оптимизация)', level=2)
doc.add_paragraph(
    'Если quality < 20 — тяжёлые проверки (momentum, NER, viral) пропускаются для экономии ресурсов. '
    'Новость всё равно получает скор (но низкий).'
)

doc.add_page_break()

# ══════════════════════════════════════════════
# 16. SCORING FORMULAS
# ══════════════════════════════════════════════

doc.add_heading('16. Формулы скоринга — все расчёты', level=1)

doc.add_heading('16.1. total_score (внутренний скор)', level=2)
add_formula(
    'Формула',
    '(quality + relevance + freshness + viral) / 4 + momentum/5 + source_weight_bonus + headline_bonus + feedback',
    'Диапазон: 0–100+. Среднее 4 основных проверок + бонусы.'
)

add_styled_table(
    ['Компонент', 'Формула', 'Диапазон', 'Пример'],
    [
        ['Базовый скор', '(quality + relevance + freshness + viral) / 4', '0–100', '(80+70+100+45)/4 = 73.75'],
        ['Моментум бонус', 'momentum_score / 5', '0–20', '60/5 = +12'],
        ['Вес источника', '(source_weight - 1.0) × 10', '-5..+10', '(1.4-1.0)×10 = +4'],
        ['Заголовок бонус', 'headline_score / 10', '0–10', '70/10 = +7'],
        ['Feedback', '±10 из истории решений', '-10..+10', '+5 (источник часто одобряется)'],
    ],
    col_widths=[3.5, 5, 2, 4]
)

add_info_box('Авто-отклонение: total_score < 15 → rejected. Дубликаты: status → duplicate.', 'red')

doc.add_heading('16.2. Quality Score', level=2)
add_formula(
    'Формула',
    'base_quality - clickbait_penalty - spam_penalty',
)
add_styled_table(
    ['Критерий', 'Условие', 'Баллы'],
    [
        ['Длина текста', '≥ 150 символов', '+50 базовых'],
        ['Длина текста', '< 150 символов', '25 базовых'],
        ['Длина заголовка', '≥ 20 символов', '+20'],
        ['Длина заголовка', '< 20 символов', '+5'],
        ['H1 присутствует', 'len(h1) > 0', '+10'],
        ['Описание', 'len(description) > 0', '+10'],
        ['Кликбейт паттерны', '"YOU WON\'T BELIEVE", "SHOCKING" и др.', '-20'],
        ['Спам паттерны', 'Повторяющиеся символы, capslock', '-15'],
    ],
    col_widths=[4, 5, 5]
)

doc.add_heading('16.3. Freshness Score', level=2)
add_styled_table(
    ['Возраст', 'Баллы', 'Статус'],
    [
        ['< 2 часов', '100', 'Горячая (hot)'],
        ['2–6 часов', '80', 'Свежая (fresh)'],
        ['6–12 часов', '65', 'Нормальная'],
        ['12–24 часа', '50', 'Вчерашняя'],
        ['24–48 часов', '30', 'Старая'],
        ['48–72 часа', '25', 'Очень старая'],
        ['> 72 часов', '10', 'Устаревшая'],
    ],
    col_widths=[3.5, 2, 5]
)

doc.add_heading('16.4. Headline Score', level=2)
add_styled_table(
    ['Критерий', 'Бонус', 'Пример'],
    [
        ['Числа в заголовке', '+10', '"10 лучших RPG 2026"'],
        ['Вопросительный знак', '+8', '"Выйдет ли GTA 6 в 2026?"'],
        ['Breaking / Срочно', '+15', '"BREAKING: Sony купила..."'],
        ['Эксклюзив', '+12', '"Exclusive: новый трейлер..."'],
        ['Длина 40–80 символов', '+10', 'Оптимальная длина заголовка'],
        ['Длина < 20 символов', '-10', 'Слишком короткий'],
        ['Длина > 120 символов', '-5', 'Слишком длинный'],
    ],
    col_widths=[4, 2, 8]
)

doc.add_heading('16.5. Viral Score', level=2)
add_formula(
    'Формула',
    'SUM(trigger_weights) + entity_boost + combo_bonus',
    'Каждый совпавший триггер добавляет свой вес. Entity boost по тиру. Комбо: 3+ триггеров = +10.'
)

doc.add_heading('16.6. Momentum Score', level=2)
add_formula(
    'Формула',
    'word_overlap_matches × multiplier',
    'Подсчёт совпадений ключевых слов с другими новостями за 1ч (×3), 6ч (×2), 24ч (×1).'
)

doc.add_page_break()

# ══════════════════════════════════════════════
# 17. FINAL SCORE
# ══════════════════════════════════════════════

doc.add_heading('17. Финальный скор — композитная формула', level=1)
doc.add_paragraph(
    'Финальный скор вычисляется после обогащения API-данными. '
    'Объединяет внутреннюю оценку и внешние сигналы в единый балл.'
)

add_info_box('Финальный скор = решающая метрика для полного автомата и вкладки «Финал».', 'blue')

doc.add_heading('Формула', level=2)
add_formula(
    'final_score',
    'internal×0.4 + viral×0.2 + keyso_bonus×0.15 + trends_bonus×0.1 + headline×0.15',
)

doc.add_heading('Компоненты подробно', level=2)

add_styled_table(
    ['Компонент', 'Вес', 'Диапазон', 'Источник', 'Описание'],
    [
        ['internal (total_score)', '40%', '0–100', 'Локальные 12 проверок', 'Базовая оценка качества, свежести, релевантности'],
        ['viral (viral_score)', '20%', '0–100', 'Триггеры + entities', 'Виральный потенциал новости'],
        ['keyso_bonus', '15%', '0–100', 'Keys.so API', 'Бонус за поисковую частоту (таблица ниже)'],
        ['trends_bonus', '10%', '0–100', 'Google Trends API', 'Бонус за трендовость (таблица ниже)'],
        ['headline (headline_score)', '15%', '0–100', 'Анализ заголовка', 'Сила заголовка (числа, вопросы, breaking)'],
    ],
    col_widths=[3.5, 1, 1.5, 3.5, 5]
)

doc.add_heading('Keys.so бонус — таблица перевода', level=2)
add_styled_table(
    ['Частота (freq/ws)', 'keyso_bonus', 'Интерпретация'],
    [
        ['≥ 10,000', '100', 'Очень популярный запрос'],
        ['5,000 – 9,999', '80', 'Популярный запрос'],
        ['1,000 – 4,999', '60', 'Средняя популярность'],
        ['100 – 999', '40', 'Низкая популярность'],
        ['1 – 99', '20', 'Минимальный трафик'],
        ['0', '0', 'Нет поискового трафика'],
    ],
    col_widths=[3, 2.5, 9]
)

doc.add_heading('Google Trends бонус — таблица перевода', level=2)
doc.add_paragraph('Берётся максимальное значение по всем регионам (RU, US, GB, DE):')
add_styled_table(
    ['Макс. Trends индекс', 'trends_bonus', 'Интерпретация'],
    [
        ['≥ 80', '100', 'Горячий тренд (пиковый интерес)'],
        ['50 – 79', '70', 'Сильный тренд'],
        ['20 – 49', '40', 'Умеренный интерес'],
        ['1 – 19', '20', 'Слабый интерес'],
        ['0', '0', 'Нет данных / нет интереса'],
    ],
    col_widths=[3.5, 2.5, 9]
)

doc.add_heading('Цветовая индикация финального скора', level=2)
add_styled_table(
    ['Диапазон', 'Цвет', 'Рекомендация', 'В полном автомате'],
    [
        ['≥ 60', 'Зелёный', 'Отличный кандидат на публикацию', '→ рерайт → Sheets Ready'],
        ['35 – 59', 'Жёлтый', 'Рассмотреть, возможно доработать', '→ пропускается (filtered)'],
        ['< 35', 'Красный', 'Слабый, лучше пропустить', '→ пропускается (filtered)'],
    ],
    col_widths=[2, 2, 5, 5.5]
)

doc.add_heading('Пример расчёта', level=2)
doc.add_paragraph('Новость: "GTA 6 trailer breaks YouTube record" от IGN')
add_styled_table(
    ['Компонент', 'Значение', '× Вес', 'Вклад'],
    [
        ['total_score (internal)', '82', '× 0.40', '32.8'],
        ['viral_score', '95 (S-tier entity + скандал)', '× 0.20', '19.0'],
        ['keyso_bonus', '100 (freq = 45,000)', '× 0.15', '15.0'],
        ['trends_bonus', '100 (Trends US = 92)', '× 0.10', '10.0'],
        ['headline_score', '75 (числа + breaking)', '× 0.15', '11.25'],
        ['ИТОГО', '', '', '88 (округлённо)'],
    ],
    col_widths=[5, 4.5, 2, 2.5]
)

doc.add_page_break()

# ══════════════════════════════════════════════
# 18. API INTEGRATIONS
# ══════════════════════════════════════════════

doc.add_heading('18. API-интеграции — подробно', level=1)

doc.add_heading('18.1. Keys.so', level=2)
doc.add_paragraph('Сервис анализа поисковых запросов Яндекс/Google.')
add_styled_table(
    ['Параметр', 'Значение'],
    [
        ['Базовый URL', 'https://api.keys.so'],
        ['Аутентификация', 'auth-token (параметр запроса)'],
        ['Кэш', '24 часа'],
        ['Rate limit', 'Встроенный в систему (2с между запросами)'],
        ['Регион', 'Авто: US для EN-источников, RU для русских'],
    ],
    col_widths=[4, 10.5]
)

doc.add_heading('Эндпоинты Keys.so', level=3)
add_styled_table(
    ['Эндпоинт', 'Описание', 'Возвращает'],
    [
        ['/report/simple/keyword_dashboard', 'Частотность ключевого слова', 'ws (частота), wsk (конкуренция)'],
        ['/report/simple/similarkeys', 'Похожие поисковые запросы (до 10)', 'word, ws, wsk, cnt'],
        ['/tools/keywords_by_list', 'Массовая проверка (создаёт задачу)', 'uid → data[] с ws, wsk'],
    ],
    col_widths=[5, 4.5, 5]
)

doc.add_heading('18.2. Google Trends', level=2)
add_styled_table(
    ['Параметр', 'Значение'],
    [
        ['Библиотека', 'pytrends (неофициальный API)'],
        ['Регионы', 'RU, US, GB, DE (все одновременно)'],
        ['Период', 'Последние 24 часа (now 1-d)'],
        ['Кэш', '6 часов'],
        ['Rate limit', '3с между запросами'],
        ['Возвращает', 'Индекс 0–100 для каждого региона'],
    ],
    col_widths=[4, 10.5]
)

doc.add_heading('18.3. LLM (OpenRouter)', level=2)
add_styled_table(
    ['Параметр', 'Значение'],
    [
        ['Провайдер', 'OpenRouter (router для 100+ моделей)'],
        ['Модель по умолч.', 'GPT-4o-mini (настраивается в панели)'],
        ['Фолбэк', '2 API-ключа с автоматическим переключением'],
        ['Rate limit', '2с между запросами'],
        ['Формат ответа', 'JSON (все промпты возвращают структурированный JSON)'],
    ],
    col_widths=[4, 10.5]
)

doc.add_heading('LLM-промпты', level=3)
add_styled_table(
    ['Промпт', 'Вход', 'Выход', 'Когда'],
    [
        ['forecast_trend', 'title, text, bigrams, keyso, trends', 'recommendation (publish_now/later/skip), trend_score (0–100)', 'Обогащение'],
        ['rewrite_news', 'title, text, style, language', 'title, text, seo_title, seo_description, tags[]', 'Рерайт'],
        ['translate', 'title, target_language', 'translated_title', 'Кнопка «Перевод»'],
        ['merge_analysis', 'news1, news2', 'merged_title, merged_text', 'Объединение дубликатов'],
        ['daily_digest', 'top_20_news, style', 'title, text, news_count', 'Авто-дайджест (23:00)'],
    ],
    col_widths=[3, 4, 4, 3]
)

doc.add_heading('18.4. Google Sheets', level=2)
add_styled_table(
    ['Параметр', 'Значение'],
    [
        ['Библиотека', 'gspread'],
        ['Аутентификация', 'Service Account JSON (base64 в env)'],
        ['Rate limit', '1.2с между записями (50 req/min safe)'],
        ['Кэш клиента', 'TTL 3000с (токены обновляются)'],
        ['Кэш листов', '5 мин'],
    ],
    col_widths=[4, 10.5]
)

doc.add_heading('Вкладки Google Sheets', level=3)
add_styled_table(
    ['Вкладка', 'Содержимое', 'Кто пишет'],
    [
        ['Лист1', 'Основной экспорт (все обработанные)', 'Кнопка «В Sheets» из Редакции/Обогащённых'],
        ['Ready', 'Рерайтнутые статьи (полный экспорт)', 'Полный автомат, авто-рерайт при publish_now'],
        ['NotReady', 'Для ручного обзора (без LLM)', 'Пайплайн «Без LLM»'],
    ],
    col_widths=[2.5, 6, 6]
)

doc.add_heading('Колонки Ready (16 колонок)', level=3)
add_styled_table(
    ['Колонка', 'Данные'],
    [
        ['A', 'parsed_at — дата парсинга'],
        ['B', 'source — источник'],
        ['C', 'original_title — оригинальный заголовок'],
        ['D', 'rewrite_title — заголовок рерайта'],
        ['E', 'rewrite_text — текст рерайта (до 5000 симв.)'],
        ['F', 'seo_title — SEO заголовок'],
        ['G', 'seo_description — мета-описание'],
        ['H', 'tags — теги через запятую'],
        ['I', 'total_score — внутренний скор'],
        ['J', 'viral_score — виральный балл'],
        ['K', 'viral_triggers — список триггеров'],
        ['L', 'keyso_freq — частота Keys.so'],
        ['M', 'trends_RU — Google Trends RU'],
        ['N', 'llm_recommendation — рекомендация LLM'],
        ['O', 'llm_trend_forecast — прогноз тренда'],
        ['P', 'url — ссылка на оригинал'],
    ],
    col_widths=[1.5, 13]
)

doc.add_page_break()

# ══════════════════════════════════════════════
# 19. SEO ANALYSIS
# ══════════════════════════════════════════════

doc.add_heading('19. SEO-анализ статей', level=1)
doc.add_paragraph(
    'Кнопка «SEO-проверка» в редакторе статей анализирует текст по 9 критериям '
    'и выдаёт итоговый балл (0–100) с рекомендациями.'
)

add_styled_table(
    ['#', 'Критерий', 'Макс. баллов', 'Условие прохождения'],
    [
        ['1', 'Длина заголовка', '10', '30–60 символов (оптимально для Google)'],
        ['2', 'SEO Title', '15', 'Заполнен, 30–60 символов'],
        ['3', 'Meta Description', '15', 'Заполнено, 120–155 символов'],
        ['4', 'Количество слов', '15', '≥ 300 слов (минимум для SEO)'],
        ['5', 'Плотность ключевиков', '10', '1–3% (ключевые слова в тексте)'],
        ['6', 'Теги', '10', '≥ 3 тегов'],
        ['7', 'Читаемость', '10', 'Средняя длина предложения < 25 слов'],
        ['8', 'Лид-абзац', '10', 'Первый абзац ≥ 50 символов'],
        ['9', 'Подзаголовки', '5', 'Наличие ## или подзаголовков в тексте'],
    ],
    col_widths=[0.7, 3.5, 2, 8]
)

add_formula('SEO Score', 'SUM(критерий_1..9)', 'Итого: 0–100 баллов.')

doc.add_page_break()

# ══════════════════════════════════════════════
# 20. PUBLICATION SCHEDULER
# ══════════════════════════════════════════════

doc.add_heading('20. Планировщик публикаций', level=1)
doc.add_paragraph(
    'Статьи можно запланировать на отложенную публикацию. '
    'Система проверяет каждую минуту и автоматически публикует статьи, у которых наступило время.'
)

doc.add_heading('Как запланировать', level=2)
add_numbered([
    'Откройте статью в редакторе (вкладка «Контент»)',
    'Нажмите кнопку «Запланировать»',
    'Выберите дату и время через datetime picker',
    'Нажмите «Сохранить расписание»',
    'Статус статьи → scheduled',
    'При наступлении времени: статус → published (автоматически)',
])

doc.add_heading('Отмена расписания', level=2)
doc.add_paragraph('Нажмите «Отменить расписание» в редакторе. Статус вернётся на draft.')

doc.add_page_break()

# ══════════════════════════════════════════════
# 21. BROWSER NOTIFICATIONS
# ══════════════════════════════════════════════

doc.add_heading('21. Уведомления браузера', level=1)
doc.add_paragraph(
    'Система может отправлять push-уведомления через браузер (Web Notifications API). '
    'Уведомления приходят когда вкладка открыта.'
)

doc.add_heading('Включение', level=2)
add_numbered([
    'Нажмите иконку колокольчика (🔔) в шапке панели',
    'Браузер запросит разрешение — нажмите «Разрешить»',
    'Иконка станет активной (зелёной)',
    'Повторный клик — отключает уведомления',
])

doc.add_heading('Когда приходят уведомления', level=2)
add_bullets([
    'Новая новость с финальным скором ≥ 60 появилась во вкладке «Финал»',
    'Несколько новых топ-новостей появились одновременно',
    'Пайплайн завершил обработку (полный автомат или Без LLM)',
])

doc.add_page_break()

# ══════════════════════════════════════════════
# 22. SMART POLLING
# ══════════════════════════════════════════════

doc.add_heading('22. Умный polling и оптимизации', level=1)

doc.add_heading('Visibility API', level=2)
doc.add_paragraph(
    'Система использует браузерный Visibility API для экономии ресурсов. '
    'Когда вкладка скрыта (переключились на другой таб), все автообновления приостанавливаются. '
    'При возвращении — данные обновляются мгновенно.'
)

doc.add_heading('TF-IDF дисковый кэш', level=2)
doc.add_paragraph(
    'Словарь TF-IDF сохраняется на диск в JSON. При перезапуске сервера не нужно '
    'переобучать модель — данные загружаются из кэша.'
)
add_bullets([
    '3-уровневый кэш: память → диск (JSON) → переобучение',
    'Файл: nlp/vocab_cache.json',
    'Переобучение: автоматически при > 100 новых документов',
])

doc.add_heading('Cache-Control заголовки', level=2)
doc.add_paragraph(
    'Все API-ответы и HTML-страницы отдаются с заголовком Cache-Control: no-cache. '
    'Это предотвращает проблемы с устаревшими данными в браузере.'
)

doc.add_page_break()

# ══════════════════════════════════════════════
# 23. PROXY ROTATION
# ══════════════════════════════════════════════

doc.add_heading('23. Прокси-ротация для парсеров', level=1)
doc.add_paragraph(
    'Система поддерживает ротацию прокси-серверов для парсинга. '
    'Полезно для обхода блокировок и распределения нагрузки.'
)

doc.add_heading('Как работает', level=2)
add_styled_table(
    ['Механизм', 'Описание'],
    [
        ['Round-robin', 'Прокси используются по кругу (1→2→3→1→...)'],
        ['Retry с backoff', '3 попытки с экспоненциальной задержкой: 2с, 4с, 8с + jitter'],
        ['Ротация на ошибке', 'При 429/503 или connection error — переключение на другой прокси'],
        ['Circuit Breaker', 'После 5 последовательных ошибок — домен блокируется на 1 час'],
        ['User-Agent ротация', '6 браузерных UA, случайный выбор при каждом запросе'],
    ],
    col_widths=[4, 10.5]
)

doc.add_heading('Настройка прокси', level=2)
doc.add_paragraph('Переменная окружения PROXY_LIST (через запятую):')
add_info_box('PROXY_LIST=http://proxy1:8080,http://proxy2:8080,socks5://proxy3:1080', 'blue')

doc.add_heading('Circuit Breaker', level=2)
add_styled_table(
    ['Параметр', 'Значение'],
    [
        ['Порог ошибок', '5 последовательных сбоев'],
        ['Время блокировки', '1 час'],
        ['Сброс', 'Автоматический после cooldown, или при успешном запросе'],
    ],
    col_widths=[4, 10.5]
)

doc.add_page_break()

# ══════════════════════════════════════════════
# 24. ENTITY SYSTEM
# ══════════════════════════════════════════════

doc.add_heading('24. Система сущностей (NER)', level=1)
doc.add_paragraph(
    'Собственная система распознавания именованных сущностей (Named Entity Recognition), '
    'оптимизированная для игровой индустрии.'
)

doc.add_heading('База сущностей', level=2)
add_styled_table(
    ['Категория', 'Количество', 'Примеры'],
    [
        ['Игры', '80+', 'GTA VI, Elden Ring, Cyberpunk 2077, Zelda: TotK'],
        ['Студии', '30+', 'Nintendo, Rockstar, CD Projekt Red, FromSoftware'],
        ['Платформы', '10+', 'PlayStation 5, Xbox Series X, Nintendo Switch, Steam Deck'],
    ],
    col_widths=[3, 2, 9.5]
)

doc.add_heading('Система тиров', level=2)
add_styled_table(
    ['Тир', 'Описание', 'Бонус к viral', 'Примеры'],
    [
        ['S (AAA)', 'Самые ожидаемые/крупные', '+30', 'GTA VI, Elder Scrolls VI, Nintendo, PlayStation'],
        ['A (Major)', 'Крупные проекты', '+25', 'Cyberpunk 2077, Elden Ring, Ubisoft, EA'],
        ['B (Notable)', 'Заметные проекты', '+15', "Baldur's Gate 3, Hollow Knight, Capcom"],
        ['C (Niche)', 'Нишевые/инди', '+5', 'Малые студии, инди-игры'],
    ],
    col_widths=[2, 3.5, 2.5, 6.5]
)

doc.add_heading('Оптимизации', level=2)
add_bullets([
    'Precompiled regex: короткие ключи (3 символа и меньше) компилируются в regex с word boundaries',
    'LRU-кэш: 256 записей, повторный анализ одного текста мгновенный',
    'Используется в: TF-IDF (entity boost), viral_score (tier bonus), deduplication (entity overlap)',
])

doc.add_page_break()

# ══════════════════════════════════════════════
# 25. DEDUPLICATION
# ══════════════════════════════════════════════

doc.add_heading('25. Дедупликация — как работает', level=1)
doc.add_paragraph('Двухуровневая система обнаружения дубликатов:')

doc.add_heading('Уровень 1: URL-дедупликация (парсинг)', level=2)
doc.add_paragraph(
    'При парсинге вычисляется MD5-хеш URL. Если хеш уже есть в базе — '
    'новость не добавляется. Это предотвращает повторный импорт.'
)
add_formula('id', 'MD5(url)')

doc.add_heading('Уровень 2: Семантическая дедупликация (ревью)', level=2)
doc.add_paragraph(
    'При авто-ревью каждая новость сравнивается с существующими за последние 48 часов:'
)
add_numbered([
    'TF-IDF vectorization: текст каждой новости преобразуется в вектор',
    'Cosine Similarity: вычисляется попарное сходство',
    'Entity Overlap: пересечение обнаруженных сущностей (игры, студии)',
    'Комбинированный порог: если cosine > 0.7 И entity_overlap > 0.5 → дубликат',
])

add_formula('is_duplicate', 'cosine_similarity > 0.7 AND entity_overlap > 0.5')

doc.add_heading('Оптимизация', level=2)
doc.add_paragraph(
    'Entity sets предвычисляются один раз для всего батча (O(n) вместо O(n²)). '
    'Это критично при проверке сотен новостей.'
)

doc.add_page_break()

# ══════════════════════════════════════════════
# 26. FEEDBACK LOOP
# ══════════════════════════════════════════════

doc.add_heading('26. Feedback Loop — обучаемость', level=1)
doc.add_paragraph(
    'Система учится на решениях редактора. Каждое одобрение или отклонение '
    'корректирует будущие оценки аналогичных новостей.'
)

doc.add_heading('Как работает', level=2)
add_numbered([
    'При одобрении/отклонении записывается: source, tags, decision, total_score',
    'При следующем скоринге система проверяет: какие решения принимались для этого источника и тегов',
    'Если источник часто одобряется (>70%) → feedback_bonus до +10',
    'Если источник часто отклоняется (<30%) → feedback_penalty до -10',
    'Бонус/штраф добавляется к total_score',
])

add_formula('feedback_bonus', '(approval_rate - 0.5) × 20', 'Диапазон: -10..+10')

doc.add_heading('Адаптация весов источников', level=2)
doc.add_paragraph(
    'Вес источника (source_weight) также корректируется:\n'
    '• >80% одобрений → +0.2\n'
    '• <30% одобрений → -0.2\n'
    '• Диапазон: 0.5 – 2.0'
)

doc.add_page_break()

# ══════════════════════════════════════════════
# 27. GOOGLE SHEETS
# ══════════════════════════════════════════════

doc.add_heading('27. Google Sheets — экспорт', level=1)

doc.add_heading('3 вкладки в таблице', level=2)
add_styled_table(
    ['Вкладка', 'Когда заполняется', 'Содержимое'],
    [
        ['Лист1', 'Ручной экспорт (кнопка «В Sheets»)', '15 колонок: дата, URL, заголовок, H1, описание, биграммы, Keys.so, Trends, LLM, текст'],
        ['Ready', 'Полный автомат / авто-рерайт при publish_now', '16 колонок: рерайтнутый текст, SEO-данные, скоры, триггеры, Tags'],
        ['NotReady', 'Пайплайн «Без LLM»', 'Базовые данные + internal score для ручного обзора'],
    ],
    col_widths=[2.5, 5, 7]
)

doc.add_heading('Дедупликация в Sheets', level=2)
doc.add_paragraph(
    'Перед записью система проверяет, есть ли URL в таблице (кэш 2 мин). '
    'Если URL уже есть — запись пропускается (sheet_row = -1).'
)

doc.add_heading('Rate Limiting', level=2)
doc.add_paragraph(
    'Google Sheets API ограничен ~60 запросов/мин. '
    'Система выдерживает паузу 1.2с между записями для безопасности.'
)

doc.add_page_break()

# ══════════════════════════════════════════════
# 28. FAQ
# ══════════════════════════════════════════════

doc.add_heading('28. Часто задаваемые вопросы', level=1)

faqs = [
    ('Как часто парсятся новости?',
     'Каждые 5–30 минут в зависимости от источника. IGN — каждые 5 минут, '
     'HTML-парсеры (iXBT, VGTimes) — каждые 30 минут. Всего 15 источников.'),

    ('Сколько стоят API-вызовы?',
     'Парсинг и авто-ревью бесплатны (12 локальных проверок). API вызываются только при полном автомате '
     'или одобрении: Keys.so (~2 запроса), Google Trends (~1 запрос), LLM (~1 запрос), '
     'LLM рерайт (~1 запрос при полном автомате). Итого: ~$0.01–0.05 за новость.'),

    ('Что делает кнопка «Полный автомат»?',
     'Берёт выбранные (или все новые) новости и проводит через 6 стадий: '
     '1) Скоринг (бесплатно) → 2) Фильтр >70 → 3) Обогащение API → '
     '4) Финальный скор → 5) Фильтр >60 → 6) Рерайт → Sheets Ready. '
     'Новости с низким скором пропускаются, экономя API-вызовы.'),

    ('Чем отличается total_score от финального скора?',
     'total_score (внутренний) — результат 12 локальных бесплатных проверок. '
     'Финальный скор — композитный балл: internal×40% + viral×20% + Keys.so×15% + Trends×10% + headline×15%. '
     'Финальный вычисляется только после обогащения API-данными.'),

    ('Почему Keys.so показывает 0?',
     'Keys.so ищет частоту в определённом регионе. Для EN-источников используется US-регион, '
     'для русских — RU. Если freq=0 — тема не имеет поискового трафика в этом регионе.'),

    ('Как работает дедупликация?',
     'Два уровня: 1) MD5-хеш URL при парсинге (точные дубли). '
     '2) TF-IDF cosine similarity + entity overlap при ревью (похожие новости). '
     'Порог: cosine > 0.7 И entity_overlap > 0.5.'),

    ('Как добавить новый источник?',
     '⚙ Настройки → Источники → «Добавить». Укажите: имя, тип (rss/html/json), '
     'URL фида, интервал парсинга. Для HTML нужен CSS-селектор списка новостей.'),

    ('Как изменить модель LLM?',
     '⚙ Настройки → Общие → Модель LLM. Введите ID любой модели из OpenRouter '
     '(gpt-4o-mini, claude-3-haiku, llama-3 и др.).'),

    ('Что делать если вкладка пустая?',
     'Обогащённые: одобрите новости в Редакции. Финал: нужны processed-новости '
     'с рекомендацией publish_now. Контент: отправьте новости на рерайт. '
     'Также попробуйте Ctrl+Shift+R (жёсткая перезагрузка).'),

    ('Как перезапустить ошибочные задачи?',
     'Вкладка «Очередь» → отметить задачи с ошибкой → «Повторить выбранные». '
     'Или ⚙ Настройки → Инструменты → «Повторить все ошибки».'),

    ('Как работает SEO-анализ?',
     'Кнопка «SEO-проверка» в редакторе анализирует статью по 9 критериям: '
     'заголовок, SEO title, meta description, кол-во слов, плотность ключевиков, '
     'теги, читаемость, лид, подзаголовки. Максимум 100 баллов.'),

    ('Как запланировать публикацию?',
     'В редакторе статьи нажмите «Запланировать» → выберите дату/время. '
     'Статус → scheduled. Система проверяет каждую минуту и публикует автоматически.'),

    ('Что такое Circuit Breaker?',
     'Защита от перегрузки. Если домен (например, ign.com) 5 раз подряд не отвечает, '
     'он блокируется на 1 час. После остывания — запросы возобновляются.'),

    ('Что такое feedback loop?',
     'Система учится на решениях: часто одобряемые источники/теги получают бонус до +10, '
     'часто отклоняемые — штраф до -10. Это автоматически улучшает скоринг.'),
]

for q, a in faqs:
    p = doc.add_paragraph()
    run = p.add_run(f'Q: {q}')
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
    p = doc.add_paragraph()
    run = p.add_run(f'A: {a}')
    run.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(10)

doc.add_page_break()

# ══════════════════════════════════════════════
# FOOTER
# ══════════════════════════════════════════════

for _ in range(4):
    doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('IgroNews v3.0')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Полное руководство пользователя')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('28 разделов  •  Все формулы  •  Все функции')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Документация актуальна на март 2026 года')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Railway  •  github.com/staurus86/IgroNews')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ── Save ──
out_path = os.path.join(os.path.dirname(__file__), 'IgroNews_Manual.docx')
doc.save(out_path)
print(f'Saved: {out_path}')
print(f'Size: {os.path.getsize(out_path) / 1024:.1f} KB')
